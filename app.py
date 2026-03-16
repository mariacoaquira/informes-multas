import ssl
#ssl._create_default_https_context = ssl._create_unverified_context
import streamlit as st
import io
import locale
from babel.dates import format_date 
from datetime import datetime, date
import pandas as pd
import importlib
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH                  
from docxcompose.composer import Composer
from docxtpl import DocxTemplate, RichText
from num2words import num2words
import requests # <--- AÑADIR
import traceback # <--- AÑADIR
from jinja2 import Environment
import copy #

# Importaciones de nuestros módulos
from sheets import (conectar_gsheet, cargar_hoja_a_df, 
                    get_person_details_by_base_name, descargar_archivo_drive,
                    calcular_beneficio_ilicito, calcular_multa,
                    actualizar_hoja_con_df,
                    guardar_datos_caso, cargar_datos_caso) # <-- AÑADIR ESTAS DOS
from funciones import (combinar_con_composer, create_table_subdoc, 
                     create_main_table_subdoc, create_summary_table_subdoc, create_personal_table_subdoc, NumberingManager,
                     texto_con_numero, format_decimal_dinamico, get_initials_from_name, formatear_lista_hechos, redondeo_excel, create_capacitacion_table_subdoc)

def preparar_datos_para_json(obj):
    """
    Limpia recursivamente un objeto para que sea compatible con JSON.
    Convierte fechas a string y elimina objetos binarios o complejos.
    """
    if isinstance(obj, dict):
        # Eliminamos claves conocidas que causan errores de guardado (objetos binarios/clases)
        claves_invalidas = {'resultados', 'anexos_ce', 'tabla_detalle_personal', 'acronyms', 'numbering', 'doc_tpl'}
        return {k: preparar_datos_para_json(v) for k, v in obj.items() if k not in claves_invalidas}
    elif isinstance(obj, list):
        return [preparar_datos_para_json(i) for i in obj]
    elif isinstance(obj, (datetime, date)):
        return obj.isoformat() # "2024-05-20"
    elif isinstance(obj, (RichText, io.BytesIO)):
        return None # No se pueden guardar estos objetos
    return obj


def actualizar_y_recalcular_prorrateos(cliente_gspread, NOMBRE_GSHEET_MAESTRO):
    """
    Analiza todos los hechos calculados, detecta colisiones de capacitación por año,
    actualiza los factores de prorrateo y vuelve a procesar los resultados para que
    la interfaz muestre los montos actualizados (1/N).
    """
    if 'imputaciones_data' not in st.session_state:
        return

    # 1. Contar cuántos extremos hay por cada año (solo para los que usan capacitación)
    conteo_por_anio = {}
    for datos in st.session_state.imputaciones_data:
        id_inf = datos.get('id_infraccion', '')
        for extremo in datos.get('extremos', []):
            tipo = extremo.get('tipo_extremo', '') or extremo.get('tipo_presentacion', '')
            # Lógica: INF003 siempre cuenta, otros solo si es "No presentó"
            if 'INF003' in id_inf or tipo in ["No presentó", "No remitió", "No remitió información / Remitió incompleto"]:
                fecha = extremo.get('fecha_incumplimiento') or extremo.get('fecha_incumplimiento_extremo')
                if fecha:
                    conteo_por_anio[fecha.year] = conteo_por_anio.get(fecha.year, 0) + 1

    # 2. Actualizar factores y disparar recálculo de los hechos afectados
    for i, datos in enumerate(st.session_state.imputaciones_data):
        id_inf = datos.get('id_infraccion')
        if not id_inf or 'resultados' not in datos:
            continue

        mapa_nuevo = {}
        for extremo in datos.get('extremos', []):
            fecha = extremo.get('fecha_incumplimiento') or extremo.get('fecha_incumplimiento_extremo')
            if fecha and fecha.year in conteo_por_anio:
                mapa_nuevo[fecha.year] = 1.0 / conteo_por_anio[fecha.year]
        
        # Si el factor cambió, recalculamos este hecho individual
        if datos.get('mapa_factores_prorrateo') != mapa_nuevo:
            datos['mapa_factores_prorrateo'] = mapa_nuevo
            
            # --- RE-EJECUCIÓN DEL PROCESAMIENTO ---
            try:
                modulo = importlib.import_module(f"infracciones.{id_inf}")
                # Reconstruir datos comunes (similar al botón de calcular del Paso 3)
                df_tip = st.session_state.datos_calculo['df_tipificacion']
                id_plantilla = df_tip[df_tip['ID_Infraccion'] == id_inf].iloc[0]['ID_Plantilla_BI']
                
                datos_comunes = {
                    **st.session_state.datos_calculo,
                    'datos_hecho_completos': datos,
                    'fecha_emision_informe': st.session_state.get('fecha_emision_informe', date.today()),
                    'id_infraccion': id_inf,
                    'rubro': st.session_state.rubro_seleccionado,
                    'id_rubro_seleccionado': st.session_state.get('id_rubro_seleccionado'),
                    'numero_hecho_actual': i + 1,
                    'doc_tpl': DocxTemplate(descargar_archivo_drive(id_plantilla)),
                    'context_data': st.session_state.get('context_data', {}),
                    'acronym_manager': st.session_state.context_data.get('acronyms')
                }
                
                nuevos_res = modulo.procesar_infraccion(datos_comunes, datos)
                if not nuevos_res.get('error'):
                    st.session_state.imputaciones_data[i]['resultados'] = nuevos_res
            except Exception as e:
                st.error(f"Error al actualizar prorrateo del Hecho {i+1}: {e}")

# --- CONSTANTES: FACTORES DE GRADUACIÓN (f1 a f7) ---
FACTORES_GRADUACION = {
    "f1": {
        "titulo": "GRAVEDAD DEL DAÑO AL AMBIENTE",
        "criterios": {
            "1.1 Componentes Ambientales": {
                "No determinado / No aplica": 0.0,
                "El daño afecta a un (01) componente ambiental.": 0.10,
                "El daño afecta a dos (02) componentes ambientales.": 0.20,
                "El daño afecta a tres (03) componentes ambientales.": 0.30,
                "El daño afecta a cuatro (04) componentes ambientales.": 0.40,
                "El daño afecta a cinco (05) componentes ambientales.": 0.50
            },
            "1.2 Incidencia en la calidad": {
                "No determinado / No aplica": 0.0,
                "Impacto mínimo.": 0.06,
                "Impacto regular.": 0.12,
                "Impacto alto.": 0.18,
                "Impacto total.": 0.24
            },
            "1.3 Extensión geográfica": {
                "No determinado / No aplica": 0.0,
                "El impacto está localizado en el área de influencia directa.": 0.10,
                "El impacto está localizado en el área de influencia indirecta.": 0.20
            },
            "1.4 Reversibilidad/Recuperabilidad": {
                "No determinado / No aplica": 0.0,
                "Reversible en el corto plazo.": 0.06,
                "Recuperable en el corto plazo.": 0.12,
                "Recuperable en el mediano plazo.": 0.18,
                "Recuperable en el largo plazo o irrecuperable.": 0.24
            },
            "1.5 Afectación recursos/áreas protegidas": {
                "No existe afectación o esta es indeterminable...": 0.0,
                "El impacto se ha producido en un área natural protegida...": 0.40
            },
             "1.6 Afectación comunidades": {
                "No afecta a comunidades nativas o campesinas.": 0.0,
                "Afecta a una comunidad nativa o campesina.": 0.15,
                "Afecta a más de una comunidad nativa o campesina.": 0.30
            },
            "1.7 Afectación salud": {
                "No afecta a la salud de las personas...": 0.0,
                "Afecta la salud de las personas.": 0.60
            }
        }
    },
    "f2": {
        "titulo": "PERJUICIO ECONÓMICO CAUSADO (Pobreza)",
        "criterios": {
            "Incidencia de pobreza total": {
                "No determinado / No aplica": 0.0,
                "Incidencia de pobreza total hasta 19,6%.": 0.04,
                "Incidencia de pobreza total mayor a 19,6% hasta 39,1%.": 0.08,
                "Incidencia de pobreza total mayor a 39,1% hasta 58,7%.": 0.12,
                "Incidencia de pobreza total mayor a 58,7% hasta 78,2%.": 0.16,
                "Incidencia de pobreza total mayor a 78,2%.": 0.20
            }
        }
    },
    "f3": {
        "titulo": "ASPECTOS AMBIENTALES O FUENTES DE CONTAMINACIÓN",
        "criterios": {
             "Cantidad de aspectos": {
                "No determinado / No aplica": 0.0,
                "El impacto involucra un (01) aspecto ambiental...": 0.06,
                "El impacto involucra dos (02) aspectos ambientales...": 0.12,
                "El impacto involucra tres (03) aspectos ambientales...": 0.18,
                "El impacto involucra cuatro (04) aspectos ambientales...": 0.24,
                "El impacto involucra cinco (05) aspectos ambientales...": 0.30
             }
        }
    },
    "f4": {
        "titulo": "REINCIDENCIA",
        "criterios": {
            "Reincidencia": {
                "No existe reincidencia": 0.0,
                "Existe reincidencia (comisión de misma infracción en 1 año)": 0.20
            }
        }
    },
    "f5": {
        "titulo": "CORRECCIÓN DE LA CONDUCTA INFRACTORA (Atenuante)",
        "criterios": {
             "Subsanación/Corrección": {
                "No corrige / No aplica": 0.0,
                "Subsana voluntariamente antes del inicio del PAS (Eximente)": "Eximente",
                "Corrige (leve) a requerimiento, antes del inicio del PAS (Eximente)": "Eximente",
                "Corrige (trascendente) a requerimiento, antes del inicio del PAS (-40%)": -0.40,
                "Corrige luego del inicio del PAS, antes de resolución (-20%)": -0.20
             }
        }
    },
    "f6": {
        "titulo": "ADOPCIÓN DE MEDIDAS PARA REVERTIR CONSECUENCIAS",
        "criterios": {
            "Medidas adoptadas": {
                "No ejecutó ninguna medida (+30%)": 0.30,
                "Ejecutó medidas tardías (+20%)": 0.20,
                "Ejecutó medidas parciales (+10%)": 0.10,
                "No aplica / Neutro": 0.0,
                "Ejecutó medidas necesarias e inmediatas (-10%)": -0.10
            }
        }
    },
    "f7": {
        "titulo": "INTENCIONALIDAD",
        "criterios": {
             "Intencionalidad": {
                 "No se acredita intencionalidad": 0.0,
                 "Se acredita intencionalidad (+72%)": 0.72
             }
        }
    }
}

#st.write(st.session_state)

# --- INICIALIZACIÓN DE LA APLICACIÓN ---
#st.set_page_config(layout="wide", page_title="Asistente de Multas")
st.title("🤖 Asistente para la elaboración de informes de multa")

# --- INICIO: Lógica de Actualización BCRP ---
def actualizar_datos_bcrp(cliente_gspread):
    """
    Función principal para conectarse a la API del BCRP,
    descargar datos y actualizar la hoja de Google Sheets.
    """
    
    # 1. --- ¡DEBES REEMPLAZAR ESTOS CÓDIGOS! ---
    # Búscalos en la web de BCRPData. (Ej: 'PN01288PM', 'PN01270PM')
    CODIGO_IPC_MENSUAL = "PN38705PM" 
    CODIGO_TC_MENSUAL = "PN01210PM"

    # Nombres de tus hojas y columnas
    NOMBRE_ARCHIVO = "Base de datos"
    NOMBRE_HOJA = "Indices_BCRP"
    COLUMNAS_API = [CODIGO_IPC_MENSUAL, CODIGO_TC_MENSUAL]
    COLUMNAS_HOJA = ['IPC_Mensual', 'TC_Mensual']
    
    # Unir códigos para la API [cite: 74]
    series_a_pedir = "-".join(COLUMNAS_API)
    
    with st.spinner("Actualizando datos del BCRP..."):
        try:
            # 2. Cargar datos existentes de Google Sheets
            df_existente = cargar_hoja_a_df(cliente_gspread, NOMBRE_ARCHIVO, NOMBRE_HOJA)
            if df_existente is None:
                st.error("No se pudo cargar la hoja 'Indices_BCRP'. Abortando.")
                return

            df_existente['Indice_Mes'] = pd.to_datetime(df_existente['Indice_Mes'], dayfirst=True, errors='coerce')
            ultima_fecha = df_existente['Indice_Mes'].max()
            
            # 3. Determinar el rango de fechas para la API
            fecha_hoy_str = pd.to_datetime('today').strftime('%Y-%m')
            
            if pd.isna(ultima_fecha):
                # Si la hoja está vacía, pedimos los últimos 5 años
                periodo_inicial_str = (pd.to_datetime('today') - pd.DateOffset(years=5)).strftime('%Y-%m')
            else:
                # Pedimos desde el mes SIGUIENTE al último que tenemos
                periodo_inicial_str = (ultima_fecha + pd.DateOffset(months=1)).strftime('%Y-%m')

            if periodo_inicial_str > fecha_hoy_str:
                st.success("¡Datos actualizados! No se encontraron nuevos registros.")
                st.cache_data.clear() # Limpiar caché por si acaso
                return

            # 4. Construir y llamar a la API del BCRP [cite: 71, 97]
            url_api = f"https://estadisticas.bcrp.gob.pe/estadisticas/series/api/{series_a_pedir}/json/{periodo_inicial_str}/{fecha_hoy_str}/"
            
            response = requests.get(url_api)
            response.raise_for_status() # Lanza un error si la petición falla
            
            data = response.json()
            
            # 5. Procesar y parsear la respuesta JSON
            nuevos_registros = []
            meses_map = {'Ene': 1, 'Feb': 2, 'Mar': 3, 'Abr': 4, 'May': 5, 'Jun': 6,
                         'Jul': 7, 'Ago': 8, 'Set': 9, 'Sep': 9, 'Oct': 10, 'Nov': 11, 'Dic': 12}
            
            # --- INICIO: Lógica robusta de mapeo de API (v2 - por NOMBRE) ---
            # 1. Definir palabras clave ÚNICAS para cada serie
            MAP_PALABRA_CLAVE_A_COLUMNA = {
                'IPC': 'IPC_Mensual', # Si 'name' contiene "IPC"
                'Tipo de cambio': 'TC_Mensual' # Si 'name' contiene "Tipo de cambio"
            }
            
            # 2. Leer la configuración de la API para saber el orden real de los valores
            map_index_a_columna = {}
            api_series_config = data.get('config', {}).get('series', [])
            
            for index, series_info in enumerate(api_series_config):
                nombre_api = series_info.get('name', '') # <-- OBTENER EL NOMBRE
                
                # Buscar la palabra clave en el nombre
                for palabra_clave, nombre_columna_hoja in MAP_PALABRA_CLAVE_A_COLUMNA.items():
                    if palabra_clave in nombre_api:
                        map_index_a_columna[index] = nombre_columna_hoja # Ej: {0: 'TC_Mensual', 1: 'IPC_Mensual'}
                        break # Salir del bucle interno una vez encontrada
            # --- FIN: Lógica robusta de mapeo (v2) ---

            # --- INICIO: Verificación de Mapeo ---
            if not map_index_a_columna or len(map_index_a_columna) < len(COLUMNAS_API):
                st.error("Error: La respuesta de la API del BCRP no incluyó una configuración de series válida o faltan códigos.")
                st.warning(f"Mapeo de columnas resultante: {map_index_a_columna}")
                st.warning(f"Respuesta de 'config' de la API: {data.get('config', 'No encontrada')}")
                return # Detener la ejecución de la función
            # --- FIN: Verificación de Mapeo ---

            for periodo in data.get('periods', []):
                try:
                    # Parsear la fecha "Ene.2024"
                    mes_str, anio_str = periodo['name'].split('.')
                    mes_num = meses_map[mes_str.capitalize()]
                    fecha_registro = pd.to_datetime(f"{anio_str}-{mes_num}-01")
                    
                    # Crear el diccionario de datos
                    registro = {'Indice_Mes': fecha_registro}
                    valores = periodo['values'] # Ej: ["3.5432", "115.5868"]
                    
                    # --- INICIO: Asignación corregida ---
                    # Asignar valores usando el mapeo que creamos
                    for index, col_nombre in map_index_a_columna.items():
                        try:
                            # Convertir a float, manejando comas si las hubiera
                            valor_limpio = valores[index].replace(',', '.') if isinstance(valores[index], str) else valores[index]
                            registro[col_nombre] = float(valor_limpio)
                        except (IndexError, ValueError, TypeError):
                            registro[col_nombre] = None # Poner nulo si el valor no es numérico
                    # --- FIN: Asignación corregida ---
                    
                    nuevos_registros.append(registro)

                except Exception as e_parse:
                    st.warning(f"No se pudo procesar el periodo '{periodo.get('name', 'N/A')}'. Error: {e_parse}")
            
            if not nuevos_registros:
                st.success("¡Datos actualizados! No se encontraron nuevos registros.")
                st.cache_data.clear()
                return
                
            # 6. Convertir a DataFrame y filtrar duplicados (por si acaso)
            df_nuevos = pd.DataFrame(nuevos_registros)
            df_nuevos = df_nuevos.dropna(subset=COLUMNAS_HOJA) # Eliminar filas donde falten datos
            df_nuevos_filtrados = df_nuevos[~df_nuevos['Indice_Mes'].isin(df_existente['Indice_Mes'])]

            if df_nuevos_filtrados.empty:
                st.success("¡Datos actualizados! No se encontraron nuevos registros.")
                st.cache_data.clear()
                return

            # 7. Enviar los datos nuevos a Google Sheets
            # (Solo enviamos las 3 columnas que nos importan)
            df_final_a_subir = df_nuevos_filtrados[['Indice_Mes', 'IPC_Mensual', 'TC_Mensual']]
            
            num_filas_anadidas = actualizar_hoja_con_df(cliente_gspread, NOMBRE_ARCHIVO, NOMBRE_HOJA, df_final_a_subir)
            
            if num_filas_anadidas > 0:
                st.success(f"¡Éxito! Se añadieron {num_filas_anadidas} nuevos registros a '{NOMBRE_HOJA}'.")
                st.cache_data.clear() # MUY IMPORTANTE: Limpia el caché para que el resto de la app lea los datos nuevos.
            else:
                st.error("No se pudo añadir los nuevos registros a la hoja de cálculo.")

        except requests.exceptions.HTTPError as e_http:
            st.error(f"Error al contactar la API del BCRP: {e_http}")
            st.error(f"URL consultada: {url_api}")
        except Exception as e:
            st.error(f"Ocurrió un error inesperado durante la actualización: {e}")
            traceback.print_exc()

# --- FIN: Lógica de Actualización BCRP ---

# --- INICIO: Botón de Actualización BCRP ---
if st.button("Sincronizar datos del BCRP"):
    # Esta función la definiremos a continuación
    actualizar_datos_bcrp(conectar_gsheet()) 
# --- FIN: Botón ---

if 'app_inicializado' not in st.session_state:
    st.session_state.clear()
    st.session_state.app_inicializado = True

cliente_gspread = conectar_gsheet()
NOMBRE_GSHEET_MAESTRO = "Base de datos"
NOMBRE_GSHEET_ASIGNACIONES = "Base de asignaciones de multas"

# --- CUERPO DE LA APLICACIÓN ---
if cliente_gspread:
    # --- PASO 1: BÚSQUEDA DE EXPEDIENTE ---
    st.header("Paso 1: Búsqueda del Expediente")
    col1, col2 = st.columns([1, 2])
    with col1:
        try:
            locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
        except locale.Error:
            locale.setlocale(locale.LC_TIME, '')
        # En app.py, dentro del Paso 1
        hojas_disponibles = [
            format_date(datetime.now() - pd.DateOffset(months=i), "MMMM yyyy", locale='es').capitalize().replace(
                "Septiembre", "Setiembre") for i in range(3)]
        mes_seleccionado = st.selectbox("Selecciona el mes de la asignación:", options=hojas_disponibles)
    with col2:
        df_asignaciones = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_ASIGNACIONES, mes_seleccionado)
        if df_asignaciones is not None:
            num_expediente_input = st.text_input("Ingresa el N° de Expediente:", placeholder="Ej: 1234-2023 o 1234-2023-OEFA/DFAI/PAS")
            
            if st.button("Buscar Expediente", type="primary"):
                # --- LIMPIEZA PROFUNDA DE SESIÓN ---
                claves_a_limpiar = [
                    'info_expediente', 'imputaciones_data', 'context_data', 
                    'num_expediente_formateado', 'rubro_seleccionado', 
                    'id_rubro_seleccionado', 'id_subdireccion_seleccionada',
                    'confiscatoriedad', 'numero_rsd_base', 'fecha_rsd',
                    'numero_ifi', 'fecha_ifi', 'num_informe_multa_ifi',
                    'monto_multa_ifi', 'num_imputaciones_ifi'
                ]
                for clave in claves_a_limpiar:
                    if clave in st.session_state:
                        del st.session_state[clave]
                st.rerun() # Forzamos recarga para asegurar limpieza
            
            if num_expediente_input:
                num_formateado = ""
                if "OEFA" in num_expediente_input.upper():
                    num_formateado = num_expediente_input
                elif "-" in num_expediente_input:
                    num_formateado = f"{num_expediente_input}-OEFA/DFAI/PAS"

                if num_formateado:
                    resultado = df_asignaciones[df_asignaciones['EXPEDIENTE'] == num_formateado]
                    if not resultado.empty:
                        # --- CORRECCIÓN: Actualizar si el número es diferente al actual ---
                        if st.session_state.get('num_expediente_formateado') != num_formateado:
                            st.session_state.num_expediente_formateado = num_formateado
                            st.session_state.info_expediente = resultado.iloc[0].to_dict()
                            
                            # Limpiar datos de hechos anteriores al cambiar de expediente
                            if 'imputaciones_data' in st.session_state:
                                del st.session_state['imputaciones_data']
                            
                            st.success(f"¡Expediente '{num_formateado}' cargado correctamente!")
                        # --- BLOQUE DE CARGA OPTIMIZADO ---
                        if st.button("📥 Cargar Avance Guardado"):
                            expediente_a_cargar = st.session_state.num_expediente_formateado
                            with st.spinner("Buscando avance..."):
                                datos_cargados, mensaje = cargar_datos_caso(cliente_gspread, expediente_a_cargar)
                            
                            if datos_cargados:
                                # Función interna para reconstruir fechas
                                def restaurar_fechas(obj):
                                    if isinstance(obj, dict):
                                        for k, v in obj.items():
                                            if isinstance(v, str) and len(v) == 10 and v.count('-') == 2:
                                                try:
                                                    obj[k] = date.fromisoformat(v)
                                                except: pass
                                            else: restaurar_fechas(v)
                                    elif isinstance(obj, list):
                                        for i in obj: restaurar_fechas(i)

                                restaurar_fechas(datos_cargados)
                                
                                # Inyectar en el session_state
                                for key, value in datos_cargados.items():
                                    st.session_state[key] = value
                                
                                st.success("Datos cargados. Los cálculos se actualizarán al presionar 'Generar Informe'.")
                                st.rerun()
                            else:
                                st.warning(mensaje)

                        
                        # --- CORRECCIÓN CLAVE ---
                        # Solo inicializamos la lista de hechos si no existe previamente
                        if 'imputaciones_data' not in st.session_state:
                            num_imputaciones = int(st.session_state.info_expediente.get('IMPUTACIONES', 1))
                            st.session_state.imputaciones_data = [{} for _ in range(num_imputaciones)]
                    else:
                        st.error(f"No se encontró el expediente '{num_expediente_input}'.")
                else:
                    st.warning("Ingresa un número de expediente en un formato válido.")
    st.divider()

# Reemplaza todo desde aquí en tu app.py

    # --- PASO 2 Y LÓGICA SUBSIGUIENTE ---
    if st.session_state.get('info_expediente'):
        st.header("Paso 2: Detalles del Expediente")
        info_caso = st.session_state.info_expediente

        # --- Subsección: Datos del Expediente ---
        st.subheader("Datos del Expediente")
        nombre_completo_analista = "No encontrado"
        df_analistas = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Analistas")
        if df_analistas is not None:
            nombre_base_analista = info_caso.get('ANALISTA ECONÓMICO')
            if nombre_base_analista:
                analista_encontrado = df_analistas[df_analistas['Nombre_Base_Analista'] == nombre_base_analista]
                if not analista_encontrado.empty:
                    nombre_completo_analista = analista_encontrado.iloc[0]['Nombre_Analista']

        col_info1, col_info2 = st.columns(2)
        with col_info1:
            st.text_input("Nombre o Razón Social del administrado", value=info_caso.get('ADMINISTRADO'), disabled=True)
            st.text_input("Producto", value=info_caso.get('PRODUCTO'), disabled=True)

        with col_info2:
            st.text_input("Analista Económico", value=nombre_completo_analista, disabled=True)
            st.text_input("Sector", value=info_caso.get('SECTOR'), disabled=True)

            df_sector_subdireccion = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Sector_Subdireccion")
            if df_sector_subdireccion is not None and 'ID_Rubro' in df_sector_subdireccion.columns:
                sector_del_caso = info_caso.get('SECTOR')
                if sector_del_caso:
                    rubros_filtrados_df = df_sector_subdireccion[df_sector_subdireccion['Sector_Base'] == sector_del_caso]
                    if not rubros_filtrados_df.empty:
                        lista_rubros = rubros_filtrados_df['Sector_Rubro'].tolist()
                        index_seleccionado = None
                        rubro_guardado = st.session_state.get('rubro_seleccionado')
                        if rubro_guardado and rubro_guardado in lista_rubros:
                            index_seleccionado = lista_rubros.index(rubro_guardado)
                        nombre_rubro_seleccionado = st.selectbox("Elige el subsector/rubro", options=lista_rubros, index=index_seleccionado, placeholder="Selecciona una opción...")
                        if nombre_rubro_seleccionado:
                            st.session_state.rubro_seleccionado = nombre_rubro_seleccionado
                            fila_rubro = rubros_filtrados_df[rubros_filtrados_df['Sector_Rubro'] == nombre_rubro_seleccionado]
                            if not fila_rubro.empty:
                                id_rubro = fila_rubro.iloc[0]['ID_Rubro']
                                id_subdireccion = fila_rubro.iloc[0]['ID_Subdireccion'] # <--- AÑADIR
                                st.session_state.id_rubro_seleccionado = id_rubro
                                st.session_state.id_subdireccion_seleccionada = id_subdireccion # <--- AÑADIR
                    else:
                        st.warning(f"No hay rubros para el sector '{sector_del_caso}'.")
            else:
                st.error("No se pudo cargar la hoja 'Sector_Subdireccion'.")

        st.subheader("Fecha del Informe")

        # Inicializamos el valor en el estado solo si no existe
        if 'fecha_emision_informe' not in st.session_state:
            st.session_state['fecha_emision_informe'] = date.today()

        # El widget ahora solo usa el 'key', Streamlit manejará el valor automáticamente
        fecha_emision_informe = st.date_input(
            "Selecciona la fecha de emisión del informe:",
            key='fecha_emision_informe',
            format="DD/MM/YYYY"
        )

        # --- Subsección: Resolución Subdirectoral (RSD) ---
        producto_caso = st.session_state.info_expediente.get('PRODUCTO', '')

        # Solo mostramos estos campos si el producto requiere una RSD de base (IFI y RD)
        if producto_caso != 'COERCITIVA':
            st.subheader("Resolución Subdirectoral (RSD)")
            col_rsd1, col_rsd2 = st.columns([1, 1])
            
            with col_rsd1:
                if 'numero_rsd_base' not in st.session_state:
                    st.session_state['numero_rsd_base'] = ''

                st.text_input(
                    "N° de RSD:", 
                    key='numero_rsd_base',
                    placeholder="Ej. 00245-2025"
                )

                numero_rsd_base = st.session_state.get('numero_rsd_base', '')
                numero_rsd_completo = ""
                if numero_rsd_base:
                    suffix_sub = st.session_state.get('id_subdireccion_seleccionada', 'ERROR_SUB')
                    numero_rsd_completo = f"{numero_rsd_base}-OEFA/DFAI-{suffix_sub}"
                
                if numero_rsd_completo:
                    st.info(f"**RSD:** {numero_rsd_completo}")
                st.session_state.numero_rsd = numero_rsd_completo 

            with col_rsd2:
                if 'fecha_rsd' not in st.session_state:
                    st.session_state['fecha_rsd'] = date.today()

                st.date_input(
                    "Fecha de notificación de la RSD:", 
                    key='fecha_rsd',
                    format="DD/MM/YYYY"
                )

        # --- INICIO: SECCIÓN EXCLUSIVA PARA RD (IFI + INFORME MULTA) ---
        producto_caso = st.session_state.info_expediente.get('PRODUCTO', '')
        
        # --- NUEVO ORDEN: Primero capturamos los inputs del IFI ---
        if producto_caso == 'RD':
            st.divider()
            st.subheader("Informe Final de Instrucción (IFI)")
            
            col_ifi1, col_ifi2 = st.columns(2)
            with col_ifi1:
                if 'numero_ifi_base' not in st.session_state: 
                    st.session_state.numero_ifi_base = ''
                
                st.text_input("N° del IFI:", key="numero_ifi_base", placeholder="Ej: 00123-2023")
                
                # Lógica de autocompletado
                ifi_base = st.session_state.get('numero_ifi_base', '')
                ifi_completo = ""
                if ifi_base:
                    suffix_sub = st.session_state.get('id_subdireccion_seleccionada', 'SECTOR')
                    ifi_completo = f"{ifi_base}-OEFA/DFAI-{suffix_sub}"
                
                if ifi_completo:
                    st.info(f"**IFI:** {ifi_completo}")
                
                # Guardamos el resultado final para que lo use el informe
                st.session_state.numero_ifi = ifi_completo
            with col_ifi2:
                if 'fecha_ifi' not in st.session_state: st.session_state.fecha_ifi = date.today()
                st.date_input("Fecha de notificación del IFI:", key="fecha_ifi", format="DD/MM/YYYY")
            
            with st.container(border=True):
                st.markdown("###### Datos del Informe de Multa (IFI)")
                col_im1, col_im2, col_im3 = st.columns(3)
                with col_im1:
                    if 'num_informe_multa_ifi_base' not in st.session_state: 
                        st.session_state.num_informe_multa_ifi_base = ''
                    
                    st.text_input("N° de Informe de Multa:", key="num_informe_multa_ifi_base", placeholder="Ej: 00045-2024")
                    
                    # Lógica de autocompletado
                    im_base = st.session_state.get('num_informe_multa_ifi_base', '')
                    im_completo = ""
                    if im_base:
                        im_completo = f"{im_base}-OEFA/DFAI-SSAG"
                    
                    if im_completo:
                        st.info(f"**Inf. Multa:** {im_completo}")
                    
                    # Guardamos el resultado final para que lo use el informe
                    st.session_state.num_informe_multa_ifi = im_completo
                with col_im2:
                    if 'monto_multa_ifi' not in st.session_state: st.session_state.monto_multa_ifi = 0.0
                    st.number_input("Monto Total Propuesto (UIT):", key="monto_multa_ifi", format="%.3f")
                with col_im3:
                    if 'num_imputaciones_ifi' not in st.session_state: st.session_state.num_imputaciones_ifi = 1
                    st.number_input("Nº Imputaciones:", key="num_imputaciones_ifi", min_value=1)

        # --- RE-CALCULAR resolucion_ok AQUÍ ---
        resolucion_ok = False
        if producto_caso == 'RD':
            if (st.session_state.get('numero_rsd') and st.session_state.get('fecha_rsd') and 
                st.session_state.get('numero_ifi') and st.session_state.get('fecha_ifi')):
                resolucion_ok = True
        elif producto_caso == 'COERCITIVA':
            resolucion_ok = True
        else: # IFI
            if st.session_state.get('numero_rsd') and st.session_state.get('fecha_rsd'):
                resolucion_ok = True

        resolucion_ok = False
        if st.session_state.get('info_expediente'):
            producto_caso = st.session_state.info_expediente.get('PRODUCTO', '')
            
            if producto_caso == 'RD':
                # Para RD seguimos exigiendo todo (IFI y RSD)
                if (st.session_state.get('numero_rsd') and 
                    st.session_state.get('fecha_rsd') and 
                    st.session_state.get('numero_ifi') and 
                    st.session_state.get('fecha_ifi')):
                    resolucion_ok = True
            
            elif producto_caso == 'COERCITIVA':
                # No es necesario ingresar RSD para avanzar
                resolucion_ok = True
                # Aseguramos que las variables no sean None para evitar errores en el contexto
                if 'numero_rsd' not in st.session_state: st.session_state.numero_rsd = ""
                if 'fecha_rsd' not in st.session_state: st.session_state.fecha_rsd = date.today()
            
            else: # IFI
                # Para IFI exigimos solo la RSD
                if st.session_state.get('numero_rsd') and st.session_state.get('fecha_rsd'):
                    resolucion_ok = True

        if st.session_state.get('rubro_seleccionado') and resolucion_ok:
            with st.spinner("Preparando datos generales..."):

                from funciones import AcronymManager

                df_analistas = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Analistas")
                df_subdirecciones = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Subdirecciones")
                df_sector_sub = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Sector_Subdireccion")
                df_producto_asunto = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Producto_Asunto")
                df_administrados = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Administrados")
                df_indices_final = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Indices_BCRP")

                # --- INICIO: (REQ 1) USAR FECHA DE EMISIÓN ---
                fecha_actual = st.session_state.get('fecha_emision_informe', date.today())
                # --- FIN: (REQ 1) ---
                
                fecha_rsd_dt = st.session_state.get('fecha_rsd') or fecha_actual
                context_data = {
                    'fecha_hoy': format_date(fecha_actual, "d 'de' MMMM 'de' yyyy", locale='es').replace("septiembre", "setiembre").replace("Septiembre", "Setiembre"),
                    'mes_hoy': format_date(fecha_actual, "MMMM 'de' yyyy", locale='es').lower().replace("septiembre", "setiembre"),
                    'fecha_rsd_texto': format_date(fecha_rsd_dt, "d 'de' MMMM 'de' yyyy", locale='es').replace("septiembre", "setiembre").replace("Septiembre", "Setiembre"),
                    'acronyms': AcronymManager(),
                    'numbering': NumberingManager(),
                    'nombre_administrado': info_caso.get('ADMINISTRADO', '')
                }
                nombre_base_administrado = info_caso.get('ADMINISTRADO', '')
                nombre_final_administrado = nombre_base_administrado
                if df_administrados is not None:
                    admin_info = df_administrados[df_administrados['Nombre_Administrado_Base'] == nombre_base_administrado]
                    if not admin_info.empty:
                        nombre_final_administrado = admin_info.iloc[0].get('Nombre_Administrado', nombre_base_administrado)
                context_data['administrado'] = nombre_final_administrado
                analista_details = get_person_details_by_base_name(info_caso.get('ANALISTA ECONÓMICO'), df_analistas)
                revisor_details = get_person_details_by_base_name(info_caso.get('REVISOR'), df_analistas)
                context_data.update({
                    'titulo_analista': analista_details['titulo'], 'nombre_analista': analista_details['nombre'],
                    'cargo_analista': analista_details['cargo'], 'colegiatura_analista': analista_details['colegiatura'],
                    'titulo_revisor': revisor_details['titulo'], 'nombre_revisor': revisor_details['nombre'],
                    'cargo_revisor': revisor_details['cargo'], 'colegiatura_revisor': revisor_details['colegiatura']
                })
                num_imputaciones = int(info_caso.get('IMPUTACIONES', 1))
                inf_texto = "una" if num_imputaciones == 1 else num2words(num_imputaciones, lang='es')
                
                # --- INICIO DE LA ADICIÓN ---
                texto_incumplimiento_plural = "el incumplimiento" if num_imputaciones == 1 else "los incumplimientos"
                # --- FIN DE LA ADICIÓN ---

                # --- INICIO DE LA ADICIÓN (REQ. 1) ---
                plural_infraccion_analizada = "la infracción analizada" if num_imputaciones == 1 else "las infracciones analizadas"
                plural_cada_infraccion = "la infracción" if num_imputaciones == 1 else "cada infracción"
                # --- FIN DE LA ADICIÓN ---

                # --- INICIO: OPTIMIZACIÓN 1 (Placeholders S/P) ---
                ph_hecho_imputado = "el hecho imputado" if num_imputaciones == 1 else "los hechos imputados"
                ph_conducta_infractora = "la conducta infractora" if num_imputaciones == 1 else "las conductas infractoras"
                ph_calculo_multa = "cálculo de multa" if num_imputaciones == 1 else "cálculo de multas"
                ph_la_infraccion = "la infracción" if num_imputaciones == 1 else "las infracciones"
                ph_hecho_imputado_ant = "del hecho imputado referido" if num_imputaciones == 1 else "de los hechos imputados referidos"
                # --- FIN: OPTIMIZACIÓN 1 ---
                ph_multa_propuesta = "la multa propuesta resulta" if num_imputaciones == 1 else "las multas propuestas resultan"
                
                context_data.update({
                    'inf_numero': num_imputaciones, 'inf_texto': inf_texto,
                    'inf_oracion': "presunta infracción administrativa" if num_imputaciones == 1 else "presuntas infracciones administrativas",
                    'hechos_plural_objeto': "al hecho imputado mencionado" if num_imputaciones == 1 else "a los hechos imputados mencionados", # <-- AÑADIR ESTA LÍNEA
                    'texto_incumplimiento': texto_incumplimiento_plural, # <-- AÑADIR ESTA LÍNEA
                    # --- AÑADIR ESTAS LÍNEAS ---
                    'plural_infraccion_analizada': plural_infraccion_analizada,
                    'plural_cada_infraccion': plural_cada_infraccion,
                    # --- FIN ---
                    # --- INICIO: AÑADIR NUEVOS PLACEHOLDERS ---
                    'ph_hecho_imputado': ph_hecho_imputado,
                    'ph_conducta_infractora': ph_conducta_infractora,
                    'ph_calculo_multa': ph_calculo_multa,
                    'ph_la_infraccion': ph_la_infraccion,
                    'ph_hecho_imputado_ant': ph_hecho_imputado_ant,
                    'ph_multa_propuesta': ph_multa_propuesta
                    # --- FIN: AÑADIR NUEVOS PLACEHOLDERS ---
                })
                context_data.update({
                    'expediente': st.session_state.get('num_expediente_formateado', ''),
                    'ht': info_caso.get('HT', ''),
                    'numero_rsd': st.session_state.get('numero_rsd', ''),
                    'numero_ifi': st.session_state.get('numero_ifi', ''),
                    'fecha_ifi': format_date(st.session_state.get('fecha_ifi'), "d 'de' MMMM 'de' yyyy", locale='es') if st.session_state.get('fecha_ifi') else '',
                    
                    # Datos del Informe de Multa previo
                    'num_informe_multa_ifi': st.session_state.get('num_informe_multa_ifi', ''),
                    'monto_multa_ifi': f"{st.session_state.get('monto_multa_ifi', 0.0):,.3f} UIT",
                    'num_imputaciones_ifi': st.session_state.get('num_imputaciones_ifi', 0)
                })
                # Lógica para obtener datos de la subdirección y SSAG
                # --- Lógica de Encargado Principal (sub1) ---
                # Para RD y COERCITIVA siempre se dirige a la DFAI (Director)
                if producto_caso in ['RD', 'COERCITIVA']:
                    sub_row = df_subdirecciones[df_subdirecciones['ID_Subdireccion'].astype(str).str.strip().str.upper() == 'DFAI']
                else:
                    # Para IFI y otros, se mantiene el encargado del sector (Subdirector)
                    id_sub_row = df_sector_sub[df_sector_sub['Sector_Rubro'] == st.session_state.rubro_seleccionado]
                    sub_row = pd.DataFrame()
                    if not id_sub_row.empty:
                        id_sub_id = id_sub_row.iloc[0].get('ID_Subdireccion')
                        sub_row = df_subdirecciones[df_subdirecciones['ID_Subdireccion'] == id_sub_id]

                if not sub_row.empty:
                    # Estos placeholders ahora tendrán los datos de la DFAI en RD/COERCITIVA
                    context_data['nombre_encargado_sub1'] = sub_row.iloc[0].get('Encargado_Sub', '')
                    context_data['cargo_encargado_sub1'] = sub_row.iloc[0].get('Cargo_Encargado_Sub', '')
                    context_data['titulo_encargado_sub1'] = sub_row.iloc[0].get('Titulo_Encargado_Sub', '')
                    context_data['subdireccion'] = sub_row.iloc[0].get('Subdireccion', '')
                    context_data['id_subdireccion'] = sub_row.iloc[0].get('ID_Subdireccion', '')

                ssag_row = df_subdirecciones[df_subdirecciones['ID_Subdireccion'].astype(str).str.strip().str.upper() == 'SSAG']
                if not ssag_row.empty:
                    nombre_enc_ssag = ssag_row.iloc[0].get('Encargado_Sub')
                    context_data.update({
                        'nombre_encargado_sub2': nombre_enc_ssag,
                        'titulo_encargado_sub2': ssag_row.iloc[0].get('Titulo_Encargado_Sub', ''),
                        'cargo_encargado_sub2': ssag_row.iloc[0].get('Cargo_Encargado_Sub', '')
                    })
                    
                    # --- INICIO REQ 4: Placeholders de Iniciales ---
                    ssag_base_name = ''
                    if nombre_enc_ssag and df_analistas is not None:
                        enc_ssag_analista_row = df_analistas[df_analistas['Nombre_Analista'] == nombre_enc_ssag]
                        if not enc_ssag_analista_row.empty:
                            context_data['colegiatura_encargado_sub2'] = enc_ssag_analista_row.iloc[0].get('Colegiatura_Analista', '')
                            # Extraer el Nombre_Base_Analista (ej: RMACHUCA)
                            ssag_base_name = enc_ssag_analista_row.iloc[0].get('Nombre_Base_Analista', '')
                    
                    # --- INICIO REQ 4: Placeholders de Iniciales (Corrección v3) ---
                    
                    # Obtener nombres completos
                    nombre_completo_sub = nombre_enc_ssag
                    nombre_completo_rev = revisor_details.get('nombre', '')
                    nombre_completo_ana = analista_details.get('nombre', '')

                    # --- Placeholder 1: [RMACHUCA] ---
                    # Requerimiento: [PrimeraLetraNombre][PrimerApellido]
                    placeholder_corchetes = "[SSAG_SUBDIRECTOR]" # Default
                    if nombre_completo_sub:
                        parts = nombre_completo_sub.split()
                        # Asegurarse de que parts[0] no esté vacío
                        primera_letra = parts[0][0].upper() if parts and parts[0] else ''
                        apellido = ""

                        # Lógica para encontrar el primer apellido
                        if len(parts) == 2: # Ej: Ricardo Machuca
                            apellido = parts[1].upper()
                        elif len(parts) == 3: # Ej: Ricardo Machuca Bravo
                            apellido = parts[1].upper() # Asume que el 2do es el primer apellido
                        elif len(parts) >= 4: # Ej: Ricardo Oscar Machuca Bravo
                            apellido = parts[2].upper() # Asume que el 3ro es el primer apellido
                        
                        if primera_letra and apellido:
                            placeholder_corchetes = f"[{primera_letra}{apellido}]"
                        elif primera_letra: # Fallback si solo hay un nombre/palabra
                            placeholder_corchetes = f"[{parts[0].upper()}]"
                            
                    context_data['ssag_iniciales_corchetes'] = placeholder_corchetes
                    
                    # --- Placeholder 2: ROMB/tv/ec ---
                    # Requerimiento: INICIALES_COMPLETAS_SUB (MAYUS) / iniciales_completas_rev (minus) / iniciales_completas_ana (minus)
                    
                    # Calcular iniciales completas usando la función de funciones.py
                    ssag_iniciales = get_initials_from_name(nombre_completo_sub, to_lower=False) # MAYUS
                    revisor_iniciales = get_initials_from_name(nombre_completo_rev, to_lower=True) # minus
                    analista_iniciales = get_initials_from_name(nombre_completo_ana, to_lower=True) # minus
                    
                    context_data['ssag_iniciales_linea'] = f"{ssag_iniciales or 'SSAG'}/{revisor_iniciales or 'rev'}/{analista_iniciales or 'ana'}"
                    # --- FIN REQ 4 (Corrección v3) ---

                if producto_caso and df_producto_asunto is not None:
                    asunto_row = df_producto_asunto[df_producto_asunto['Producto'] == producto_caso]
                    if not asunto_row.empty:
                        context_data['asunto'] = asunto_row.iloc[0].get('Producto_Asunto', '')
                mes_indice_texto = "No disponible"
                if df_indices_final is not None and not df_indices_final.empty:
                    try:
                        df_indices_final['Indice_Mes_dt'] = pd.to_datetime(df_indices_final['Indice_Mes'], dayfirst=True, errors='coerce')
                        latest_date = df_indices_final['Indice_Mes_dt'].max()
                        if pd.notna(latest_date):
                            mes_indice_texto = format_date(latest_date, 'MMMM yyyy', locale='es').lower().replace("septiembre", "setiembre")
                    except Exception: pass
                context_data['mes_indice'] = mes_indice_texto
                st.session_state.context_data = context_data
                st.success("Datos generales preparados.")

        st.divider()

        # --- PASO 3: LÓGICA CONDICIONAL POR TIPO DE PRODUCTO ---
        if st.session_state.get('context_data'):
            
            # Obtener el tipo de producto del expediente para decidir qué hacer
            producto_caso = st.session_state.info_expediente.get('PRODUCTO', '')

            # ----------------------------------------------------
            # ---- CASO 1: El producto es "COERCITIVA" ----
            # ----------------------------------------------------
            if producto_caso == 'COERCITIVA':
                
                # Importar las funciones del nuevo módulo
                from producto_coercitiva import renderizar_inputs_coercitiva, validar_inputs_coercitiva, procesar_coercitiva

                # Inicializar el diccionario de estado para coercitiva si no existe
                if 'datos_informe_coercitiva' not in st.session_state:
                    st.session_state.datos_informe_coercitiva = {}

                # 1. RENDERIZAR LA INTERFAZ
                st.session_state.datos_informe_coercitiva = renderizar_inputs_coercitiva(st.session_state.datos_informe_coercitiva)

                # 2. VALIDAR LOS INPUTS
                boton_habilitado = validar_inputs_coercitiva(st.session_state.datos_informe_coercitiva)

                st.divider()
                st.header("Paso 4: Generar Informe Coercitivo")

                # 3. BOTÓN PARA PROCESAR Y GENERAR
                if st.button("🚀 Generar Informe Coercitivo", type="primary", disabled=(not boton_habilitado)):
                    with st.spinner("Generando informe..."):
                        # Preparar datos comunes que necesita el módulo de coercitiva
                        datos_comunes_coercitiva = {
                            'cliente_gspread': cliente_gspread,
                            'context_data': st.session_state.get('context_data', {}),
                            'df_productos': cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Productos"),
                        }
                        # Guardar resultados en el estado de la sesión
                        st.session_state.resultados_coercitiva = procesar_coercitiva(datos_comunes_coercitiva, st.session_state.datos_informe_coercitiva)

                # Mostrar resultados y botón de descarga si existen
                if 'resultados_coercitiva' in st.session_state:
                    resultados_finales = st.session_state.resultados_coercitiva
                    if resultados_finales and not resultados_finales.get('error'):
                        st.success("¡Informe Coercitivo generado con éxito!")
                        
                        # Mostrar tabla resumen en la app
                        resultados_app = resultados_finales.get('resultados_para_app', {})
                        if resultados_app.get('tabla_resumen_coercitivas'):
                            st.markdown("###### Resumen de Multas Coercitivas Calculadas")
                            df_resumen = pd.DataFrame(resultados_app['tabla_resumen_coercitivas'])
                            # Formatear columnas para visualización
                            df_resumen_display = df_resumen.rename(columns={
                                'num_medida': 'N° Medida',
                                'multa_base_uit': 'Multa Base (UIT)',
                                'num_coercitiva_texto': 'N° Coercitiva',
                                'multa_coercitiva_final_uit': 'Coercitiva a Imponer (UIT)'
                            })
                            st.dataframe(
                                df_resumen_display[['N° Medida', 'Multa Base (UIT)', 'N° Coercitiva', 'Coercitiva a Imponer (UIT)']],
                                column_config={
                                    "Multa Base (UIT)": st.column_config.NumberColumn(format="%.3f"),
                                    "Coercitiva a Imponer (UIT)": st.column_config.NumberColumn(format="%.3f"),
                                },
                                use_container_width=True, 
                                hide_index=True
                            )
                            st.metric("Multa Coercitiva Total (UIT)", f"{resultados_app.get('total_uit', 0):,.3f}")

                        # Botón de descarga
                        nombre_exp = st.session_state.get('num_expediente_formateado', 'EXPEDIENTE')
                        st.download_button(
                            label="✅ Descargar Informe Coercitivo (.docx)",
                            data=resultados_finales['doc_pre_compuesto'].getvalue(),
                            file_name=f"Informe_Coercitiva_{nombre_exp.replace('/', '-')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                        # Limpiar estado para la siguiente ejecución
                        del st.session_state.resultados_coercitiva
                    elif resultados_finales.get('error'):
                        st.error(f"Error: {resultados_finales['error']}")
                        del st.session_state.resultados_coercitiva

            # ----------------------------------------------------
            # ---- LÓGICA EXISTENTE PARA IFI/RD (INFRACCIONES) ----
            # ----------------------------------------------------
            elif producto_caso in ['IFI', 'RD']: # O los productos que ya manejas
                st.header("Paso 3: Detalles de Hechos Imputados")

                # --- INICIO DE LA OPTIMIZACIÓN ---
                # Cargamos todos los DataFrames necesarios para los cálculos UNA SOLA VEZ
                with st.spinner("Cargando datos para cálculos..."):
                    if 'datos_calculo' not in st.session_state:
                        st.session_state.datos_calculo = {
                            'df_tipificacion': cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Tipificacion_Infracciones"),
                            'df_items_infracciones': cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Items_Infracciones"),
                            'df_costos_items': cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Costos_Items"),
                            'df_salarios_general': cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Salarios_General"),
                            'df_cos': cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "COS"),
                            'df_uit': cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "UIT"),
                            'df_coti_general': pd.to_datetime(cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Cotizaciones_General")['Fecha_Costeo'], dayfirst=True, errors='coerce').to_frame().join(cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Cotizaciones_General").drop('Fecha_Costeo', axis=1)),
                            'df_indices': pd.to_datetime(cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Indices_BCRP")['Indice_Mes'], dayfirst=True, errors='coerce').to_frame().join(cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Indices_BCRP").drop('Indice_Mes', axis=1)),
                            'df_dias_no_laborables': cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Dias_No_Laborables")
                        }
                
                # Obtenemos los dataframes desde el estado de la sesión para usarlos
                datos_calculo = st.session_state.datos_calculo
                df_tipificacion = datos_calculo['df_tipificacion']
                # --- FIN DE LA OPTIMIZACIÓN ---
                
                for i in range(len(st.session_state.get('imputaciones_data', []))):
                    with st.expander(f"Hecho imputado n.° {i + 1}", expanded=(i == 0)):
                        st.subheader(f"Detalles del Hecho {i + 1}")
                        
                        # --- Lógica de widgets con estado corregida ---
                        texto_guardado = st.session_state.imputaciones_data[i].get('texto_hecho', '')
                        texto_ingresado = st.text_area("Redacta aquí el hecho imputado:", value=texto_guardado, key=f"texto_hecho_{i}", height=150)
                        st.session_state.imputaciones_data[i]['texto_hecho'] = texto_ingresado
                        st.divider()

                        if df_tipificacion is not None:
                            try:
                                lista_tipos_infraccion = df_tipificacion['Tipo_Infraccion'].unique().tolist()
                                index_tipo = None
                                tipo_guardado = st.session_state.imputaciones_data[i].get('tipo_seleccionado')
                                if tipo_guardado in lista_tipos_infraccion:
                                    index_tipo = lista_tipos_infraccion.index(tipo_guardado)
                                tipo_seleccionado = st.radio("**Selecciona el tipo de infracción:**", options=lista_tipos_infraccion, index=index_tipo, horizontal=True, key=f"radio_tipo_infraccion_{i}")
                                st.session_state.imputaciones_data[i]['tipo_seleccionado'] = tipo_seleccionado

                                if tipo_seleccionado:
                                    subtipos_df = df_tipificacion[df_tipificacion['Tipo_Infraccion'] == tipo_seleccionado]
                                    lista_subtipos = subtipos_df['Descripcion_Infraccion'].tolist()
                                    index_subtipo = None
                                    subtipo_guardado = st.session_state.imputaciones_data[i].get('subtipo_seleccionado')
                                    if subtipo_guardado in lista_subtipos:
                                        index_subtipo = lista_subtipos.index(subtipo_guardado)
                                    subtipo_seleccionado = st.selectbox("**Selecciona la descripción de la infracción:**", options=lista_subtipos, index=index_subtipo, placeholder="Elige una descripción específica...", key=f"subtipo_infraccion_{i}")
                                    st.session_state.imputaciones_data[i]['subtipo_seleccionado'] = subtipo_seleccionado

                                    # En app.py, dentro del bucle for del Paso 3

                                    if subtipo_seleccionado:
                                        # --- INICIO DE LA CORRECCIÓN ---
                                        # Paso 1: Buscamos la fila y guardamos el resultado (que puede estar vacío)
                                        fila_encontrada = subtipos_df[subtipos_df['Descripcion_Infraccion'] == subtipo_seleccionado]
                                        
                                        # Paso 2: Solo si la búsqueda tuvo éxito (no está vacía), extraemos los datos
                                        if not fila_encontrada.empty:
                                            id_infraccion = fila_encontrada.iloc[0]['ID_Infraccion']
                                            st.session_state.imputaciones_data[i]['id_infraccion'] = id_infraccion
                                        # --- FIN DE LA CORRECCIÓN ---
                            except KeyError as e:
                                st.error(f"Error en la hoja 'Tipificacion_Infracciones'. Falta la columna: {e}")
                        else:
                            st.error("Error crítico: No se pudo cargar la hoja 'Tipificacion_Infracciones'.")
                            
                        # --- Lógica de cálculo del hecho (ya corregida) ---
                        id_infraccion = st.session_state.imputaciones_data[i].get('id_infraccion')
                        if id_infraccion:
                            try:
                                modulo_especialista = importlib.import_module(f"infracciones.{id_infraccion}")
                                datos_especificos = modulo_especialista.renderizar_inputs_especificos(i, datos_calculo.get('df_dias_no_laborables'))
                                st.session_state.imputaciones_data[i].update(datos_especificos)
                                
                                # --- INICIO: SECCIÓN GRADUACIÓN DE SANCIONES (Ubicación Correcta) ---
                                st.divider()
                                st.subheader("Factores de graduación")
                                
                                datos_hecho_actual = st.session_state.imputaciones_data[i]
                                
                                # 1. Botón de activación inicial
                                aplica_graduacion = st.radio(
                                    "¿Existen factores agravantes o atenuantes para graduar la sanción?",
                                    ["No", "Sí"],
                                    key=f"aplica_grad_{i}",
                                    index=0 if datos_hecho_actual.get('aplica_graduacion', 'No') == 'No' else 1,
                                    horizontal=True
                                )
                                datos_hecho_actual['aplica_graduacion'] = aplica_graduacion
                                
                                factor_f_final = 1.0 # Valor por defecto (Neutro)
                                es_eximente = False
                                
                                if aplica_graduacion == "Sí":
                                    with st.expander("Configurar Factores de Graduación", expanded=True):
                                        st.info("Seleccione los criterios. Por defecto, todos inician en 0%.")
                                        
                                        # Recuperar o inicializar el diccionario de graduación
                                        if 'graduacion' not in datos_hecho_actual:
                                            datos_hecho_actual['graduacion'] = {}
                                        seleccion_graduacion = datos_hecho_actual['graduacion']
                                        
                                        total_porcentaje_f = 0.0

                                        # Iterar sobre f1, f2, ... f7 (Usando la constante FACTORES_GRADUACION)
                                        for codigo_f, data_f in FACTORES_GRADUACION.items():
                                            st.markdown(f"**{codigo_f.upper()}: {data_f['titulo']}**")
                                            
                                            subtotal_f = 0.0
                                            
                                            for subcriterio, opciones in data_f['criterios'].items():
                                                key_widget = f"grad_{i}_{codigo_f}_{subcriterio}"
                                                val_guardado = seleccion_graduacion.get(key_widget)
                                                
                                                # Asegurar que la opción por defecto sea la primera (que suele ser 0.0)
                                                lista_opciones = list(opciones.keys())
                                                idx_sel = 0
                                                if val_guardado in lista_opciones:
                                                    idx_sel = lista_opciones.index(val_guardado)

                                                opcion_sel = st.selectbox(
                                                    subcriterio,
                                                    options=lista_opciones,
                                                    key=key_widget,
                                                    index=idx_sel
                                                )
                                                
                                                valor = opciones[opcion_sel]
                                                
                                                # Guardar selección
                                                seleccion_graduacion[key_widget] = opcion_sel
                                                
                                                # Lógica Eximente
                                                if valor == "Eximente":
                                                    es_eximente = True
                                                    valor_num = 0.0
                                                    seleccion_graduacion[f"{key_widget}_valor"] = "Eximente"
                                                else:
                                                    valor_num = float(valor)
                                                    seleccion_graduacion[f"{key_widget}_valor"] = valor_num

                                                subtotal_f += valor_num
                                            
                                            # Visualización del subtotal
                                            if subtotal_f != 0:
                                                color_txt = "red" if subtotal_f > 0 else "green" # Rojo=Agrava, Verde=Atenúa
                                                st.caption(f":{color_txt}[Subtotal {codigo_f}: {subtotal_f:+.0%}]")
                                            
                                            total_porcentaje_f += subtotal_f
                                            seleccion_graduacion[f"subtotal_{codigo_f}"] = subtotal_f

                                        # Cálculo final de F
                                        if es_eximente:
                                            st.warning("⚠️ **SE APLICA EXIMENTE DE RESPONSABILIDAD** (Multa = 0)")
                                            factor_f_final = 0.0
                                        else:
                                            factor_f_final = 1.0 + total_porcentaje_f
                                            if factor_f_final < 0: factor_f_final = 0.0 # No puede ser negativo
                                            
                                            col_res1, col_res2 = st.columns([3, 1])
                                            with col_res1:
                                                st.markdown(f"**Ajuste total acumulado:** {total_porcentaje_f:+.0%}")
                                            with col_res2:
                                                st.metric("Factor F Final", f"{factor_f_final:.2%}")
                                else:
                                    # Si dice "No", limpiamos o reseteamos
                                    # Opcional: datos_hecho_actual['graduacion'] = {} 
                                    pass

                                # Guardar en el estado para el cálculo posterior
                                datos_hecho_actual['factor_f_calculado'] = factor_f_final
                                datos_hecho_actual['es_eximente'] = es_eximente
                                # --- FIN: SECCIÓN GRADUACIÓN ---

                                # --- INICIO: (REQ 1) AÑADIR INPUTS DE REDUCCIÓN DE MULTA ---
                                st.divider()
                                st.subheader("Reconocimiento de responsabilidad")
                                datos_hecho_actual = st.session_state.imputaciones_data[i]
                                
                                aplica_reduccion = st.radio(
                                    "¿Aplica reducción de multa por reconocimiento de responsabilidad?",
                                    ["No", "Sí"],
                                    key=f"aplica_reduccion_{i}",
                                    index=0 if datos_hecho_actual.get('aplica_reduccion', 'No') == 'No' else 1,
                                    horizontal=True
                                )
                                datos_hecho_actual['aplica_reduccion'] = aplica_reduccion
                                
                                if aplica_reduccion == "Sí":
                                    porcentaje_reduccion = st.radio(
                                        "Seleccione el porcentaje:",
                                        ["50%", "30%"],
                                        key=f"porcentaje_reduccion_{i}",
                                        index=0 if datos_hecho_actual.get('porcentaje_reduccion') == '50%' else 1,
                                        horizontal=True
                                    )
                                    datos_hecho_actual['porcentaje_reduccion'] = porcentaje_reduccion
                                    
                                    # Definir el texto del placeholder basado en el porcentaje
                                    if porcentaje_reduccion == "50%":
                                        datos_hecho_actual['texto_reduccion'] = "en la presentación de los descargos contra la imputación de cargos"
                                    else: # 30%
                                        datos_hecho_actual['texto_reduccion'] = "mediante los descargos al IFI"
                                    
                                    st.info(f"Se aplicará reducción del **{porcentaje_reduccion}** ({datos_hecho_actual['texto_reduccion']}).")
                                    
                                    col_memo1, col_memo2 = st.columns(2)
                                    with col_memo1:
                                        datos_hecho_actual['memo_num'] = st.text_input("N.° de Memorando:", value=datos_hecho_actual.get('memo_num', ''), key=f"memo_num_{i}")
                                    with col_memo2:
                                        datos_hecho_actual['memo_fecha'] = st.date_input("Fecha del Memorando:", value=datos_hecho_actual.get('memo_fecha'), key=f"memo_fecha_{i}", format="DD/MM/YYYY")
                                        
                                    col_esc1, col_esc2 = st.columns(2)
                                    with col_esc1:
                                        datos_hecho_actual['escrito_num'] = st.text_input("N.° de Escrito (Administrado):", value=datos_hecho_actual.get('escrito_num', ''), key=f"escrito_num_{i}")
                                    with col_esc2:
                                        datos_hecho_actual['escrito_fecha'] = st.date_input("Fecha del Escrito:", value=datos_hecho_actual.get('escrito_fecha'), key=f"escrito_fecha_{i}", format="DD/MM/YYYY")
                                
                                # --- FIN: (REQ 1) ---

                                datos_generales_ok = st.session_state.imputaciones_data[i].get('texto_hecho') and st.session_state.imputaciones_data[i].get('subtipo_seleccionado')
                                datos_especificos_ok = modulo_especialista.validar_inputs(st.session_state.imputaciones_data[i])
                                
                                # --- INICIO: (REQ 1) VALIDACIÓN DE REDUCCIÓN ---
                                datos_reduccion_ok = True
                                datos_hecho_actual = st.session_state.imputaciones_data[i]
                                if datos_hecho_actual.get('aplica_reduccion') == 'Sí':
                                    if not all([
                                        datos_hecho_actual.get('porcentaje_reduccion'),
                                        datos_hecho_actual.get('memo_num'),
                                        datos_hecho_actual.get('memo_fecha'),
                                        datos_hecho_actual.get('escrito_num'),
                                        datos_hecho_actual.get('escrito_fecha')
                                    ]):
                                        datos_reduccion_ok = False
                                        st.warning("Para aplicar la reducción, debe completar todos los campos del memorando y del escrito.")
                                # --- FIN: (REQ 1) ---
                                
                                boton_habilitado = datos_generales_ok and datos_especificos_ok and datos_reduccion_ok

                                st.divider()
                                if st.button(f"Calcular Hecho {i+1}", key=f"calc_btn_{i}", disabled=(not boton_habilitado)):
                                    with st.spinner(f"Calculando hecho {i+1}..."):
                                        # --- INICIO DE LA OPTIMIZACIÓN ---
                                        # 1. Buscamos el ID de la plantilla (esto es rápido)
                                        df_plantillas = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Productos")
                                        producto_caso = st.session_state.info_expediente.get('PRODUCTO', '')
                                        plantilla_row = df_plantillas[df_plantillas['Producto'] == producto_caso]
                                        id_plantilla_inicio_calc = plantilla_row.iloc[0].get('ID_Plantilla_Inicio') if not plantilla_row.empty else df_plantillas[df_plantillas['Producto'] == 'DEFAULT'].iloc[0].get('ID_Plantilla_Inicio')

                                        if not id_plantilla_inicio_calc:
                                            st.error("No se pudo encontrar la plantilla de inicio para el cálculo.")
                                            st.stop()
                                        
                                        buffer_tpl_estilos = descargar_archivo_drive(id_plantilla_inicio_calc)
                                        if not buffer_tpl_estilos:
                                            st.error("No se pudo descargar la plantilla de inicio para el cálculo.")
                                            st.stop()
                                        
                                        # 2. Preparar datos comunes USANDO LOS DATOS YA CARGADOS
                                        acronym_manager = st.session_state.context_data.get('acronyms')
                                        datos_comunes = {
                                            **datos_calculo, # <-- Usamos todos los DFs precargados
                                            'datos_hecho_completos': st.session_state.imputaciones_data[i], # <-- AÑADIR
                                            'fecha_emision_informe': st.session_state.get('fecha_emision_informe', date.today()),
                                            'id_infraccion': id_infraccion,
                                            'rubro': st.session_state.rubro_seleccionado,
                                            'id_rubro_seleccionado': st.session_state.get('id_rubro_seleccionado'),
                                            'numero_hecho_actual': i + 1,
                                            'doc_tpl': DocxTemplate(buffer_tpl_estilos),
                                            'context_data': st.session_state.get('context_data', {}),
                                            'acronym_manager': acronym_manager
                                        }
                                        # --- FIN DE LA OPTIMIZACIÓN ---

                                        resultados_completos = modulo_especialista.procesar_infraccion(
                                            datos_comunes, 
                                            st.session_state.imputaciones_data[i]
                                        )
                                        if resultados_completos.get('error'):
                                            st.error(f"Error en el cálculo del Hecho {i+1}: {resultados_completos['error']}")
                                        else:
                                            st.session_state.imputaciones_data[i]['resultados'] = resultados_completos
                                            st.session_state.imputaciones_data[i]['es_extemporaneo'] = resultados_completos.get('es_extemporaneo', False)
                                            st.session_state.imputaciones_data[i]['usa_capacitacion'] = resultados_completos.get('usa_capacitacion', False)
                                            st.success(f"Hecho {i+1} calculado.")
                                            st.session_state.imputaciones_data[i]['anexos_ce'] = resultados_completos.get('anexos_ce_generados', [])
                                            # --- NUEVA LÍNEA: Actualizar todos los hechos para reflejar el prorrateo ---
                                            actualizar_y_recalcular_prorrateos(cliente_gspread, NOMBRE_GSHEET_MAESTRO)
                                            
                                            st.success(f"Hecho {i+1} calculado y montos prorrateados actualizados.")
                                            st.rerun() # Forzar refresco para ver los cambios
                            except ImportError:
                                st.error(f"El módulo para '{id_infraccion}' no está implementado.")


                            # --- Sección para mostrar resultados ya calculados ---
                            if 'resultados' in st.session_state.imputaciones_data[i]:
                                resultados_app = st.session_state.imputaciones_data[i]['resultados'].get('resultados_para_app', {}) # Use .get for safety
                                id_infraccion_actual = st.session_state.imputaciones_data[i].get('id_infraccion', '')
                                st.subheader(f"Resultados del Cálculo para el Hecho {i + 1}")

                                totales_finales = {}

                                # Lógica específica para INF003
                                if 'INF003' in id_infraccion_actual or 'INF006' in id_infraccion_actual:
                                    if 'extremos' in resultados_app and len(resultados_app['extremos']) > 1:
                                        st.markdown("#### Desglose de Costo Evitado (CE) por Extremo")
                                        for j, extremo_data in enumerate(resultados_app['extremos']):
                                            st.markdown(f"##### Extremo {j + 1}: {extremo_data.get('tipo', 'N/A')}")
                                            df_ce_extremo = pd.DataFrame(extremo_data.get('ce_data', []))
                                            if not df_ce_extremo.empty:
                                                # Asumiendo que ce_data tiene las columnas correctas para INF003
                                                df_ce_extremo = df_ce_extremo.drop(columns=['id_anexo'], errors='ignore')
                                                df_ce_renamed = df_ce_extremo.rename(columns={
                                                    'descripcion': 'Descripción',
                                                    'precio_dolares': 'Precio Base (US$)', # Ajustar nombre si es diferente
                                                    'precio_soles': 'Precio Base (S/)',   # Ajustar nombre si es diferente
                                                    'factor_ajuste': 'Factor de Ajuste',
                                                    'monto_soles': 'Monto (S/)',
                                                    'monto_dolares': 'Monto (US$)'
                                                })
                                                numeric_cols_ce = ['Precio Base (US$)', 'Precio Base (S/)', 'Factor de Ajuste', 'Monto (S/)', 'Monto (US$)']
                                                for col in numeric_cols_ce:
                                                    if col in df_ce_renamed.columns: df_ce_renamed[col] = pd.to_numeric(df_ce_renamed[col], errors='coerce')
                                                st.dataframe(df_ce_renamed.style.format("{:,.3f}", subset=[c for c in numeric_cols_ce if c in df_ce_renamed.columns], na_rep='').hide(axis="index"), use_container_width=True)

                                        st.markdown("#### Beneficio Ilícito (BI) Consolidado")
                                        df_bi_total = pd.DataFrame(resultados_app.get('totales', {}).get('bi_data_raw', []))
                                        if not df_bi_total.empty:
                                            # Usar descripcion_texto si existe (para superíndices), sino 'descripcion'
                                            desc_col = 'descripcion_texto' if 'descripcion_texto' in df_bi_total.columns else 'descripcion'
                                            df_bi_display = df_bi_total.rename(columns={desc_col: 'Descripción', 'monto': 'Monto'})
                                            # Seleccionar solo las columnas a mostrar
                                            cols_to_show = ['Descripción', 'Monto']
                                            if 'descripcion_superindice' in df_bi_total.columns: # Añadir superíndice si existe
                                                df_bi_display['Descripción'] = df_bi_display['Descripción'] + df_bi_total['descripcion_superindice'].fillna('')
                                            st.dataframe(df_bi_display[cols_to_show].style.hide(axis="index"), use_container_width=True)

                                        totales_finales = resultados_app.get('totales', {})
                                    else:
                                        totales_finales = resultados_app # Caso simple INF003
                                        st.markdown("###### Costo Evitado (CE)")
                                        df_ce_total = pd.DataFrame(totales_finales.get('ce_data_raw', []))
                                        if not df_ce_total.empty:
                                            df_ce_total = df_ce_total.drop(columns=['id_anexo'], errors='ignore')
                                            df_ce_renamed = df_ce_total.rename(columns={
                                                'descripcion': 'Descripción',
                                                'precio_dolares': 'Precio Base (US$)', # Ajustar nombre si es diferente
                                                'precio_soles': 'Precio Base (S/)',   # Ajustar nombre si es diferente
                                                'factor_ajuste': 'Factor de Ajuste',
                                                'monto_soles': 'Monto (S/)',
                                                'monto_dolares': 'Monto (US$)'
                                            })
                                            numeric_cols_ce = ['Precio Base (US$)', 'Precio Base (S/)', 'Factor de Ajuste', 'Monto (S/)', 'Monto (US$)']
                                            for col in numeric_cols_ce:
                                                if col in df_ce_renamed.columns: df_ce_renamed[col] = pd.to_numeric(df_ce_renamed[col], errors='coerce')
                                            st.dataframe(df_ce_renamed.style.format("{:,.3f}", subset=[c for c in numeric_cols_ce if c in df_ce_renamed.columns], na_rep='').hide(axis="index"), use_container_width=True)

                                        st.markdown("###### Beneficio Ilícito (BI)")
                                        df_bi_total = pd.DataFrame(totales_finales.get('bi_data_raw', []))
                                        if not df_bi_total.empty:
                                            # Usar descripcion_texto si existe (para superíndices), sino 'descripcion'
                                            desc_col = 'descripcion_texto' if 'descripcion_texto' in df_bi_total.columns else 'descripcion'
                                            df_bi_display = df_bi_total.rename(columns={desc_col: 'Descripción', 'monto': 'Monto'})
                                            cols_to_show = ['Descripción', 'Monto']
                                            if 'descripcion_superindice' in df_bi_total.columns:
                                                df_bi_display['Descripción'] = df_bi_display['Descripción'] + df_bi_total['descripcion_superindice'].fillna('')
                                            st.dataframe(df_bi_display[cols_to_show].style.hide(axis="index"), use_container_width=True)

                                # Lógica específica para INF002
                                elif 'INF002' in id_infraccion_actual:
                                    if 'extremos' in resultados_app and isinstance(resultados_app['extremos'], list):
                                        st.markdown("#### Desglose por Extremo de Monitoreo")

                                        columnas_map_ce1 = {'descripcion': 'Descripción', 'horas': 'Horas', 'cantidad': 'Cantidad', 'precio_unitario': 'Precio Unitario (S/)', 'factor_ajuste': 'Factor Ajuste', 'monto_soles': 'Monto (S/)', 'monto_dolares': 'Monto (US$)'}
                                        columnas_map_ce2 = {'descripcion': 'Descripción', 'unidad': 'Unidad', 'cantidad': 'Cantidad', 'precio_unitario': 'Precio Unitario (S/)', 'factor_ajuste': 'Factor Ajuste', 'monto_soles': 'Monto (S/)', 'monto_dolares': 'Monto (US$)'}
                                        cols_numericas = ['Cantidad', 'Precio Unitario (S/)', 'Factor Ajuste', 'Monto (S/)', 'Monto (US$)', 'Horas']

                                        for j, extremo_data in enumerate(resultados_app['extremos']):
                                            st.markdown(f"##### Extremo {j + 1}: {extremo_data.get('tipo', 'N/A')}")

                                            if extremo_data.get('ce1_data'):
                                                st.markdown("###### Costo Evitado por Muestreo (CE1)")
                                                df_ce1 = pd.DataFrame(extremo_data['ce1_data']).rename(columns=columnas_map_ce1)
                                                cols_existentes = [col for col in cols_numericas if col in df_ce1.columns]
                                                st.dataframe(df_ce1.style.format("{:,.3f}", subset=cols_existentes, na_rep='').hide(axis="index"), use_container_width=True)

                                            if extremo_data.get('ce2_envio_data'):
                                                st.markdown("###### Costo Evitado por Envío de Muestras")
                                                df_ce2_envio = pd.DataFrame(extremo_data['ce2_envio_data']).rename(columns=columnas_map_ce2)
                                                cols_existentes = [col for col in cols_numericas if col in df_ce2_envio.columns]
                                                st.dataframe(df_ce2_envio.style.format("{:,.3f}", subset=cols_existentes, na_rep='').hide(axis="index"), use_container_width=True)

                                            if extremo_data.get('ce2_lab_data'):
                                                st.markdown("###### Costo Evitado por Análisis de Laboratorio")
                                                df_ce2_lab = pd.DataFrame(extremo_data['ce2_lab_data']).rename(columns=columnas_map_ce2)
                                                cols_existentes = [col for col in cols_numericas if col in df_ce2_lab.columns]
                                                st.dataframe(df_ce2_lab.style.format("{:,.3f}", subset=cols_existentes, na_rep='').hide(axis="index"), use_container_width=True)

                                            total_ce2_soles_extremo = extremo_data.get('ce2_envio_soles', 0) + extremo_data.get('ce2_lab_soles', 0)
                                            total_ce2_dolares_extremo = extremo_data.get('ce2_envio_dolares', 0) + extremo_data.get('ce2_lab_dolares', 0)

                                            if total_ce2_soles_extremo > 0:
                                                st.markdown("###### Resumen de Costo Evitado (CE2)")
                                                resumen_data_ce2 = []
                                                if extremo_data.get('ce2_envio_soles', 0) > 0:
                                                    resumen_data_ce2.append({'Componente': 'Subtotal Envío de Muestras', 'Monto (S/)': f"{extremo_data['ce2_envio_soles']:,.3f}", 'Monto (US$)': f"{extremo_data['ce2_envio_dolares']:,.3f}"})
                                                if extremo_data.get('ce2_lab_soles', 0) > 0:
                                                    resumen_data_ce2.append({'Componente': 'Subtotal Análisis de Laboratorio', 'Monto (S/)': f"{extremo_data['ce2_lab_soles']:,.3f}", 'Monto (US$)': f"{extremo_data['ce2_lab_dolares']:,.3f}"})
                                                resumen_data_ce2.append({'Componente': 'Total Costo Evitado (CE2)', 'Monto (S/)': f"{total_ce2_soles_extremo:,.3f}", 'Monto (US$)': f"{total_ce2_dolares_extremo:,.3f}"})
                                                st.dataframe(pd.DataFrame(resumen_data_ce2).style.hide(axis="index"), use_container_width=True)

                                            st.markdown("###### Resumen Total del Costo Evitado del Extremo")
                                            resumen_data_total = [
                                                {'Componente': 'Total Costo Evitado (CE1)', 'Monto (S/)': f"{extremo_data.get('ce1_soles', 0):,.3f}", 'Monto (US$)': f"{extremo_data.get('ce1_dolares', 0):,.3f}"},
                                                {'Componente': 'Total Costo Evitado (CE2)', 'Monto (S/)': f"{total_ce2_soles_extremo:,.3f}", 'Monto (US$)': f"{total_ce2_dolares_extremo:,.3f}"},
                                                {'Componente': 'Costo Evitado Total del Extremo', 'Monto (S/)': f"{extremo_data.get('total_soles_extremo', 0):,.3f}", 'Monto (US$)': f"{extremo_data.get('total_dolares_extremo', 0):,.3f}"}
                                            ]
                                            st.dataframe(pd.DataFrame(resumen_data_total).style.hide(axis="index"), use_container_width=True)

                                        totales_finales = resultados_app.get('totales', {})
                                        st.markdown("--- \n#### Totales Consolidados del Hecho")
                                        st.markdown("###### Beneficio Ilícito (BI)")
                                        df_bi_total = pd.DataFrame(totales_finales.get('bi_data_raw', []))
                                        if not df_bi_total.empty:
                                            # Usar descripcion_texto si existe (para superíndices), sino 'descripcion'
                                            desc_col = 'descripcion_texto' if 'descripcion_texto' in df_bi_total.columns else 'descripcion'
                                            df_bi_display = df_bi_total.rename(columns={desc_col: 'Descripción', 'monto': 'Monto'})
                                            cols_to_show = ['Descripción', 'Monto']
                                            if 'descripcion_superindice' in df_bi_total.columns:
                                                df_bi_display['Descripción'] = df_bi_display['Descripción'] + df_bi_total['descripcion_superindice'].fillna('')
                                            st.dataframe(df_bi_display[cols_to_show].style.hide(axis="index"), use_container_width=True)

                                # --- Lógica específica para INF004 ---
                                elif 'INF004' in id_infraccion_actual  or 'INF009' in id_infraccion_actual or 'INF010' in id_infraccion_actual:
                                    if 'extremos' in resultados_app and isinstance(resultados_app['extremos'], list):
                                        # --- Lógica para Múltiples Extremos (INF004) ---
                                        st.markdown("#### Desglose por Extremo")
                                        # Definir mapeo COMPLETO de columnas para INF004
                                        columnas_map_inf004 = {
                                            'descripcion': 'Descripción',
                                            'cantidad': 'Cantidad',
                                            'horas': 'Horas',
                                            'precio_soles': 'Precio asociado (S/)',
                                            'factor_ajuste': 'Factor Ajuste',
                                            'monto_soles': 'Monto (S/)',
                                            'monto_dolares': 'Monto (US$)'
                                        }
                                        # Columnas que SÍ deben tener 3 decimales fijos
                                        cols_formato_numerico = ['Precio asociado (S/)', 'Factor Ajuste', 'Monto (S/)', 'Monto (US$)']

                                        for j, extremo_data in enumerate(resultados_app['extremos']):
                                            st.markdown(f"##### Extremo {j + 1}: {extremo_data.get('tipo', 'N/A')}")
                                            st.markdown("###### Costo Evitado (CE)")
                                            datos_ce_crudos = extremo_data.get('ce_data', [])
                                            if datos_ce_crudos:
                                                df_ce_display = pd.DataFrame(datos_ce_crudos)

                                                # --- INICIO Req 2: Formato dinámico ---
                                                if 'cantidad' in df_ce_display.columns:
                                                    df_ce_display['cantidad'] = df_ce_display['cantidad'].apply(format_decimal_dinamico)
                                                if 'horas' in df_ce_display.columns:
                                                    df_ce_display['horas'] = df_ce_display['horas'].apply(format_decimal_dinamico)
                                                # --- FIN Req 2 ---
                                                
                                                # --- INICIO: AÑADIR FILA TOTAL ---
                                                total_monto_soles = sum(item.get('monto_soles', 0) for item in datos_ce_crudos)
                                                total_monto_dolares = sum(item.get('monto_dolares', 0) for item in datos_ce_crudos)
                                                total_df = pd.DataFrame([{'descripcion': 'Total', 'monto_soles': total_monto_soles, 'monto_dolares': total_monto_dolares}])
                                                df_ce_display = pd.concat([df_ce_display, total_df], ignore_index=True)
                                                # --- FIN: AÑADIR FILA TOTAL ---
                                                
                                                # --- INICIO Req 3: Arreglar tabla deformada ---
                                                columnas_a_mostrar = [col for col in columnas_map_inf004.keys() if col in df_ce_display.columns]
                                                df_ce_display = df_ce_display[columnas_a_mostrar].rename(columns=columnas_map_inf004)
                                                # Reindexar para asegurar todas las columnas y evitar deformación
                                                df_ce_display = df_ce_display.reindex(columns=columnas_map_inf004.values())
                                                # --- FIN Req 3 ---

                                                # Formatear columnas numéricas (moneda/factor)
                                                cols_numericas_existentes = [col for col in cols_formato_numerico if col in df_ce_display.columns]
                                                st.dataframe(
                                                    df_ce_display.style.format("{:,.3f}", subset=cols_numericas_existentes, na_rep='').hide(axis="index"), 
                                                    use_container_width=True
                                                )
                                            else:
                                                st.warning("No hay columnas válidas para mostrar en el Costo Evitado.")
                                            
                                            # Mostrar BI del extremo (si existe en los datos)
                                            st.markdown("###### Beneficio Ilícito (BI) del Extremo")
                                            df_bi_extremo = pd.DataFrame(extremo_data.get('bi_data', []))
                                            if not df_bi_extremo.empty:
                                                # Usar descripcion_texto si existe, sino 'descripcion'
                                                desc_col = 'descripcion_texto' if 'descripcion_texto' in df_bi_extremo.columns else 'descripcion'
                                                df_bi_display = df_bi_extremo.rename(columns={desc_col: 'Descripción', 'monto': 'Monto'})
                                                cols_to_show = ['Descripción', 'Monto']
                                                
                                                # --- INICIO DE LA CORRECCIÓN ---
                                                # 1. Verificar si 'descripcion_superindice' existe en el DF *renombrado*
                                                if 'descripcion_superindice' in df_bi_display.columns:
                                                    # 2. Usar las columnas del DF *renombrado* para la operación
                                                    df_bi_display['Descripción'] = df_bi_display['Descripción'] + df_bi_display['descripcion_superindice'].fillna('')
                                                
                                                # 3. Asegurarse de que las columnas existan ANTES de seleccionarlas
                                                cols_finales = [col for col in cols_to_show if col in df_bi_display.columns]
                                                st.dataframe(df_bi_display[cols_finales].style.hide(axis="index"), use_container_width=True)
                                                # --- FIN DE LA CORRECCIÓN ---
                                            else:
                                                st.warning("No hay datos de Beneficio Ilícito para este extremo.")

                                        totales_finales = resultados_app.get('totales', {}) # Obtener totales consolidados

                                    else:
                                    # --- Lógica para Hecho Simple (INF004) ---
                                        totales_finales = resultados_app 
                                        st.markdown("###### Costo Evitado (CE)")
                                        datos_ce_crudos = totales_finales.get('ce_data_raw', [])
                                        if datos_ce_crudos:
                                            df_ce_display = pd.DataFrame(datos_ce_crudos)
                                            # Definir mapeo COMPLETO de columnas para INF004
                                            columnas_map_inf004 = {
                                                'descripcion': 'Descripción',
                                                'cantidad': 'Cantidad',
                                                'horas': 'Horas',
                                                'precio_soles': 'Precio asociado (S/)',
                                                'factor_ajuste': 'Factor Ajuste',
                                                'monto_soles': 'Monto (S/)',
                                                'monto_dolares': 'Monto (US$)'
                                            }
                                            # Columnas que SÍ deben tener 3 decimales fijos
                                            cols_formato_numerico = ['Precio asociado (S/)', 'Factor Ajuste', 'Monto (S/)', 'Monto (US$)']

                                            # --- INICIO Req 2: Formato dinámico ---
                                            if 'cantidad' in df_ce_display.columns:
                                                df_ce_display['cantidad'] = df_ce_display['cantidad'].apply(format_decimal_dinamico)
                                            if 'horas' in df_ce_display.columns:
                                                df_ce_display['horas'] = df_ce_display['horas'].apply(format_decimal_dinamico)
                                            # --- FIN Req 2 ---
                                            
                                            # --- INICIO: AÑADIR FILA TOTAL ---
                                            total_monto_soles = sum(item.get('monto_soles', 0) for item in datos_ce_crudos)
                                            total_monto_dolares = sum(item.get('monto_dolares', 0) for item in datos_ce_crudos)
                                            total_df = pd.DataFrame([{'descripcion': 'Total', 'monto_soles': total_monto_soles, 'monto_dolares': total_monto_dolares}])
                                            df_ce_display = pd.concat([df_ce_display, total_df], ignore_index=True)
                                            # --- FIN: AÑADIR FILA TOTAL ---
                                            
                                            # --- INICIO Req 3: Arreglar tabla deformada ---
                                            columnas_a_mostrar = [col for col in columnas_map_inf004.keys() if col in df_ce_display.columns]
                                            df_ce_display = df_ce_display[columnas_a_mostrar].rename(columns=columnas_map_inf004)
                                            # Reindexar para asegurar todas las columnas y evitar deformación
                                            df_ce_display = df_ce_display.reindex(columns=columnas_map_inf004.values())
                                            # --- FIN Req 3 ---

                                            # Formatear columnas numéricas (moneda/factor)
                                            cols_numericas_existentes = [col for col in cols_formato_numerico if col in df_ce_display.columns]
                                            st.dataframe(
                                                df_ce_display.style.format("{:,.3f}", subset=cols_numericas_existentes, na_rep='').hide(axis="index"), 
                                                use_container_width=True
                                            )
                                        else:
                                            st.warning("No hay columnas válidas para mostrar en el Costo Evitado.")

                                        st.markdown("###### Beneficio Ilícito (BI)")
                                        df_bi_total = pd.DataFrame(totales_finales.get('bi_data_raw', []))
                                        if not df_bi_total.empty:
                                            # Usar descripcion_texto si existe, sino 'descripcion'
                                            desc_col = 'descripcion_texto' if 'descripcion_texto' in df_bi_total.columns else 'descripcion'
                                            df_bi_display = df_bi_total.rename(columns={desc_col: 'Descripción', 'monto': 'Monto'})
                                            cols_to_show = ['Descripción', 'Monto']

                                            # --- INICIO DE LA CORRECCIÓN ---
                                            # 1. Verificar si 'descripcion_superindice' existe en el DF *renombrado*
                                            if 'descripcion_superindice' in df_bi_display.columns:
                                                # 2. Usar las columnas del DF *renombrado* para la operación
                                                df_bi_display['Descripción'] = df_bi_display['Descripción'] + df_bi_display['descripcion_superindice'].fillna('')
                                            
                                            # 3. Asegurarse de que las columnas existan ANTES de seleccionarlas
                                            cols_finales = [col for col in cols_to_show if col in df_bi_display.columns]
                                            st.dataframe(df_bi_display[cols_finales].style.hide(axis="index"), use_container_width=True)
                                            # --- FIN DE LA CORRECCIÓN ---

                                # --- REVISED BLOCK FOR INF005 ---
                                elif any(inf in id_infraccion_actual for inf in ['INF001', 'INF005', 'INF007', 'INF008', 'INF009', 'INF011']):
                                    st.markdown("#### Desglose por Extremo")

                                    # Define column mappings and numeric columns
                                    # --- INICIO CORRECCIÓN 2: Columnas CE1 ---
                                    columnas_map_ce1 = {'descripcion': 'Descripción', 'cantidad': 'Cantidad', 'horas': 'Horas', 'precio_soles': 'Precio Base (S/)', 'factor_ajuste': 'Factor Ajuste', 'monto_soles': 'Monto (S/)', 'monto_dolares': 'Monto (US$)'}
                                    cols_num_ce1 = ['Cantidad', 'Horas', 'Precio Base (S/)', 'Factor Ajuste', 'Monto (S/)', 'Monto (US$)']
                                    # --- FIN CORRECCIÓN 2 ---
                                    
                                    columnas_map_ce2 = {'descripcion': 'Descripción', 'precio_soles': 'Precio Base (S/)','precio_dolares': 'Precio Base (US$)', 'factor_ajuste': 'Factor Ajuste', 'monto_soles': 'Monto (S/)','monto_dolares': 'Monto (US$)'}
                                    cols_num_ce2 = ['Precio Base (S/)', 'Precio Base (US$)', 'Factor Ajuste', 'Monto (S/)', 'Monto (US$)']

                                    # --- Data Handling for Simple vs Multiple Extremes ---
                                    if 'extremos' in resultados_app and isinstance(resultados_app['extremos'], list):
                                        extremos_a_mostrar = resultados_app['extremos']
                                        totales_finales = resultados_app.get('totales', {}) # Obtener totales consolidados
                                    else:
                                        extremos_a_mostrar = [resultados_app]
                                        if 'ce1_total_soles' in resultados_app: extremos_a_mostrar[0]['ce1_soles'] = resultados_app.get('ce1_total_soles', 0); extremos_a_mostrar[0]['ce1_dolares'] = resultados_app.get('ce1_total_dolares', 0)
                                        if 'ce2_total_soles_calculado' in resultados_app: extremos_a_mostrar[0]['ce2_soles_calculado'] = resultados_app.get('ce2_total_soles_calculado', 0); extremos_a_mostrar[0]['ce2_dolares_calculado'] = resultados_app.get('ce2_total_dolares_calculado', 0)
                                        if 'ce_total_soles_para_bi' in resultados_app: extremos_a_mostrar[0]['ce_soles_para_bi'] = resultados_app.get('ce_total_soles_para_bi', 0); extremos_a_mostrar[0]['ce_dolares_para_bi'] = resultados_app.get('ce_dolares_para_bi', 0)
                                        if 'aplicar_ce2_a_bi' not in extremos_a_mostrar[0]:
                                            totales_dict = resultados_app.get('totales', resultados_app) 
                                            extremos_a_mostrar[0]['aplicar_ce2_a_bi'] = totales_dict.get('aplicar_ce2_a_bi', False) 
                                        if 'ce1_data_raw' in resultados_app: extremos_a_mostrar[0]['ce1_data'] = resultados_app.get('ce1_data_raw', [])
                                        if 'ce2_data_raw' in resultados_app: extremos_a_mostrar[0]['ce2_data'] = resultados_app.get('ce2_data_raw', [])
                                        if 'bi_data_raw' in resultados_app:
                                            totales_dict = resultados_app.get('totales', resultados_app)
                                            extremos_a_mostrar[0]['bi_data'] = totales_dict.get('bi_data_raw', [])
                                        if 'tipo' not in extremos_a_mostrar[0]: extremos_a_mostrar[0]['tipo'] = "Incumplimiento Único"
                                        
                                        totales_finales = resultados_app.get('totales', resultados_app) # Totales para caso simple
                                    # --- End Data Handling ---

                                    # --- Loop through each extreme ---
                                    for j, extremo_data in enumerate(extremos_a_mostrar):
                                        st.markdown(f"##### Extremo {j + 1}: {extremo_data.get('tipo', 'N/A')}")
                                        aplicar_ce2_bi_extremo = extremo_data.get('aplicar_ce2_a_bi', False)
                                        datos_ce1_crudos = extremo_data.get('ce1_data', [])
                                        datos_ce2_crudos = extremo_data.get('ce2_data', [])

                                        st.markdown("###### CE1: Remisión")
                                        if datos_ce1_crudos:
                                            df_ce1_display = pd.DataFrame(datos_ce1_crudos)
                                            # --- INICIO CORRECCIÓN 2: Añadir Total CE1 ---
                                            total_ce1_s = df_ce1_display['monto_soles'].sum()
                                            total_ce1_d = df_ce1_display['monto_dolares'].sum()
                                            df_ce1_display = pd.concat([df_ce1_display, pd.DataFrame([{'descripcion': 'Total', 'monto_soles': total_ce1_s, 'monto_dolares': total_ce1_d}])], ignore_index=True)
                                            # --- FIN CORRECCIÓN 2 ---
                                            cols_exist_ce1 = [col for col in columnas_map_ce1.keys() if col in df_ce1_display.columns]
                                            if cols_exist_ce1:
                                                df_ce1_display = df_ce1_display[cols_exist_ce1].rename(columns=columnas_map_ce1)
                                                # --- INICIO CORRECCIÓN 2: Formato CE1 ---
                                                df_ce1_display = df_ce1_display.reindex(columns=columnas_map_ce1.values()) # Asegurar todas las columnas
                                                cols_num_exist_ce1 = [col for col in cols_num_ce1 if col in df_ce1_display.columns]
                                                st.dataframe(df_ce1_display.style.format("{:,.3f}", subset=cols_num_exist_ce1, na_rep='').hide(axis="index"), use_container_width=True)
                                                # --- FIN CORRECCIÓN 2 ---
                                            else: st.warning("No hay columnas válidas para mostrar en CE1.")
                                        else:
                                            st.warning("No hay datos de CE1 para este extremo.")

                                        if aplicar_ce2_bi_extremo and datos_ce2_crudos:
                                            st.markdown("###### CE2: Capacitación")
                                            df_ce2_display = pd.DataFrame(datos_ce2_crudos)
                                            # --- INICIO CORRECCIÓN 2: Añadir Total CE2 ---
                                            total_ce2_s = df_ce2_display['monto_soles'].sum()
                                            total_ce2_d = df_ce2_display['monto_dolares'].sum()
                                            df_ce2_display = pd.concat([df_ce2_display, pd.DataFrame([{'descripcion': 'Total', 'monto_soles': total_ce2_s, 'monto_dolares': total_ce2_d}])], ignore_index=True)
                                            # --- FIN CORRECCIÓN 2 ---
                                            cols_exist_ce2 = [col for col in columnas_map_ce2.keys() if col in df_ce2_display.columns]
                                            if cols_exist_ce2:
                                                df_ce2_display = df_ce2_display[cols_exist_ce2].rename(columns=columnas_map_ce2)
                                                # --- INICIO CORRECCIÓN 2: Formato CE2 ---
                                                df_ce2_display = df_ce2_display.reindex(columns=columnas_map_ce2.values()) # Asegurar todas las columnas
                                                cols_num_exist_ce2 = [col for col in cols_num_ce2 if col in df_ce2_display.columns]
                                                st.dataframe(df_ce2_display.style.format("{:,.3f}", subset=cols_num_exist_ce2, na_rep='').hide(axis="index"), use_container_width=True)
                                                # --- FIN CORRECCIÓN 2 ---
                                            else: st.warning("No hay columnas válidas para mostrar en CE2.")

                                        if aplicar_ce2_bi_extremo:
                                            st.markdown("###### Resumen CE del Extremo")
                                            resumen_ext_data = []
                                            ce1_s = extremo_data.get('ce1_soles', 0); ce1_d = extremo_data.get('ce1_dolares', 0)
                                            ce2_s_calc = extremo_data.get('ce2_soles_calculado', 0); ce2_d_calc = extremo_data.get('ce2_dolares_calculado', 0)
                                            total_ce_s = ce1_s + ce2_s_calc; total_ce_d = ce1_d + ce2_d_calc
                                            if ce1_s > 0 or ce1_d > 0: resumen_ext_data.append({'Componente': 'Subtotal Remisión (CE1)', 'Monto (S/)': f"{ce1_s:,.3f}", 'Monto (US$)': f"{ce1_d:,.3f}"})
                                            if ce2_s_calc > 0 or ce2_d_calc > 0: resumen_ext_data.append({'Componente': 'Subtotal Capacitación (CE2)', 'Monto (S/)': f"{ce2_s_calc:,.3f}", 'Monto (US$)': f"{ce2_d_calc:,.3f}"})
                                            resumen_ext_data.append({'Componente': 'Total CE Calculado (Extremo)', 'Monto (S/)': f"{total_ce_s:,.3f}", 'Monto (US$)': f"{total_ce_d:,.3f}"})

                                            if resumen_ext_data: 
                                                # --- INICIO CORRECCIÓN 2: Mostrar US$ en Resumen ---
                                                st.dataframe(pd.DataFrame(resumen_ext_data)[['Componente', 'Monto (S/)', 'Monto (US$)']].style.hide(axis="index"), use_container_width=True)
                                                # --- FIN CORRECCIÓN 2 ---

                                        st.markdown("###### Beneficio Ilícito (BI) del Extremo")
                                        df_bi_extremo = pd.DataFrame(extremo_data.get('bi_data', []))
                                        if not df_bi_extremo.empty:
                                            desc_col_bi = 'descripcion_texto' if 'descripcion_texto' in df_bi_extremo.columns else 'descripcion'
                                            df_bi_display_ext = df_bi_extremo.rename(columns={desc_col_bi: 'Descripción', 'monto': 'Monto'})
                                            cols_show_bi = ['Descripción', 'Monto']
                                            if 'descripcion_superindice' in df_bi_display_ext.columns: # <--- Usar df_bi_display_ext
                                                df_bi_display_ext['Descripción'] = df_bi_display_ext['Descripción'] + df_bi_display_ext['descripcion_superindice'].fillna('')
                                            st.dataframe(df_bi_display_ext[cols_show_bi].style.hide(axis="index"), use_container_width=True)
                                        else:
                                            st.warning("No hay datos de Beneficio Ilícito para mostrar para este extremo.")
                                        st.markdown("---") # Separator between extremes

                                # Lógica general para otras infracciones (sin desglose de extremos específico)
                                else:
                                    totales_finales = resultados_app # Asumir estructura simple
                                    st.markdown("###### Costo Evitado (CE)")
                                    datos_ce_crudos = totales_finales.get('ce_data_raw', [])
                                    if datos_ce_crudos:
                                        df_ce_display = pd.DataFrame(datos_ce_crudos)
                                        # Mapeo genérico (puede necesitar ajustes por infracción)
                                        columnas_map_generico = {
                                            'grupo': 'Grupo', 'subgrupo': 'Subgrupo', 'descripcion': 'Descripción',
                                            'unidad': 'Unidad', 'cantidad': 'Cantidad',
                                            'precio_unitario': 'Precio Unitario (S/)', # Asumiendo Soles
                                            'factor_ajuste': 'Factor Ajuste', 'monto_soles': 'Monto (S/)',
                                            'monto_dolares': 'Monto (US$)'
                                        }
                                        columnas_existentes = [col for col in columnas_map_generico.keys() if col in df_ce_display.columns]
                                        if columnas_existentes:
                                            df_ce_display = df_ce_display[columnas_existentes].rename(columns=columnas_map_generico)
                                            cols_numericas_generico = ['Cantidad', 'Precio Unitario (S/)', 'Factor Ajuste', 'Monto (S/)', 'Monto (US$)']
                                            cols_numericas_existentes = [col for col in cols_numericas_generico if col in df_ce_display.columns]
                                            st.dataframe(df_ce_display.style.format("{:,.3f}", subset=cols_numericas_existentes, na_rep='').hide(axis="index"), use_container_width=True)
                                        else:
                                            st.warning("No hay columnas válidas para mostrar en el Costo Evitado.")
                                    else:
                                        st.warning("No hay datos de Costo Evitado para mostrar.")

                                    st.markdown("###### Beneficio Ilícito (BI)")
                                    df_bi_total = pd.DataFrame(totales_finales.get('bi_data_raw', []))
                                    if not df_bi_total.empty:
                                        desc_col = 'descripcion_texto' if 'descripcion_texto' in df_bi_total.columns else 'descripcion'
                                        df_bi_display = df_bi_total.rename(columns={desc_col: 'Descripción', 'monto': 'Monto'})
                                        cols_to_show = ['Descripción', 'Monto']
                                        if 'descripcion_superindice' in df_bi_total.columns:
                                            df_bi_display['Descripción'] = df_bi_display['Descripción'] + df_bi_total['descripcion_superindice'].fillna('')
                                        st.dataframe(df_bi_display[cols_to_show].style.hide(axis="index"), use_container_width=True)

                                # --- SECCIÓN DE TOTALES Y MULTA (COMÚN Y SEGURA) ---
                                if totales_finales:
                                    st.markdown("---")
                                    st.markdown("#### Totales del Hecho Imputado")

                                    bi_uit_final = totales_finales.get('beneficio_ilicito_uit', 0)
                                    st.metric("Beneficio Ilícito Total (UIT)", f"{bi_uit_final:,.3f}")

                                    st.markdown("###### Multa Propuesta")
                                    df_multa = pd.DataFrame(totales_finales.get('multa_data_raw', []))
                                    if not df_multa.empty:
                                        df_multa_display = df_multa.rename(columns={'Componentes': 'Componentes', 'Monto': 'Monto'})
                                        st.dataframe(df_multa_display.style.hide(axis="index"), use_container_width=True)
                                    
                                    # --- INICIO: (REQ 1) MOSTRAR REDUCCIÓN EN UI (CORREGIDO) ---
                                    # Esta lógica ahora está fuera del 'if' de 'INF004' 
                                    # y se aplica a TODAS las infracciones.
                                    if totales_finales.get('aplica_reduccion') == 'Sí':
                                        porcentaje = totales_finales.get('porcentaje_reduccion', 'N/A')
                                        # Usamos la clave 'multa_con_reduccion_uit' que es la multa intermedia
                                        multa_con_reduccion = totales_finales.get('multa_con_reduccion_uit', 0) 
                                        st.info(f"Se aplica reducción del **{porcentaje}**.")
                                        st.metric("Multa con Reducción (UIT)", f"{multa_con_reduccion:,.3f}")
                                    # --- FIN: (REQ 1) ---

                # --- FIN DEL PASO 3 ---
                
                # --- INICIO: PASO 3.5: ANÁLISIS DE NO CONFISCATORIEDAD (CORREGIDO) ---
                all_steps_complete_check = all('resultados' in d for d in st.session_state.imputaciones_data)
                
                if all_steps_complete_check:
                    st.divider()
                    st.header("Análisis de no confiscatoriedad")
                    
                    if 'confiscatoriedad' not in st.session_state:
                        st.session_state.confiscatoriedad = {'aplica': 'No', 'datos_por_anio': {}}

                    aplica_conf = st.radio(
                        "¿El administrado presentó sus ingresos para el análisis de no confiscatoriedad?",
                        ["No", "Sí"],
                        key='confiscatoriedad_aplica',
                        index=0 if st.session_state.confiscatoriedad.get('aplica', 'No') == 'No' else 1
                    )
                    st.session_state.confiscatoriedad['aplica'] = aplica_conf

                    # --- INICIO CORRECCIÓN NameError (Paso 3.5) ---

                    if aplica_conf == 'Sí':
                        st.info("Sume los montos (Ventas Netas + Otros Ingresos Gravados + Otros Ingresos No Gravados) del año anterior a cada infracción.")
                        
                        # --- INICIO DE LA CORRECCIÓN (Req. 1) ---
                        # Pedir los datos del escrito UNA SOLA VEZ, fuera del bucle
                        
                        datos_conf_global = st.session_state.confiscatoriedad # Referencia al dict
                        
                        st.markdown("##### Documento de Acreditación (Único)")
                        col_e1, col_e2 = st.columns(2)
                        with col_e1:
                            # Guardar en el nivel superior de 'confiscatoriedad'
                            datos_conf_global['escrito_num_conf'] = st.text_input(
                                "N.° de Escrito (Ingresos)",
                                value=datos_conf_global.get('escrito_num_conf', ''),
                                key='conf_escrito_num_global_input'
                            )
                        with col_e2:
                            # Guardar en el nivel superior de 'confiscatoriedad'
                            datos_conf_global['escrito_fecha_conf'] = st.date_input(
                                "Fecha del Escrito (Ingresos)",
                                value=datos_conf_global.get('escrito_fecha_conf'),
                                key='conf_escrito_fecha_global_input',
                                format="DD/MM/YYYY"
                            )
                        # --- FIN DE LA CORRECCIÓN (Req. 1) ---

                        # 1. Identificar todos los años de incumplimiento
                        anios_incumplimiento = set()
                        
                        # --- INICIO DE LA CORRECCIÓN ---
                        # La lógica anterior solo buscaba 'fecha_incumplimiento' (usada por INF008).
                        # Ahora buscamos todas las claves posibles (ej: 'fecha_incumplimiento_extremo' de INF004).
                        
                        for i, datos_hecho in enumerate(st.session_state.imputaciones_data):
                            # Asegurarnos que el hecho tenga extremos antes de iterar
                            extremos_del_hecho = datos_hecho.get('extremos', [])
                            if not extremos_del_hecho:
                                st.warning(f"Hecho {i+1} no tiene extremos. Saltando para análisis de confiscatoriedad.")
                                continue
                                
                            for j, extremo in enumerate(extremos_del_hecho):
                                
                                # Clave 1: Usada por INF008 (y otros)
                                fecha_inc = extremo.get('fecha_incumplimiento') 
                                
                                # Clave 2: Usada por INF004
                                if not fecha_inc:
                                    fecha_inc = extremo.get('fecha_incumplimiento_extremo')
                                
                                # (Puedes añadir más 'elif not fecha_inc:' si otros módulos usan nombres diferentes)

                                if fecha_inc:
                                    try:
                                        # Asegurarnos que sea un objeto 'date' o 'datetime'
                                        anios_incumplimiento.add(fecha_inc.year)
                                    except AttributeError:
                                        st.warning(f"Hecho {i+1}, Extremo {j+1}: La fecha encontrada no es un objeto de fecha válido.")
                                else:
                                    # Si no encontramos NINGUNA clave de fecha
                                    st.warning(f"Hecho {i+1}, Extremo {j+1}: No se pudo encontrar una clave de fecha de incumplimiento ('fecha_incumplimiento' o 'fecha_incumplimiento_extremo').")

                        # --- FIN DE LA CORRECCIÓN ---
                        
                        anios_ordenados = sorted(list(anios_incumplimiento))
                        
                        if not anios_ordenados:
                            # Este error solo saltará si DE VERDAD no se encontró ninguna fecha en ningún hecho
                            st.error("Error: No se pudieron determinar los años de incumplimiento de ningún hecho.")
                        
                        # 2. Pedir ingresos para cada grupo de años
                        datos_por_anio_guardados = st.session_state.confiscatoriedad.get('datos_por_anio', {})
                        
                        for anio_incumplimiento in anios_ordenados:
                            anio_ingresos = anio_incumplimiento - 1
                            st.markdown(f"--- \n**Datos para incumplimientos del año {anio_incumplimiento} (se usarán ingresos de {anio_ingresos}):**")
                            
                            key_s = f"conf_soles_{anio_incumplimiento}"
                            key_a = f"conf_anio_{anio_incumplimiento}"
                            
                            datos_guardados_anio = datos_por_anio_guardados.get(anio_incumplimiento, {})
                            
                            # --- INICIO DE LA CORRECCIÓN (Req. 1) ---
                            # Ahora solo hay 2 columnas en el bucle
                            col_c1, col_c2 = st.columns(2) 
                            
                            with col_c1:
                                ingreso_total_soles = st.number_input(
                                    f"Ingreso Bruto Total {anio_ingresos} (S/)", 
                                    min_value=0.0, 
                                    value=datos_guardados_anio.get('ingreso_total_soles', 0.0),
                                    key=key_s,
                                    format="%.2f"
                                )
                            with col_c2:
                                anio_uit = st.number_input(
                                    f"Año de UIT (para ingresos {anio_ingresos})", 
                                    min_value=2000, 
                                    max_value=date.today().year, 
                                    value=datos_guardados_anio.get('anio_ingresos', anio_ingresos),
                                    key=key_a
                                )
                            
                            # Guardar los datos en el estado de sesión
                            if anio_incumplimiento not in st.session_state.confiscatoriedad['datos_por_anio']:
                                st.session_state.confiscatoriedad['datos_por_anio'][anio_incumplimiento] = {}
                            
                            st.session_state.confiscatoriedad['datos_por_anio'][anio_incumplimiento]['ingreso_total_soles'] = ingreso_total_soles
                            st.session_state.confiscatoriedad['datos_por_anio'][anio_incumplimiento]['anio_ingresos'] = anio_uit
                            # (Se elimina el guardado del N°/Fecha de escrito aquí)
                            # --- FIN DE LA CORRECCIÓN (Req. 1) ---

                            # --- SISTEMA DE GUARDADO UBICADO AL FINAL DE LOS INPUTS ---
                st.divider()
                col_save_title, col_save_btn = st.columns([2, 1])
                with col_save_title:
                    st.markdown("### 💾 Guardar Todo el Avance")
                    st.caption("Guarda hechos, graduación, reducciones y análisis de confiscatoriedad.")
                
                with col_save_btn:
                    if st.button("Guardar Caso en la Nube", type="secondary", use_container_width=True):
                        with st.spinner("Sincronizando con la base de datos..."):
                            # Lista maestra de claves a persistir
                            claves_sesion = [
                                'fecha_emision_informe', 'numero_rsd_base', 'fecha_rsd',
                                'confiscatoriedad', 'rubro_seleccionado', 'imputaciones_data',
                                'numero_ifi', 'fecha_ifi', 'num_informe_multa_ifi',
                                'monto_multa_ifi', 'num_imputaciones_ifi'
                            ]

                            # Limpiar y serializar
                            estado_sucio = {k: st.session_state[k] for k in claves_sesion if k in st.session_state}
                            estado_a_guardar = preparar_datos_para_json(estado_sucio)

                            expediente = st.session_state.num_expediente_formateado
                            producto = st.session_state.info_expediente.get('PRODUCTO', 'IFI')

                            exito, mensaje = guardar_datos_caso(cliente_gspread, expediente, producto, estado_a_guardar)

                        if exito:
                            st.success(f"✅ {mensaje}")
                        else:
                            st.error(f"❌ {mensaje}")
                                        
    # --- PASO 4: GENERAR INFORME FINAL ---
    all_steps_complete = False
    if 'imputaciones_data' in st.session_state and st.session_state.imputaciones_data:
        all_steps_complete = all('resultados' in d for d in st.session_state.imputaciones_data)

    if all_steps_complete:
        st.divider()
        st.header("Paso 4: Generar Informe Final")
        if st.button("🚀 Generar Informe", type="primary"):

            # --- CAMBIO PUNTUAL: Resetear Acrónimos ---
            # Limpiamos la memoria de acrónimos para que el documento se genere limpio desde cero
            if 'context_data' in st.session_state and 'acronyms' in st.session_state.context_data:
                st.session_state.context_data['acronyms'].defined_acronyms = set()
            # ------------------------------------------

            with st.status("Generando informe... Este proceso puede tardar un momento.", expanded=True) as status:
                try:
                    # ETAPA 1: RECOLECCIÓN DE RESULTADOS Y ANEXOS
                    status.update(label="🔄 Etapa 1: Recolectando y ensamblando secciones...")
                    
                    secciones_hechos_listas = []
                    anexos_ce_finales = []
                    ids_anexos_sustento = []
                    
                    df_tipificacion_final = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Tipificacion_Infracciones")

                    for datos_hecho in st.session_state.imputaciones_data:
                        resultados = datos_hecho.get('resultados', {})
                        if not resultados: continue

                        if resultados.get('anexos_ce_generados'):
                            anexos_ce_finales.extend(resultados['anexos_ce_generados'])
                        if resultados.get('ids_anexos'):
                            ids_anexos_sustento.extend(resultados['ids_anexos'])

                        buffer_seccion = io.BytesIO()
                        if 'doc_pre_compuesto' in resultados:
                            buffer_seccion = resultados['doc_pre_compuesto']
                        else:
                            contexto_a_renderizar = resultados.get('contexto_final_word', {})
                            if contexto_a_renderizar:
                                id_infraccion_actual = datos_hecho.get('id_infraccion')
                                fila_infraccion = df_tipificacion_final[df_tipificacion_final['ID_Infraccion'] == id_infraccion_actual]
                                if not fila_infraccion.empty:
                                    id_plantilla_infraccion = fila_infraccion.iloc[0].get('ID_Plantilla_BI')
                                    buffer_plantilla_infraccion = descargar_archivo_drive(id_plantilla_infraccion)
                                    if buffer_plantilla_infraccion:
                                        jinja_env = Environment(trim_blocks=True, lstrip_blocks=True)
                                        doc_seccion_tpl = DocxTemplate(buffer_plantilla_infraccion)
                                        doc_seccion_tpl.render(contexto_a_renderizar, autoescape=True, jinja_env=jinja_env)
                                        doc_seccion_tpl.save(buffer_seccion)
                        
                        secciones_hechos_listas.append(buffer_seccion)

                    # ETAPA 2: ENSAMBLAJE DE PLANTILLAS EN BRUTO
                    context_data = st.session_state.get('context_data', {})
                    info_caso = st.session_state.info_expediente
                    producto_caso = info_caso.get('PRODUCTO', '')
                    
                    df_plantillas = cargar_hoja_a_df(cliente_gspread, NOMBRE_GSHEET_MAESTRO, "Productos")
                    plantilla_row = df_plantillas[df_plantillas['Producto'] == producto_caso]
                    id_plantilla_inicio = plantilla_row.iloc[0].get('ID_Plantilla_Inicio')
                    id_plantilla_fin = plantilla_row.iloc[0].get('ID_Plantilla_Fin')
                    
                    # --- AÑADE ESTA LÍNEA ---
                    id_plantilla_anexo_cap = plantilla_row.iloc[0].get('ID_Plantilla_Anexo_Cap')
                    # --- INICIO: NUEVA PLANTILLA ---
                    id_plantilla_graduacion = plantilla_row.iloc[0].get('ID_Plantilla_Graduacion')
                    # --- FIN ---
                    buffer_inicio_tpl = descargar_archivo_drive(id_plantilla_inicio)
                    buffer_fin_tpl = descargar_archivo_drive(id_plantilla_fin)

                    # Se inicia el ensamblador con la plantilla de inicio (aún sin datos)
                    compositor_final = Composer(Document(buffer_inicio_tpl))
                    spacer_doc = Document()
                    spacer_doc.add_paragraph()

                    # Se añaden las secciones de los hechos (aún con placeholders)
                    for i, buffer_hecho in enumerate(secciones_hechos_listas):
                        if buffer_hecho.getbuffer().nbytes > 0:
                            compositor_final.append(spacer_doc)
                            buffer_hecho.seek(0)
                            compositor_final.append(Document(buffer_hecho))
                    
                    # Se añade la plantilla de fin (aún con placeholders)
                    compositor_final.append(spacer_doc)
                    compositor_final.append(Document(buffer_fin_tpl))
                    
                    # Guardamos este documento ensamblado pero "vacío" en un buffer
                    buffer_documento_completo = io.BytesIO()
                    compositor_final.save(buffer_documento_completo)
                    buffer_documento_completo.seek(0)

                    # ETAPA 3: RENDERIZADO FINAL DEL DOCUMENTO COMPLETO
                    status.update(label="🔄 Etapa 2: Rellenando los datos del informe...")
                    jinja_env = Environment(trim_blocks=True, lstrip_blocks=True)
                    doc_a_renderizar = DocxTemplate(buffer_documento_completo)

                    # a. Preparar datos condicionales (capacitación, extemporáneo)
                    se_usa_capacitacion = any(d.get('resultados', {}).get('usa_capacitacion', False) for d in st.session_state.imputaciones_data)
                    hubo_extemporaneo = any(d.get('resultados', {}).get('es_extemporaneo', False) for d in st.session_state.imputaciones_data)

                    tabla_personal_subdoc = None
                    if se_usa_capacitacion:
                        # Buscamos los datos de la tabla en el primer hecho que los contenga
                        for datos_hecho in st.session_state.imputaciones_data:
                            resultados_app = datos_hecho.get('resultados', {}).get('resultados_para_app', {})
                            
                            # Lógica de búsqueda mejorada
                            tabla_personal_data = None
                            # Primero, intenta la ruta simple
                            if 'tabla_personal_data' in resultados_app:
                                tabla_personal_data = resultados_app['tabla_personal_data']
                            # Si no la encuentra, intenta la ruta anidada para casos múltiples
                            elif 'totales' in resultados_app and 'tabla_personal_data' in resultados_app['totales']:
                                tabla_personal_data = resultados_app['totales']['tabla_personal_data']

                            # Si se encontraron los datos por cualquiera de las dos vías, crea la tabla
                            if tabla_personal_data:
                                # LLAMAMOS A LA NUEVA FUNCIÓN
                                tabla_personal_subdoc = create_personal_table_subdoc(
                                    doc_a_renderizar,
                                    headers=["Perfil", "Descripción", "Cantidad"],
                                    data=tabla_personal_data,
                                    keys=['Perfil', 'Descripción', 'Cantidad'],
                                    texto_posterior='Elaboración: Subdirección de Sanción y Gestión de Incentivos (SSAG) - DFAI.',
                                    column_widths=(2, 3, 1) # Ajusta los anchos a tu gusto
                                )
                                break
                    
                    # b. Preparar tabla de resumen final
                    
                    # --- INICIO: LÓGICA DE CONFISCATORIEDAD POR AÑO ---
                    
                    # 1. Agrupar multas por año de incumplimiento (LÓGICA DE PRORRATEO PRECISA)
                    multas_por_anio = {}
                    mapa_hecho_a_multa_final = {} # {0: 5.2, 1: 3.1, ...}

                    for i, d in enumerate(st.session_state.imputaciones_data):
                        # --- a. Obtener datos del HECHO ---
                        resultados_hecho = d.get('resultados', {}).get('resultados_para_app', {})
                        totales_hecho = resultados_hecho.get('totales', resultados_hecho) # Maneja ambos formatos
                        
                        monto_multa_final_hecho = totales_hecho.get('multa_final_aplicada', 0)
                        mapa_hecho_a_multa_final[i] = monto_multa_final_hecho
                        
                        total_bi_del_hecho = totales_hecho.get('beneficio_ilicito_uit', 0)
                        
                        # --- b. Obtener listas de EXTREMOS (Input y Output) ---
                        input_extremos = d.get('extremos', []) # De session_state, tienen las FECHAS
                        output_extremos = resultados_hecho.get('extremos', []) # De resultados, tienen el BI
                        
                        # --- c. Comprobar si podemos prorratear ---
                        if (not input_extremos or not output_extremos or 
                            len(input_extremos) != len(output_extremos) or 
                            total_bi_del_hecho == 0):
                            
                            # --- d. FALLBACK: Usar lógica antigua (asignar 100% al primer año) ---
                            anio_incumplimiento_hecho = None
                            if input_extremos:
                                primer_extremo = input_extremos[0]
                                fecha_inc = primer_extremo.get('fecha_incumplimiento') or primer_extremo.get('fecha_incumplimiento_extremo')
                                if fecha_inc:
                                    anio_incumplimiento_hecho = fecha_inc.year
                            
                            if anio_incumplimiento_hecho:
                                if anio_incumplimiento_hecho not in multas_por_anio:
                                    multas_por_anio[anio_incumplimiento_hecho] = 0.0
                                multas_por_anio[anio_incumplimiento_hecho] += monto_multa_final_hecho
                        
                        else:
                            # --- e. ÉXITO: Usar lógica NUEVA (prorrateo) ---
                            for j, out_ext in enumerate(output_extremos):
                                in_ext = input_extremos[j]
                                
                                bi_del_extremo = out_ext.get('bi_uit', 0.0)
                                if bi_del_extremo == 0:
                                    continue
                                
                                # Encontrar el año del extremo
                                fecha_inc = in_ext.get('fecha_incumplimiento') or in_ext.get('fecha_incumplimiento_extremo')
                                if not fecha_inc:
                                    continue
                                
                                anio_incumplimiento_extremo = fecha_inc.year
                                
                                # Calcular proporción y multa
                                proporcion_bi = bi_del_extremo / total_bi_del_hecho
                                multa_proporcional = monto_multa_final_hecho * proporcion_bi
                                
                                # Añadir al diccionario
                                if anio_incumplimiento_extremo not in multas_por_anio:
                                    multas_por_anio[anio_incumplimiento_extremo] = 0.0
                                multas_por_anio[anio_incumplimiento_extremo] += multa_proporcional

                    # 1.1. Calcular el total ANTES del tope de confiscatoriedad
                    multa_total_pre_confiscatoriedad = sum(mapa_hecho_a_multa_final.values())
                    # --- FIN DE LA CORRECCIÓN ---

                    # 2. Aplicar lógica de confiscatoriedad por cada año
                    multa_total_uit_final = 0.0
                    textos_confiscatoriedad_word = [] # Almacenará los párrafos para Word
                    conf_data = st.session_state.get('confiscatoriedad', {})

                    # --- INICIO DE LA ADICIÓN ---
                    mapa_anio_a_factor_reduccion = {}
                    # --- FIN DE LA ADICIÓN ---

                    # --- INICIO DE LA CORRECCIÓN ---
                    # Inicializar las variables ANTES del if/else
                    lista_datos_confiscatoriedad = []
                    escrito_num_conf_global = ""
                    escrito_fecha_conf_global = None
                    # --- FIN DE LA CORRECCIÓN ---
                    
                    if conf_data.get('aplica') == 'No':
                        # Si no se aplica, la multa total es la suma simple
                        multa_total_uit_final = sum(mapa_hecho_a_multa_final.values())
                        
                        # --- INICIO DE LA ADICIÓN ---
                        # Si no hay análisis, el factor de reducción es 1 (o sea, 100%)
                        for anio_inc in multas_por_anio.keys():
                            mapa_anio_a_factor_reduccion[anio_inc] = 1.0
                        # --- FIN DE LA ADICIÓN ---
                        
                        textos_confiscatoriedad_word.append(
                            "De acuerdo con lo establecido en el numeral 12.2 del artículo 12° del RPAS, la multa a imponerse (...) no puede ser mayor al diez por ciento (10 %) del ingreso bruto anual (...). "
                            "No obstante, el administrado no ha presentado la documentación que acredite los ingresos brutos anuales del año anterior a la comisión de la infracción, por lo que no se ha podido realizar el análisis de no confiscatoriedad."
                        )
                    
                    else:
                        # Si SÍ se aplica, iteramos por cada año de incumplimiento
                        df_uit = datos_calculo.get('df_uit')
                        datos_ingresos_ui = conf_data.get('datos_por_anio', {})
                        
                        lista_datos_confiscatoriedad = []
                        escrito_num_conf_global = conf_data.get('escrito_num_conf', '')
                        escrito_fecha_conf_global = conf_data.get('escrito_fecha_conf')
                        
                        for anio_inc, multa_sumada_del_anio in multas_por_anio.items():
                            ingresos_de_este_grupo = datos_ingresos_ui.get(anio_inc, {})
                            ingreso_total_soles = ingresos_de_este_grupo.get('ingreso_total_soles', 0.0)
                            anio_ingresos_uit = ingresos_de_este_grupo.get('anio_ingresos', anio_inc - 1)
                            
                            uit_de_ese_anio = 0.0
                            if df_uit is not None and anio_ingresos_uit > 0:
                                uit_row = df_uit[df_uit['Año_UIT'] == anio_ingresos_uit]
                                if not uit_row.empty:
                                    uit_de_ese_anio = float(uit_row.iloc[0]['Valor_UIT'])

                            if ingreso_total_soles > 0 and uit_de_ese_anio > 0:
                                ingreso_bruto_uit = ingreso_total_soles / uit_de_ese_anio
                                tope_10_porciento_uit = redondeo_excel(ingreso_bruto_uit * 0.10, 3)
                                
                                multa_final_del_anio = multa_sumada_del_anio
                                texto_parrafo = ""
                                
                                # --- INICIO DE LA CORRECCIÓN ---
                                # 1. Determinar si es confiscatoria
                                es_confiscatoria_este_anio = (multa_sumada_del_anio > tope_10_porciento_uit)
                                # --- FIN DE LA CORRECCIÓN ---

                                if es_confiscatoria_este_anio:
                                    # Aplicar el tope
                                    multa_final_del_anio = tope_10_porciento_uit
                                    
                                    # --- INICIO DE LA ADICIÓN ---
                                    # Calcular el factor de reducción (ej. 0.590 / 1.237 = 0.477)
                                    mapa_anio_a_factor_reduccion[anio_inc] = multa_final_del_anio / multa_sumada_del_anio if multa_sumada_del_anio > 0 else 0
                                    # --- FIN DE LA ADICIÓN ---

                                    texto_parrafo = (
                                        f"Para los hechos imputados con incumplimiento en el año {anio_inc}, la suma de las multas ({multa_sumada_del_anio:,.3f} UIT) "
                                        f"supera el tope del 10% de los ingresos brutos del año {anio_ingresos_uit} (equivalente a {tope_10_porciento_uit:,.3f} UIT). "
                                        f"Por lo tanto, la multa para este grupo de hechos se topa a {multa_final_del_anio:,.3f} UIT."
                                    )
                                else:
                                    # No aplicar el tope
                                    texto_parrafo = (
                                        f"Para los hechos imputados con incumplimiento en el año {anio_inc}, la suma de las multas ({multa_sumada_del_anio:,.3f} UIT) "
                                        f"no supera el tope del 10% de los ingresos brutos del año {anio_ingresos_uit} (equivalente a {tope_10_porciento_uit:,.3f} UIT). "
                                        f"Por lo tanto, se mantiene la multa calculada de {multa_final_del_anio:,.3f} UIT."
                                    )
                                
                                    # --- INICIO DE LA ADICIÓN ---
                                    # No hay reducción, el factor es 1.0
                                    mapa_anio_a_factor_reduccion[anio_inc] = 1.0
                                    # --- FIN DE LA ADICIÓN ---

                                textos_confiscatoriedad_word.append(texto_parrafo)
                                multa_total_uit_final += multa_final_del_anio
                            
                                # --- INICIO DE LA CORRECCIÓN ---
                                # 2. Añadir los placeholders faltantes al diccionario del bucle
                                datos_para_bucle = {
                                    'anio_ingresos': anio_ingresos_uit,
                                    'ingreso_bruto_total_s': f"S/ {ingreso_total_soles:,.2f}",
                                    'uit_anio_ingreso': f"S/ {uit_de_ese_anio:,.2f}",
                                    'ingreso_bruto_uit': f"{ingreso_bruto_uit:,.3f} UIT",
                                    'tope_10_porciento_uit': f"{tope_10_porciento_uit:,.3f} UIT",
                                    
                                    # --- PLACEHOLDERS FALTANTES AÑADIDOS ---
                                    'anio_incumplimiento': anio_inc,
                                    'multa_total_del_anio': f"{multa_sumada_del_anio:,.3f} UIT",
                                    'es_confiscatoria': es_confiscatoria_este_anio
                                }
                                lista_datos_confiscatoriedad.append(datos_para_bucle)
                                # --- FIN DE LA CORRECCIÓN ---
                            else:
                                # ... (lógica de 'else' no cambia) ...
                                multa_total_uit_final += multa_sumada_del_anio # Sumar la multa sin topar
                                
                                # --- INICIO DE LA ADICIÓN ---
                                # No hay datos, no hay reducción, el factor es 1.0
                                mapa_anio_a_factor_reduccion[anio_inc] = 1.0
                                # --- FIN DE LA ADICIÓN ---
                                
                                textos_confiscatoriedad_word.append(
                                    f"Para los hechos imputados con incumplimiento en el año {anio_inc}, no se ingresaron datos de ingresos o UIT válidos. "
                                    f"Por lo tanto, se asume la multa de {multa_sumada_del_anio:,.3f} UIT sin análisis de confiscatoriedad."
                                )

                    # 3. Crear los placeholders para Word
                    # (RichText maneja los saltos de línea \n)
                    context_data['texto_confiscatoriedad_final'] = RichText("\n".join(textos_confiscatoriedad_word))
                    context_data['aplica_confiscatoriedad'] = conf_data.get('aplica') == 'Sí'
                    
                    # 4. Crear la tabla resumen (LÓGICA CORREGIDA)
                    
                    # --- INICIO DE LA CORRECCIÓN ---
# --- 4a. Crear un nuevo mapa de multas finales por HECHO, aplicando los factores por AÑO ---
                    mapa_hecho_a_multa_final_topped = {}
                    
                    for i, d in enumerate(st.session_state.imputaciones_data):
                        # Obtener la multa *original* del hecho (ej. 1.237 UIT)
                        monto_multa_original_hecho = mapa_hecho_a_multa_final.get(i, 0)
                        if monto_multa_original_hecho == 0:
                            mapa_hecho_a_multa_final_topped[i] = 0.0
                            continue
                        
                        # Obtener datos de BI y extremos para decidir si prorrateamos
                        resultados_hecho = d.get('resultados', {}).get('resultados_para_app', {})
                        totales_hecho = resultados_hecho.get('totales', resultados_hecho)
                        total_bi_del_hecho = totales_hecho.get('beneficio_ilicito_uit', 0)
                        
                        input_extremos = d.get('extremos', [])
                        output_extremos = resultados_hecho.get('extremos', [])

                        multa_final_para_este_hecho = 0.0

                        # --- CORRECCIÓN LÓGICA: CASO SIMPLE VS MÚLTIPLE ---
                        # Si es un solo extremo o el BI es 0, no se puede prorratear; se asigna al año del hecho.
                        if (len(input_extremos) <= 1 or not output_extremos or 
                            len(input_extremos) != len(output_extremos) or 
                            total_bi_del_hecho == 0):
                            
                            anio_incumplimiento_hecho = None
                            if input_extremos:
                                primer_extremo = input_extremos[0]
                                # Buscar la fecha en las distintas claves posibles según el módulo
                                fecha_inc = primer_extremo.get('fecha_incumplimiento') or primer_extremo.get('fecha_incumplimiento_extremo')
                                if fecha_inc: 
                                    anio_incumplimiento_hecho = fecha_inc.year
                            
                            # Obtener el factor de reducción para ese año (calculado en el análisis de confiscatoriedad)
                            factor_reduccion = mapa_anio_a_factor_reduccion.get(anio_incumplimiento_hecho, 1.0)
                            multa_final_para_este_hecho = monto_multa_original_hecho * factor_reduccion
                        
                        else:
                            # CASO MÚLTIPLE: Prorrateo por el BI de cada extremo
                            for j, out_ext in enumerate(output_extremos):
                                in_ext = input_extremos[j]
                                bi_del_extremo = out_ext.get('bi_uit', 0.0)
                                if bi_del_extremo == 0: continue
                                
                                # Encontrar el año del extremo específico
                                fecha_inc_ext = in_ext.get('fecha_incumplimiento') or in_ext.get('fecha_incumplimiento_extremo')
                                if not fecha_inc_ext: continue
                                
                                anio_incumplimiento_extremo = fecha_inc_ext.year
                                
                                # Obtener el factor de reducción para el año de este extremo
                                factor_reduccion = mapa_anio_a_factor_reduccion.get(anio_incumplimiento_extremo, 1.0)
                                
                                # Calcular qué parte de la multa corresponde a este extremo y aplicar su tope
                                proporcion_bi = bi_del_extremo / total_bi_del_hecho
                                multa_original_aporte_extremo = monto_multa_original_hecho * proporcion_bi
                                multa_final_para_este_hecho += (multa_original_aporte_extremo * factor_reduccion)

                        mapa_hecho_a_multa_final_topped[i] = multa_final_para_este_hecho
                    
                    # --- 4b. Crear las filas de la tabla resumen final ---
                    summary_rows = []
                    for i, monto_final_topped in mapa_hecho_a_multa_final_topped.items():
                        summary_rows.append({
                            'Numeral': f"IV.{i + 2}", 
                            'Infracciones': f"Hecho imputado n.° {i + 1}", 
                            'Monto': f"{monto_final_topped:,.3f} UIT"
                        })

                    # Añadir la fila de Total (esta ya usaba el valor final correcto)
                    summary_rows.append({'Numeral': 'Total', 'Infracciones': '', 'Monto': f"{multa_total_uit_final:,.3f} UIT"})
                    
                    tabla_resumen_final_subdoc = create_summary_table_subdoc(
                        doc_a_renderizar, ["Numeral", "Infracciones", "Monto"], summary_rows, 
                        ['Numeral', 'Infracciones', 'Monto'],
                        texto_posterior="Elaboración: Subdirección de Sanción y Gestión de Incentivos (SSAG) - DFAI.",
                        column_widths=(1, 4, 1.5)
                    )
                    
                    # --- INICIO: Calcular años anteriores a incumplimientos (CORREGIDO) ---
# --- INICIO: Calcular años anteriores a incumplimientos (CORRECCIÓN v2) ---
                    anios_incumplimiento_set = set()
                    for datos_hecho in st.session_state.imputaciones_data:
                        
                        # La fecha está en los 'extremos', que es la ESTRUCTURA DE INPUTS guardada en session_state
                        extremos_del_hecho = datos_hecho.get('extremos', []) 
                        
                        if not extremos_del_hecho:
                            # Fallback para infracciones simples (que guardan la fecha en el nivel superior)
                            fecha_inc_simple = datos_hecho.get('fecha_incumplimiento')
                            if fecha_inc_simple:
                                try:
                                    anios_incumplimiento_set.add(fecha_inc_simple.year - 1)
                                except AttributeError:
                                    pass # Ignorar si no es una fecha
                            continue # Pasar al siguiente hecho

                        # Si hay extremos (como en INF004), iterar sobre ellos
                        for extremo in extremos_del_hecho:
                            
                            # Clave 1: Usada por INF008 (y otros)
                            fecha_inc = extremo.get('fecha_incumplimiento') 
                            
                            # Clave 2: Usada por INF004
                            if not fecha_inc:
                                fecha_inc = extremo.get('fecha_incumplimiento_extremo')
                            
                            if fecha_inc:
                                try:
                                    # 'fecha_inc.year' es el año de incumplimiento
                                    # 'fecha_inc.year - 1' es el año de ingresos (que necesitamos para el placeholder)
                                    anios_incumplimiento_set.add(fecha_inc.year - 1)
                                except AttributeError:
                                    pass # Ignorar si no es un objeto fecha válido

                    # Formatear la salida
                    anios_ordenados = sorted(list(anios_incumplimiento_set))
                    anios_str_lista = [str(a) for a in anios_ordenados]

                    if len(anios_str_lista) == 0:
                        anios_incumplimiento_texto = "" # O un valor por defecto, ej. "N/A"
                    elif len(anios_str_lista) == 1:
                        anios_incumplimiento_texto = anios_str_lista[0]
                    elif len(anios_str_lista) == 2:
                        anios_incumplimiento_texto = " y ".join(anios_str_lista)
                    else: # 3 o más años
                        anios_incumplimiento_texto = ", ".join(anios_str_lista[:-1]) + " y " + anios_str_lista[-1]
                    # --- FIN: Calcular años ---

                    # --- INICIO: (REQ 2) Lógica de Escenarios Periódicos ---
                    
                    periodicas_ids = ['INF001', 'INF002', 'INF005', 'INF007', 'INF008', 'INF004']
                    
                    hechos_periodicos = []
                    hechos_no_periodicos = []

                    # 1. Clasificar todos los hechos (solo guardamos el número como string)
                    for i, datos_hecho in enumerate(st.session_state.imputaciones_data):
                        id_infraccion = datos_hecho.get('id_infraccion', '')
                        num_hecho_solo = f"{i + 1}" # Guardamos solo el número
                        
                        if any(pid in id_infraccion for pid in periodicas_ids):
                            hechos_periodicos.append(num_hecho_solo)
                        else:
                            hechos_no_periodicos.append(num_hecho_solo)
                    
                    # 2. Construir el texto final
                    texto_final_escenario = ""
                    
                    # --- Párrafo para HECHOS PERIÓDICOS (Escenario 2) ---
                    if hechos_periodicos:
                        # 1. Llamamos a la función con los prefijos correctos ("al" / "a los")
                        # Nota: Pasa los números con "n.° " ya incluido para que la función los una
                        lista_con_tag = [f"n.° {n}" for n in hechos_periodicos]
                        hechos_listos = formatear_lista_hechos(
                            lista_con_tag, 
                            singular_prefix="al hecho imputado", 
                            plural_prefix="a los hechos imputados"
                        )
                        
                        # 2. Definimos el verbo según la cantidad
                        verbo = "corresponden" if len(hechos_periodicos) > 1 else "corresponde"
                        
                        # 3. Construimos el párrafo final
                        texto_final_escenario += (
                            f"Sobre ello, respecto {hechos_listos}, de la revisión de los documentos "
                            f"obrantes en el presente PAS, se advierte que, el administrado se encontraría en "
                            f"un escenario del tipo 2, toda vez que habría realizado actividades iguales o "
                            f"semejantes a los costos evitados asociados a las obligaciones incumplidas, dado "
                            f"que {verbo} a incumplimientos formales y/u obligaciones periódicas. No "
                            f"obstante, hasta la emisión del presente informe, el administrado no ha presentado "
                            f"ningún comprobante de pago ni factura para poder ser evaluada."
                        )

                    # --- Párrafo para hechos NO PERIÓDICOS (Incierto) ---
                    if hechos_no_periodicos:
                        if texto_final_escenario:
                            texto_final_escenario += "\n\n"
                        
                        # 1. Preparamos la lista con el formato 'n.° X'
                        lista_con_tag_np = [f"n.° {n}" for n in hechos_no_periodicos]
                        
                        # 2. Llamamos a la función pasando los prefijos para que no use los de por defecto
                        hechos_listos_np = formatear_lista_hechos(
                            lista_con_tag_np, 
                            singular_prefix="al hecho imputado", 
                            plural_prefix="a los hechos imputados"
                        )
                        
                        # 3. Construimos la frase (ya no necesitamos el prefix_np ni el replace)
                        texto_final_escenario += (
                            f"Además, de la revisión de los documentos obrantes en el presente PAS, en relación "
                            f"{hechos_listos_np}, no se tiene información suficiente para determinar en qué "
                            f"escenario se encontraría el administrado, toda vez que, hasta la emisión del presente informe, "
                            f"no ha presentado ningún comprobante de pago, ni factura ni boletas para poder ser evaluadas."
                        )
                    # Guardar como RichText para la plantilla Word
                    context_data['hechos_escenario'] = RichText(texto_final_escenario)
                    
                    # --- FIN: (REQ 2) ---

                    # c. Construir el contexto final completo
                    
                    # --- INICIO: Nueva Lógica de Resumen de Capacitación (v3 - Separada) ---
                    
                    hechos_con_capacitacion_info = []
                    mapa_anio_a_extremos = {}
                    tabla_personal_subdoc_final = None
                    num_personal_total_final = 0
                    
                    # --- INICIO: (NUEVO) Lista para las TABLAS (v2) ---
                    lista_tablas_capacitacion_subdoc = [] # Almacenará los objetos de tabla
                    # --- FIN: (NUEVO) ---
                    
                    # 1. Recopilar datos
                    for i, datos_hecho in enumerate(st.session_state.imputaciones_data):
                        resultados = datos_hecho.get('resultados', {})
                        if resultados.get('usa_capacitacion', False):
                            hecho_num_str = f"n.° {i + 1}"
                            hechos_con_capacitacion_info.append({'num_hecho_str': hecho_num_str})
                            
                            # --- INICIO DE LA MODIFICACIÓN ---
                            # 1. Contar cuántos extremos tiene este hecho en TOTAL
                            total_extremos_del_hecho = len(datos_hecho.get('extremos', []))
                            # --- FIN DE LA MODIFICACIÓN ---

                            if tabla_personal_subdoc_final is None:
                                tabla_personal_subdoc_final = resultados.get('tabla_detalle_personal')
                                num_personal_total_final = int(datos_hecho.get('num_personal_capacitacion', 0))

                            for j, extremo in enumerate(datos_hecho.get('extremos', [])):
                                if isinstance(extremo, dict) and extremo.get('fecha_incumplimiento'):
                                    anio = extremo['fecha_incumplimiento'].year
                                    if anio not in mapa_anio_a_extremos:
                                        mapa_anio_a_extremos[anio] = []
                                    mapa_anio_a_extremos[anio].append({
                                        'hecho': hecho_num_str,
                                        'extremo': f"extremo n.° {j + 1}",
                                        # --- INICIO DE LA MODIFICACIÓN ---
                                        # 2. Guardar ese conteo en el mapa
                                        'total_extremos_del_hecho': total_extremos_del_hecho
                                        # --- FIN DE LA MODIFICACIÓN ---
                                    })

                    # 2. Inicializar placeholders
                    placeholder_hechos_caps = ""
                    lista_bullets_prorrateo = []
                    total_personal_capacitacion_str = ""
                    
                    # --- INICIO: (NUEVO) Lista para las TABLAS (v2) ---
                    lista_tablas_capacitacion_subdoc = [] # Almacenará los objetos de tabla
                    # --- FIN: (NUEVO) ---
                    
                    # 3. Construir los valores solo si se usa capacitación
                    if hechos_con_capacitacion_info:
                        
                        # --- REQUISITO 1: Placeholder de lista de hechos ---
                        num_hechos_total = len(hechos_con_capacitacion_info)
                        lista_hechos_str = [h['num_hecho_str'] for h in hechos_con_capacitacion_info]
                        
                        if num_hechos_total == 1:
                            placeholder_hechos_caps = f"la conducta infractora {lista_hechos_str[0]}"
                        elif num_hechos_total == 2:
                            placeholder_hechos_caps = f"las conductas infractoras {lista_hechos_str[0]} y {lista_hechos_str[1]}"
                        else:
                            placeholder_hechos_caps = "las conductas infractoras " + ", ".join(lista_hechos_str[:-1]) + " y " + lista_hechos_str[-1]
                        
                        # --- REQUISITO 3: Placeholder de lista de bullets ---
                        for anio in sorted(mapa_anio_a_extremos.keys()):
                            lista_extremos_del_anio = mapa_anio_a_extremos[anio]
                            num_extremos_en_anio = len(lista_extremos_del_anio)
                            
                            texto_bullet = ""
                            if num_extremos_en_anio == 1:
                                # --- INICIO DE LA CORRECCIÓN v2 ---
                                info_extremo = lista_extremos_del_anio[0]
                                hecho_str = info_extremo['hecho']
                                
                                # 1. Comprobar el número total de extremos guardado
                                if info_extremo.get('total_extremos_del_hecho', 1) == 1:
                                    # Caso A: Hecho simple (1 extremo en total) -> No mencionar extremo
                                    texto_bullet = (
                                        f"Una (1) capacitación correspondiente al año {anio}, asociada al hecho imputado {hecho_str}, "
                                        f"dado que su incumplimiento se produjo en dicho período."
                                    )
                                else:
                                    # Caso B: Hecho múltiple (>1 extremo) -> Sí mencionar extremo
                                    extremo_str = info_extremo['extremo']
                                    texto_bullet = (
                                        f"Una (1) capacitación correspondiente al año {anio}, asociada al {extremo_str} del hecho imputado {hecho_str}, "
                                        f"dado que su incumplimiento se produjo en dicho período."
                                    )
                                # --- FIN DE LA CORRECCIÓN v2 ---
                            else: # <-- INDENTACIÓN CORREGIDA
                                hechos_unicos_en_anio = sorted(list(set(e['hecho'] for e in lista_extremos_del_anio)))
                                if len(hechos_unicos_en_anio) == 1:
                                    hecho_str = hechos_unicos_en_anio[0]
                                    num_extremos_texto = texto_con_numero(num_extremos_en_anio, genero='m')
                                    texto_bullet = (
                                        f"Una (1) capacitación correspondiente al año {anio}, prorrateada entre los {num_extremos_texto} extremos del hecho imputado {hecho_str}, "
                                        f"dado que las fechas de incumplimiento corresponden al mismo año, dotando así de mayor razonabilidad al cálculo efectuado y "
                                        f"evitando la duplicidad de costos."
                                    )
                                else:
                                    if len(hechos_unicos_en_anio) == 2:
                                        lista_hechos_str_anio = " y ".join(hechos_unicos_en_anio)
                                    else:
                                        lista_hechos_str_anio = ", ".join(hechos_unicos_en_anio[:-1]) + " y " + hechos_unicos_en_anio[-1]
                                    
                                    texto_bullet = (
                                        f"Una (1) capacitación correspondiente al año {anio}, prorrateada entre los hechos imputados {lista_hechos_str_anio}, "
                                        f"dado que las fechas de incumplimiento corresponden al mismo año, dotando así de mayor razonabilidad al cálculo efectuado y "
                                        f"evitando la duplicidad de costos."
                                    )
                            lista_bullets_prorrateo.append({'texto_bullet': texto_bullet}) # <-- INDENTACIÓN CORREGIDA
                        
                        # --- REQUISITO 2: Placeholder de total de personal ---
                        if num_personal_total_final == 1:
                            total_personal_capacitacion_str = f"{texto_con_numero(num_personal_total_final, genero='f')} persona"
                        else:
                            total_personal_capacitacion_str = f"{texto_con_numero(num_personal_total_final, genero='f')} personas"

                        # --- INICIO: (NUEVO) Generar textos para Título y Placeholder ---
                        # --- INICIO: (NUEVO) Generar textos para Título y Placeholder ---
                        texto_hechos_para_tabla = formatear_lista_hechos(lista_hechos_str)
                        titulo_para_tabla = f"Costo de capacitación para {texto_hechos_para_tabla}"
                        # --- FIN: (NUEVO) ---
                        
                        # --- INICIO: (NUEVO) Construir la data de la tabla detallada ---
                        from funciones import create_capacitacion_table_subdoc # Asegúrate de importar la función
                        
                        num_anios = len(mapa_anio_a_extremos.keys())
                        for i, anio in enumerate(sorted(mapa_anio_a_extremos.keys())):
                            
                            # --- INICIO: (NUEVO) Lógica de TABLA individual por año ---
                            
                            # A. Obtener datos (igual que antes)
                            lista_hechos = mapa_anio_a_extremos[anio]
                            es_prorrateado = len(lista_hechos) > 1
                            
                            # (Esta lógica ya existe en tu código, la usamos para obtener el precio base)
                            ce2_data_raw = []
                            primer_hecho_info = lista_hechos[0]
                            idx_hecho = int(primer_hecho_info['hecho'].replace('n.° ', '')) - 1
                            idx_extremo = int(primer_hecho_info['extremo'].replace('extremo n.° ', '')) - 1
                            datos_hecho_cap = st.session_state.imputaciones_data[idx_hecho]
                            resultados_app_cap = datos_hecho_cap.get('resultados', {}).get('resultados_para_app', {})
                            
                            if 'extremos' in resultados_app_cap and isinstance(resultados_app_cap['extremos'], list): 
                                if idx_extremo < len(resultados_app_cap['extremos']):
                                    ce2_data_raw = resultados_app_cap['extremos'][idx_extremo].get('ce2_data', [])
                            else: 
                                totales_simple = resultados_app_cap.get('totales', resultados_app_cap)
                                ce2_data_raw = totales_simple.get('ce2_data_raw', [])
                            
                            
                            # B. Construir las filas de la tabla (SOLO para este año)
                            tabla_data_este_anio = []
                            if ce2_data_raw:
                                base_item = ce2_data_raw[0]
                                
                                # --- CORRECCIÓN MATEMÁTICA DE PRORRATEO ---
                                # Obtenemos el precio ya prorrateado (Ej: 325)
                                precio_prorrateado_usd = base_item.get('precio_dolares', 0)
                                num_hechos_prorrateo = len(lista_hechos)
                                
                                # Reconstruimos el precio base REAL (Ej: 325 * 2 = 650)
                                precio_base_real_usd = precio_prorrateado_usd * num_hechos_prorrateo if es_prorrateado else precio_prorrateado_usd
                                # -------------------------------------------

                                tabla_data_este_anio.append({
                                    'descripcion': f"AÑO {anio}", 'unidad': '', 'cantidad': '', 'precio_usd': ''
                                })
                                
                                tabla_data_este_anio.append({
                                    'descripcion': f"Capacitación ({total_personal_capacitacion_str}) 1/",
                                    'unidad': 'glb.', 'cantidad': "1",
                                    'precio_usd': f"US$ {precio_base_real_usd:,.3f}"
                                })
                                tabla_data_este_anio.append({
                                    'descripcion': 'Total', 'unidad': '', 'cantidad': '', 
                                    'precio_usd': f"US$ {precio_base_real_usd:,.3f}"
                                })

                                if es_prorrateado:
                                    for h_info in lista_hechos:
                                        # --- CORRECCIÓN VISUAL: Añadir Extremo a la Fila ---
                                        if h_info.get('total_extremos_del_hecho', 1) == 1:
                                            desc_fila = f"Hecho imputado {h_info['hecho']}"
                                        else:
                                            desc_fila = f"Hecho imputado {h_info['hecho']} - {h_info['extremo'].capitalize()}"
                                        # ---------------------------------------------------

                                        tabla_data_este_anio.append({
                                            'descripcion': desc_fila,
                                            'unidad': '', 'cantidad': f"{1/num_hechos_prorrateo:.2%}",
                                            'precio_usd': f"US$ {precio_prorrateado_usd:,.3f}"
                                        })

                            # C. Generar el placeholder de texto para el TÍTULO y NOTA AL PIE
                            # --- CORRECCIÓN VISUAL: Títulos y notas exactas e inteligentes ---
                            hechos_dict = {}
                            for e in lista_hechos:
                                h = e['hecho']
                                ext = e['extremo'].replace('extremo ', '') # "n.° 1"
                                total_ext = e.get('total_extremos_del_hecho', 1)
                                if h not in hechos_dict:
                                    hechos_dict[h] = {'extremos': [], 'total': total_ext}
                                hechos_dict[h]['extremos'].append(ext)

                            partes_texto = []
                            for h, info in hechos_dict.items():
                                exts = info['extremos']
                                if info['total'] == 1:
                                    partes_texto.append(f"el hecho imputado {h}")
                                elif len(exts) == 1:
                                    partes_texto.append(f"el extremo {exts[0]} del hecho imputado {h}")
                                else:
                                    exts_str = " y ".join(exts) if len(exts) == 2 else ", ".join(exts[:-1]) + " y " + exts[-1]
                                    partes_texto.append(f"los extremos {exts_str} del hecho imputado {h}")

                            if len(partes_texto) == 1:
                                texto_placeholder_tabla = partes_texto[0]
                            else:
                                texto_placeholder_tabla = ", ".join(partes_texto[:-1]) + " y " + partes_texto[-1]
                            # --------------------------------------------------
                            
                            titulo_tabla_individual = f"Costo de capacitación para {texto_placeholder_tabla}"

                            # D. Crear el objeto subdocumento y añadirlo a la lista
                            headers_cap = ["Descripción", "Unidad", "Cantidad", "Precio (US$) 2/"]
                            keys_cap = ['descripcion', 'unidad', 'cantidad', 'precio_usd']
                            
                            if tabla_data_este_anio:
                                tabla_subdoc_individual = create_capacitacion_table_subdoc(
                                    doc_a_renderizar, 
                                    headers_cap, 
                                    tabla_data_este_anio, 
                                    keys_cap,
                                    title_text=titulo_tabla_individual,
                                    hechos_placeholder=texto_placeholder_tabla # <-- Pasa el texto específico
                                )
                                lista_tablas_capacitacion_subdoc.append(tabla_subdoc_individual)
                            # --- FIN: (NUEVO) Lógica de TABLA individual por año ---
                    
                    # --- FIN: Nueva Lógica de Resumen de Capacitación (v3 - Separada) ---

                    # --- LÓGICA DE REDUCCIÓN GLOBAL ACTUALIZADA ---
                    hechos_con_reduccion_list = []
                    aplica_50 = False
                    aplica_30 = False
                    primer_memo_num = ""
                    primer_memo_fecha = None
                    primer_escrito_num = ""
                    primer_escrito_fecha = None
                    
                    for i, datos_hecho in enumerate(st.session_state.imputaciones_data):
                        if datos_hecho.get('aplica_reduccion') == 'Sí':
                            hechos_con_reduccion_list.append(f"n.° {i + 1}")
                            
                            # Detectamos el porcentaje específico del hecho
                            porcentaje = datos_hecho.get('porcentaje_reduccion')
                            if porcentaje == "50%":
                                aplica_50 = True
                            elif porcentaje == "30%":
                                aplica_30 = True
                            
                            # Capturamos datos del primer documento de sustento encontrado
                            if not primer_memo_num:
                                primer_memo_num = datos_hecho.get('memo_num', '')
                                primer_memo_fecha = datos_hecho.get('memo_fecha')
                                primer_escrito_num = datos_hecho.get('escrito_num', '')
                                primer_escrito_fecha = datos_hecho.get('escrito_fecha')

                    # 2. Crear el placeholder booleano (para el 'if')
                    aplica_reduccion_global = len(hechos_con_reduccion_list) > 0
                    
                    # 3. Crear el placeholder de texto (usando la función que ya existe)
                    texto_hechos_con_reduccion = ""
                    if aplica_reduccion_global:
                        lista_formateada = formatear_lista_hechos(
                            hechos_con_reduccion_list, 
                            singular_prefix="hecho imputado", 
                            plural_prefix="hechos imputados"
                        )
                        # Añadir la 'a' gramatical
                        if len(hechos_con_reduccion_list) == 1:
                            texto_hechos_con_reduccion = f"al {lista_formateada}"
                        else:
                            texto_hechos_con_reduccion = f"a los {lista_formateada}"
                            
                    # --- FORMATEAR LAS FECHAS CAPTURADAS ---
                    memo_fecha_formateada = (format_date(primer_memo_fecha, "d 'de' MMMM 'de' yyyy", locale='es') if primer_memo_fecha else '').replace("septiembre", "setiembre").replace("Septiembre", "Setiembre")
                    escrito_fecha_formateado = (format_date(primer_escrito_fecha, "d 'de' MMMM 'de' yyyy", locale='es') if primer_escrito_fecha else '').replace("septiembre", "setiembre").replace("Septiembre", "Setiembre")
                    # --- FIN: (NUEVO) LÓGICA DE REDUCCIÓN GLOBAL (v2) ---

                    # --- Lógica existente para construir lista_hechos_para_plantilla ---
                    lista_hechos_para_plantilla = []
                    for i, datos_hecho in enumerate(st.session_state.imputaciones_data):
                        lista_hechos_para_plantilla.append({
                            "numero_imputado": i + 1,
                            "descripcion": datos_hecho.get('texto_hecho', '')
                        })

                    hubo_extemporaneo = any(d.get('resultados', {}).get('es_extemporaneo', False) for d in st.session_state.imputaciones_data)
                    se_usa_capacitacion = any(d.get('resultados', {}).get('usa_capacitacion', False) for d in st.session_state.imputaciones_data)
                    
                    # d. Construir el contexto final completo
                    contexto_final_completo = {
                        **context_data, 
                        'lista_hechos_imputados': lista_hechos_para_plantilla, 
                        'tabla_resumen_final': tabla_resumen_final_subdoc, 
                        'mt_uit': f"{multa_total_uit_final:,.3f} UIT",
                        'multa_total_pre_conf': f"{multa_total_pre_confiscatoriedad:,.3f} UIT",
                        'hubo_cumplimiento_extemporaneo': hubo_extemporaneo,
                        'se_usa_capacitacion': se_usa_capacitacion, 
                        'anios_incumplimiento': anios_incumplimiento_texto,
                        
                        # --- PLACEHOLDERS SEPARADOS (REQ 1, 2, 3) ---
                        'lista_hechos_capacitacion': placeholder_hechos_caps,
                        'total_personal_capacitacion': total_personal_capacitacion_str,
                        'lista_bullets_prorrateo_caps': lista_bullets_prorrateo,
                        # --- FIN ---
                        # --- INICIO: (NUEVO) Añadir la tabla detallada al contexto ---
                        'lista_tablas_capacitacion': lista_tablas_capacitacion_subdoc,
                        # --- FIN: (NUEVO) ---
                        'tabla_detalle_personal': tabla_personal_subdoc_final,
                        'texto_explicacion_prorrateo': RichText(""), # Se mantiene vacío
                        # --- INICIO: AÑADIR NUEVOS PLACEHOLDERS GLOBALES ---
                        'aplica_reduccion_global': aplica_reduccion_global,
                        'aplica_reduccion_50': aplica_50, # Nueva variable para el IF del 50%
                        'aplica_reduccion_30': aplica_30,
                        'lista_hechos_con_reduccion': texto_hechos_con_reduccion,
                        'memo_num_global': primer_memo_num,
                        'memo_fecha_global': memo_fecha_formateada,
                        'escrito_num_global': primer_escrito_num,
                        'escrito_fecha_global': escrito_fecha_formateado,
                        # --- FIN: AÑADIR NUEVOS PLACEHOLDERS GLOBALES ---
                        # --- INICIO: (REQ. 3) AÑADIR LA NUEVA LISTA PARA EL BUCLE ---
                        'lista_datos_confiscatoriedad': lista_datos_confiscatoriedad if conf_data.get('aplica') == 'Sí' else [],
                        # --- FIN: (REQ. 3) ---
                        # --- INICIO: (REQ. 2) AÑADIR NUEVOS PLACEHOLDERS GLOBALES ---
                        'conf_escrito_num_global': escrito_num_conf_global,
                        'conf_escrito_fecha_global': (format_date(escrito_fecha_conf_global, "d 'de' MMMM 'de' yyyy", locale='es') if escrito_fecha_conf_global else '').replace("septiembre", "setiembre").replace("Septiembre", "Setiembre")
                        # --- FIN: (REQ. 2) ---
                    }
                    
                    # e. Renderizar
                    doc_a_renderizar.render(contexto_final_completo, autoescape=True, jinja_env=jinja_env)

                    buffer_renderizado = io.BytesIO()
                    doc_a_renderizar.save(buffer_renderizado)
                    buffer_renderizado.seek(0)
                    
                    # ETAPA 4: POST-PROCESAMIENTO DE NUMERACIÓN
                    from funciones import NumberingManager, post_process_numbering
                    doc_a_numerar = Document(buffer_renderizado)
                    # Usamos el MISMO gestor de numeración para que continúe la secuencia
                    final_numbering_manager = context_data['numbering']
                    
                    # --- INICIO DE LA CORRECCIÓN ---
                    final_numbering_manager.table_count = 0 # Reiniciar el contador a 0
                    # --- FIN DE LA CORRECCIÓN ---
                    
                    post_process_numbering(doc_a_numerar, final_numbering_manager)

                    # El compositor final ahora es el documento ya numerado
                    compositor_final_numerado = Composer(doc_a_numerar)

                    # ETAPA 5: AÑADIR ANEXOS
                    status.update(label="📑 Etapa 3: Añadiendo anexos y finalizando...")
                    
                    # --- INICIO: LÓGICA DE CONTADOR DE ANEXOS ---
                    anexo_counter = 1
                    
                    # (Datos 'se_usa_capacitacion' y 'anexos_ce_finales' ya existen desde ETAPA 1 y 3)

                    # --- INICIO: LÓGICA DE ANEXO COMPARTIDO (Capacitación + CE) ---
                    
                    # 1. Comprobar si este anexo (Capacitación o CE) existe
                    if se_usa_capacitacion or anexos_ce_finales:
                    
                        # 2. Imprimir el Título COMPARTIDO (UNA SOLA VEZ)
                        compositor_final_numerado.doc.add_page_break()
                        h_cap_ce = compositor_final_numerado.doc.add_heading(level=1)
                        h_cap_ce.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_cap_ce = h_cap_ce.add_run(f"Anexo n.° {anexo_counter}")
                        run_cap_ce.underline = True
                        compositor_final_numerado.doc.add_paragraph() # Espacio

                        # 3. Añadir la plantilla de CAPACITACIÓN (si existe)
                        if se_usa_capacitacion:
                            if not id_plantilla_anexo_cap:
                                st.error("Error: 'ID_Plantilla_Anexo_Cap' no encontrado. Revisa el Paso 1.")
                                st.stop()

                            status.update(label=f"📑 Etapa 3a: Generando Anexo {anexo_counter} (Capacitación)...")
                            buffer_anexo_cap_tpl = descargar_archivo_drive(id_plantilla_anexo_cap)
                            if not buffer_anexo_cap_tpl:
                                st.error("Error: No se pudo descargar la plantilla del anexo de capacitación.")
                                st.stop()
                            
                            doc_anexo_cap = DocxTemplate(buffer_anexo_cap_tpl)
                            
                            # (contexto_final_completo ya tiene 'tabla_detalle_personal', 'lista_bullets_prorrateo_caps', etc. de la ETAPA 3)
                            doc_anexo_cap.render(contexto_final_completo, autoescape=True, jinja_env=jinja_env)
                            
                            buffer_cap_final = io.BytesIO()
                            doc_anexo_cap.save(buffer_cap_final)
                            buffer_cap_final.seek(0)
                            
                            # Insertar la plantilla de capacitación
                            compositor_final_numerado.append(Document(buffer_cap_final))
                            
                            # Añadir un salto de página si también hay anexos CE
                            if anexos_ce_finales:
                                compositor_final_numerado.doc.add_page_break()

                        # 4. Añadir las plantillas CE (si existen)
                        if anexos_ce_finales:
                            # (Ya NO se añade título aquí, porque se puso el título compartido arriba)
                            for i, anexo_ce_buffer in enumerate(anexos_ce_finales):
                                anexo_ce_buffer.seek(0)
                                compositor_final_numerado.append(Document(anexo_ce_buffer))
                                if i < len(anexos_ce_finales) - 1: 
                                    compositor_final_numerado.doc.add_page_break()

                        # 5. Incrementar el contador de anexos
                        anexo_counter += 1
                    # --- FIN: LÓGICA DE ANEXO COMPARTIDO ---

                    # --- INICIO FASE 4: ANEXO DE GRADUACIÓN ---
                    for i, datos_hecho in enumerate(st.session_state.imputaciones_data):
                        # Verificar si este hecho tiene graduación activada
                        if datos_hecho.get('aplica_graduacion') == 'Sí':
                            
                            if not id_plantilla_graduacion:
                                st.warning(f"Advertencia: El hecho {i+1} tiene graduación, pero no se encontró 'ID_Plantilla_Graduacion' en la hoja Productos.")
                                continue

                            status.update(label=f"📑 Generando Anexo de Graduación para Hecho {i+1}...")
                            
                            buffer_tpl_grad = descargar_archivo_drive(id_plantilla_graduacion)
                            if buffer_tpl_grad:
                                doc_tpl_grad = DocxTemplate(buffer_tpl_grad)
                                
                                # --- PREPARACIÓN DINÁMICA DE DATOS PARA EL ANEXO ---
                                grad_data = datos_hecho.get('graduacion', {})
                                factor_f = datos_hecho.get('factor_f_calculado', 1.0)
                                
                                # 1. Inicializar contexto con datos básicos
                                contexto_grad = {
                                    'ph_hecho_numero': i + 1,
                                    # Req: Formato "1.46 (146%)" para el total
                                    'ph_factor_f_final_completo': f"{factor_f:,.2f} ({factor_f:.0%})",
                                    # NUEVO: Solo el porcentaje (ej. 130%)
                                    'ph_factor_f_solo_porcentaje': f"{factor_f:.0%}",
                                    'ph_suma_f_total': f"{sum(grad_data.get(f'subtotal_f{k}', 0) for k in range(1,8)):.0%}"
                                }

                                # 2. Extraer valores individuales (1.1, 1.2, etc.) y subtotales (f1, f2...)
                                for f_key, f_info in FACTORES_GRADUACION.items():
                                    # A. Subtotal del factor (f1, f2, etc.)
                                    val_subtotal = grad_data.get(f"subtotal_{f_key}", 0.0)
                                    contexto_grad[f"ph_{f_key}_subtotal"] = f"{val_subtotal:.0%}"
                                    
                                    # B. Criterios individuales (1.1, 1.2...)
                                    # El orden de los criterios en FACTORES_GRADUACION define el índice (1, 2, 3...)
                                    for idx_crit, crit_label in enumerate(f_info["criterios"].keys(), 1):
                                        # Construimos la llave que guardamos en la interfaz
                                        key_valor_interfaz = f"grad_{i}_{f_key}_{crit_label}_valor"
                                        valor_num = grad_data.get(key_valor_interfaz, 0.0)
                                        
                                        # Creamos el placeholder: ph_f1_1_valor, ph_f1_2_valor, etc.
                                        tag_name = f"ph_{f_key}_{idx_crit}_valor"
                                        contexto_grad[tag_name] = f"{valor_num:.0%}"

                                # 3. Renderizar y procesar
                                doc_tpl_grad.render(contexto_grad)
                                buffer_grad_final = io.BytesIO()
                                doc_tpl_grad.save(buffer_grad_final)
                                buffer_grad_final.seek(0)

                                # --- Insertar en el documento maestro ---
                                compositor_final_numerado.doc.add_page_break()
                                
                                # Título del Anexo
                                h2_grad = compositor_final_numerado.doc.add_heading(level=1)
                                h2_grad.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run_grad = h2_grad.add_run(f"Anexo n.° {anexo_counter}")
                                run_grad.underline = True
                                
                                compositor_final_numerado.doc.add_paragraph() # Espacio
                                
                                compositor_final_numerado.append(Document(buffer_grad_final))
                                
                                anexo_counter += 1
                            else:
                                st.error(f"Error al descargar la plantilla de graduación ({id_plantilla_graduacion}).")
                    # --- FIN FASE 4 ---

                    # --- INICIO: Anexo de COSTOS REFERENCIALES (Sustento) ---
                    if ids_anexos_sustento:
                        lista_ids_anexos = list(dict.fromkeys(ids_anexos_sustento))
                        compositor_final_numerado.doc.add_page_break()
                        
                        # Título dinámico (usando el contador)
                        h2 = compositor_final_numerado.doc.add_heading(level=1)
                        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run1 = h2.add_run(f"Anexo n.° {anexo_counter}") # <--- Título dinámico
                        run1.underline = True
                        h2.add_run("\n") 
                        run2 = h2.add_run("(Cotizaciones y costos referenciales)")
                        run2.underline = False
                        
                        compositor_final_numerado.doc.add_paragraph()

                        for i, file_id in enumerate(lista_ids_anexos):
                    # --- FIN: Corrección Anexo 2 ---
                            anexo_drive_buffer = descargar_archivo_drive(file_id)
                            if anexo_drive_buffer:
                                compositor_final_numerado.append(Document(anexo_drive_buffer))
                                if i < len(lista_ids_anexos) - 1: compositor_final_numerado.doc.add_page_break()

                        # (Opcional: anexo_counter += 1 si hubiera más anexos después)
                    # --- FIN: Anexo de COSTOS REFERENCIALES ---

                    # ETAPA 6: GUARDAR Y DESCARGAR
                    final_buffer = io.BytesIO()
                    compositor_final_numerado.save(final_buffer)
                    final_buffer.seek(0) # <-- Asegúrate que el buffer DOCX esté rebobinado

                    status.update(label="¡Informe generado con éxito!", state="complete", expanded=False)

                    # Botón de descarga para el archivo WORD (.docx) - SIN CAMBIOS
                    nombre_exp = st.session_state.get('num_expediente_formateado', 'EXPEDIENTE_SIN_NUMERO')
                    st.download_button(
                        label="✅ Descargar Informe Final (.docx)", # <-- Etiqueta clara
                        data=final_buffer.getvalue(), # <-- Sigue usando el buffer del DOCX
                        file_name=f"Informe_Multa_{nombre_exp.replace('/', '-')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )
                    st.success("¡Informe final generado con éxito!")

                except Exception as e:
                    st.error(f"Ocurrió un error al generar el documento: {e}")
                    st.exception(e)

if not cliente_gspread:
    st.error(
        "🔴 No se pudo establecer la conexión con Google Sheets. Revisa el archivo de credenciales y la conexión a internet.")