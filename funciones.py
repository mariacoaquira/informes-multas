import io
import os
import re
import pandas as pd
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import streamlit as st 
from num2words import num2words
from docxtpl import DocxTemplate, Subdoc, RichText
from datetime import date
import io
import os
import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st 
from copy import deepcopy
from num2words import num2words
import math # <--- AÑADIR ESTA LÍNEA
from decimal import Decimal, ROUND_HALF_UP

# Importaciones de docxcompose
from docxcompose.composer import Composer

# Importaciones de docxtpl
from docxtpl import DocxTemplate, Subdoc, RichText

# Importaciones de python-docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.enum.style import WD_STYLE_TYPE

# --------------------------------------------------------------------
#  FUNCIONES DE WORD
# --------------------------------------------------------------------

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_graduation_table_subdoc(tpl, headers, data, keys, texto_posterior="", column_widths=(5.7, 0.5)):
    sd = tpl.new_subdoc()
    table = sd.add_table(rows=1, cols=len(headers))
    table.style = 'TablaCuerpo' 

    # 1. Configurar anchos (Calificación más estrecha: 0.5)
    for i, width in enumerate(column_widths):
        table.columns[i].width = Inches(width)

    # Cabecera
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 3. Fuente Arial 10 Bold en cabecera
        run = p.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True
        
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

    # Datos
    for i, item in enumerate(data):
        row = table.add_row()
        for j, key in enumerate(keys):
            cell = row.cells[j]
            cell.text = str(item.get(key, ""))
            p = cell.paragraphs[0]
            
            # 3. Forzar Fuente Arial 10 en celdas
            if p.runs:
                run = p.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            
            if j == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Lógica de bordes (f1-f7)
            if i < 7:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                if i < 6:
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'nil')
                    tcBorders.append(bottom)
                if i > 0:
                    top = OxmlElement('w:top')
                    top.set(qn('w:val'), 'nil')
                    tcBorders.append(top)
                tcPr.append(tcBorders)

            if i >= 7: # Totales en negrita y fondo
                if p.runs: p.runs[0].bold = True
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D9D9D9')
                cell._tc.get_or_add_tcPr().append(shading_elm)

    # 4. Texto posterior con estilo configurado (Arial 8)
    if texto_posterior:
        p = sd.add_paragraph(texto_posterior)
        p.style = 'FuenteTabla'  

    return sd

def create_capacitacion_table_subdoc(doc_template, headers, data, keys, title_text=None, hechos_placeholder=None):
    """
    Crea la tabla de prorrateo de capacitación con formato especial,
    replicando el estilo de 'image_53eb40.png'.
    """
    sub = doc_template.new_subdoc()

    # --- INICIO: AÑADIR TÍTULO DINÁMICO ---
    if title_text:
        p_title = sub.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(title_text)
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(10)
        run_title.bold = True
    # --- FIN: AÑADIR TÍTULO DINÁMICO ---

    table = sub.add_table(rows=1, cols=len(headers))
    table.style = 'TablaAnexo' 

    # --- Formato Cabecera (Bordes superior e inferior) ---
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]; p = cell.paragraphs[0]; p.clear()
        run = p.add_run(header_text); run.font.name = 'Arial'; run.font.size = Pt(8); run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top={"sz": 4, "val": "single"}, bottom={"sz": 4, "val": "single"})
    
    # --- Formato Filas de Datos ---
    for item in data:
        row_cells = table.add_row().cells
        desc_val = str(item.get(keys[0], ''))
        
        # Determinar el formato de la fila
        is_total_row = "Total" in desc_val
        is_year_row = "AÑO" in desc_val
        is_hecho_row = "Hecho imputado" in desc_val

        for col_idx, key in enumerate(keys):
            cell_text = str(item.get(key, ''))
            cell = row_cells[col_idx]; p = cell.paragraphs[0]; p.clear()
            run = p.add_run(cell_text); run.font.name = 'Arial'; run.font.size = Pt(8)
            
            if is_total_row or is_year_row:
                run.bold = True
            
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if is_hecho_row: p.paragraph_format.left_indent = Inches(0.25)
            elif col_idx == 1 or col_idx == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            if is_year_row:
                set_cell_border(cell, top={"sz": 2, "val": "single"}, bottom={"sz": 2, "val": "single"})
            elif is_total_row:
                set_cell_border(cell, top={"sz": 2, "val": "single"}, bottom={"sz": 4, "val": "single"})
            else:
                set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"})
    
# --- INICIO: CORRECCIÓN DE ESTILO DE FUENTE ---
    
    # Añadir texto de fuente fijo (formato manual)
    p_fuente_titulo = sub.add_paragraph()
    run_fuente = p_fuente_titulo.add_run("Fuente:") 
    run_fuente.font.name = 'Arial'
    run_fuente.font.size = Pt(8)
    
    p_f1 = sub.add_paragraph("1/ Costos sobre servicios de capacitación virtual en temas de desarrollo organizacional y en el cumplimiento de obligaciones ambientales fiscalizables para el sector público y privado alcanzados por Win Work Perú S.A.C el 1 de junio de 2020 mediante carta s/n, con Registro OEFA n.° 2020-E01-036926. (Ver Anexo n.° 2)")
    for run in p_f1.runs: # Aplicar a todo el párrafo
        run.font.name = 'Arial'
        run.font.size = Pt(8)

    # --- INICIO: MODIFICAR TEXTO DE FUENTE (2) ---
    texto_hechos_final = hechos_placeholder if hechos_placeholder else "los hechos imputados correspondientes"
    texto_f2 = (
        f"2/ Es importante precisar que el costo de capacitación corresponde a la fecha de costeo. "
        f"Posteriormente, al calcular el costo evitado para {texto_hechos_final}, "
        f"este será actualizado tomando en cuenta el IPC y el TC correspondiente al período de "
        f"incumplimiento de la infracción previamente mencionada."
    )
    p_f2 = sub.add_paragraph(texto_f2)
    for run in p_f2.runs: # Aplicar a todo el párrafo
        run.font.name = 'Arial'
        run.font.size = Pt(8)
    # --- FIN: MODIFICAR TEXTO DE FUENTE (2) ---
    
    p_elab = sub.add_paragraph("Elaboración: Subdirección de Sanción y Gestión de Incentivos (SSAG) – DFAI.")
    for run in p_elab.runs: # Aplicar a todo el párrafo
        run.font.name = 'Arial'
        run.font.size = Pt(8)
    
    # --- FIN: CORRECCIÓN DE ESTILO DE FUENTE ---
    
    return sub


def redondeo_excel(n, decimales=0):
    """
    Implementa el redondeo aritmético (round half up) usando la librería Decimal
    para evitar errores de punto flotante.
    """
    try:
        # Convertimos a string primero para que Decimal capture el valor
        # exacto que vemos, no la representación interna del float.
        n_str = str(n)
        
        # '1e-3' es lo mismo que '0.001'
        quantizer = Decimal('1e-' + str(decimales)) 
        
        # ROUND_HALF_UP es el redondeo aritmético (si es 5, redondea hacia arriba)
        rounded_decimal = Decimal(n_str).quantize(quantizer, rounding=ROUND_HALF_UP)
        
        # Devolver como float para compatibilidad con el resto del código
        return float(rounded_decimal)
        
    except Exception:
        # Fallback por si n no es un número (aunque no debería pasar)
        multiplicador = 10 ** decimales
        return math.floor(n * multiplicador + 0.5) / multiplicador


def formatear_lista_hechos(lista_hechos, singular_prefix="el hecho imputado", plural_prefix="los hechos imputados"):
    """
    Formatea una lista de números de hecho (ej: ['n.° 1', 'n.° 2']) en un texto
    gramaticalmente correcto.
    
    Ejemplos:
    - ['n.° 1'] -> "el hecho imputado n.° 1"
    - ['n.° 1', 'n.° 2'] -> "los hechos imputados n.° 1 y n.° 2"
    - ['n.° 1', 'n.° 2', 'n.° 3'] -> "los hechos imputados n.° 1, n.° 2 y n.° 3"
    """
    num_hechos = len(lista_hechos)
    
    if num_hechos == 0:
        return "" # No debería pasar, pero es un control
        
    if num_hechos == 1:
        return f"{singular_prefix} {lista_hechos[0]}"
        
    if num_hechos == 2:
        return f"{plural_prefix} {lista_hechos[0]} y {lista_hechos[1]}"
        
    # Para 3 o más
    else:
        # Une todos menos el último con ", "
        primera_parte = ", ".join(lista_hechos[:-1])
        # Añade el último con " y "
        return f"{plural_prefix} {primera_parte} y {lista_hechos[-1]}"
    

def get_initials_from_name(full_name_str, to_lower=False):
    """
    Obtiene las iniciales de un nombre completo.
    Ej: "Ricardo Machuca" -> "RM"
    Ej: "Edith Vasquez" -> "ev" (si to_lower=True)
    """
    if not full_name_str or not isinstance(full_name_str, str):
        return ""
    parts = full_name_str.split()
    if not parts:
        return ""
    # Asegurarse de que part no esté vacío antes de tomar la inicial
    initials = "".join(part[0] for part in parts if part) 
    return initials.lower() if to_lower else initials.upper()


def format_decimal_dinamico(val, decimals=3):
    """
    Formatea un número: aplica redondeo Excel (redondeo_excel) con el número
    de decimales especificado, y luego formatea dinámicamente.
    """
    if val is None: return ''
    try:
        # Aseguramos que el valor es numérico
        num = float(val)
        
        # 1. Aplicar el redondeo de Excel (la regla crítica)
        # Asumimos que redondeo_excel está disponible/importado en funciones.py
        num_rounded = redondeo_excel(num, decimals)

        # 2. Aplicar formato dinámico al valor YA REDONDEADO
        if num_rounded == int(num_rounded):
            return f"{int(num_rounded)}" # Retorna como entero sin decimales
        else:
            # Genera la cadena de formato dinámicamente (ej: "{:,.3f}")
            # Se usa coma como separador de miles si es necesario, y punto decimal
            format_str = f"{{:,.{decimals}f}}" 
            
            # Formatea, quita ceros sobrantes y el punto decimal si no quedan dígitos
            return format_str.format(num_rounded).rstrip('0').rstrip('.')
            
    except (ValueError, TypeError):
        # Maneja casos donde val no es un número
        return str(val)
    

def rt_con_superindice(texto_base, texto_super):
    """
    Crea un objeto RichText con un texto base
    y un texto en superíndice (exponente).
    
    Args:
        texto_base (str): El texto normal.
        texto_super (str): El texto que irá como superíndice.
        
    Returns:
        RichText: Un objeto RichText listo para la plantilla.
    """
    rt = RichText(str(texto_base)) # Texto normal
    rt.add(str(texto_super), superscript=True) # Texto en superíndice
    return rt

def create_main_table_coercitiva(doc_template, headers, data, keys, texto_posterior=None, estilo_texto_posterior=None, column_widths=None):
    """
    Crea una tabla principal específica para informes de coercitiva.
    - Maneja saltos de línea en las celdas.
    - Formatea en negrita y con sombreado cualquier fila donde la primera celda contenga la palabra "Total".
    """
    sub = doc_template.new_subdoc()
    table = sub.add_table(rows=1, cols=len(headers))
    table.style = 'TablaCuerpo'
    SHADING_COLOR = "D9D9D9"

    if column_widths:
        if len(column_widths) == len(headers):
            for i, width in enumerate(column_widths):
                if i < len(table.columns): table.columns[i].width = Inches(width)

    # Formatear Encabezado
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        if i < len(hdr_cells):
            cell = hdr_cells[i]; p = cell.paragraphs[0]; p.clear()
            run = p.add_run(header_text); run.font.name = 'Arial'; run.font.size = Pt(10); run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, SHADING_COLOR)

    # Formatear Filas de Datos
    for item in data:
        row_cells = table.add_row().cells
        is_total_row = 'total' in str(item.get(keys[0], '')).lower()

        for col_idx, key in enumerate(keys):
            if col_idx < len(row_cells):
                cell = row_cells[col_idx]
                p = cell.paragraphs[0]; p.clear()
                cell_value = item.get(key, '')

                # --- INICIO CORRECCIÓN Saltos de Línea ---
                if isinstance(cell_value, str) and '\n' in cell_value:
                    parts = cell_value.split('\n')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        run.font.name = 'Arial'; run.font.size = Pt(10)
                        if is_total_row: run.bold = True # Negrita para total
                        if i < len(parts) - 1:
                            run.add_break() # Añade un salto de línea de Word
                else:
                    run = p.add_run(str(cell_value))
                    run.font.name = 'Arial'; run.font.size = Pt(10)
                    if is_total_row: run.bold = True # Negrita para total
                # --- FIN CORRECCIÓN ---

                if is_total_row and col_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif col_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Alinear montos a la derecha

                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if is_total_row:
                    set_cell_shading(cell, SHADING_COLOR) # Sombreado para total
    
    if texto_posterior:
        style_to_use = estilo_texto_posterior if estilo_texto_posterior else 'Normal'
        try:
             if style_to_use in doc_template.styles: sub.add_paragraph(texto_posterior, style=style_to_use)
             else: p = sub.add_paragraph(texto_posterior)
        except Exception: p = sub.add_paragraph(texto_posterior)

    return sub


def create_considerations_table_subdoc(doc_template, headers, data, keys, texto_posterior=None, estilo_texto_posterior=None, column_widths=None):
    """
    Crea la tabla de consideraciones de muestreo con el estilo 'TablaCuerpo2'
    y permite personalizar el texto de elaboración y su estilo.
    """
    sub = doc_template.new_subdoc()

    table = sub.add_table(rows=1, cols=len(headers))
    table.style = 'TablaCuerpo2'
    SHADING_COLOR = "D9D9D9"

    if column_widths:
        if len(column_widths) == len(headers):
            for i, width in enumerate(column_widths):
                table.columns[i].width = Inches(width)

    # --- Formato del Encabezado ---
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header_text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr_cells[i], SHADING_COLOR)

    # --- Formato de las Filas de Datos ---
    for item in data:
        row_cells = table.add_row().cells
        for col_idx, key in enumerate(keys):
            p = row_cells[col_idx].paragraphs[0]
            p.clear()
            cell_text = str(item.get(key, ''))
            run = p.add_run(cell_text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- Añadir texto de elaboración (DINÁMICO) ---
    # Si no se envía texto_posterior, usa el valor por defecto
    txt_final = texto_posterior if texto_posterior else "Elaboración: Subdirección de Sanción y Gestión de Incentivos (SSAG) - DFAI."
    style_final = estilo_texto_posterior if estilo_texto_posterior else 'FuenteTabla'
    
    p = sub.add_paragraph()
    run = p.add_run(txt_final)
    p.style = style_final
    
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0)

    return sub

def formatear_periodo_monitoreo(fecha, frecuencia):
    """
    Convierte una fecha en un texto de periodo según la frecuencia especificada
    (Trimestral, Semestral o Anual).
    """
    if not isinstance(fecha, date) or not frecuencia:
        return "N/A"
    
    mes = fecha.month
    año = fecha.year
    
    if frecuencia == "Trimestral":
        if 1 <= mes <= 3:
            periodo = "Primer trimestre"
        elif 4 <= mes <= 6:
            periodo = "Segundo trimestre"
        elif 7 <= mes <= 9:
            periodo = "Tercer trimestre"
        else:
            periodo = "Cuarto trimestre"
        return f"{periodo} {año}"
        
    elif frecuencia == "Semestral":
        if 1 <= mes <= 6:
            periodo = "Primer semestre"
        else:
            periodo = "Segundo semestre"
        return f"{periodo} {año}"

    elif frecuencia == "Anual":
        return f"Anual {año}"
        
    else:
        return "Frecuencia no especificada"

from docx.enum.text import WD_ALIGN_PARAGRAPH # For paragraph alignment (e.g., center, left)
from docx.enum.table import WD_ALIGN_VERTICAL # For vertical alignment *in tables*
from docx.shared import Pt, Inches
import pandas as pd
# Asumo que set_cell_border y format_decimal_dinamico están importados o definidos en funciones.py

def create_ce2_envio_table_subdoc(doc_template, data, total_soles, total_dolares):
    """
    Tabla específica para CE2 Envío:
    - Sin columna Unidad.
    - Columna Descripción más amplia.
    """
    sub = doc_template.new_subdoc()
    
    headers = ["Descripción", "Cantidad", "Precio Unitario (S/)", "Factor de Ajuste 2/", "Monto (*) (S/)", "Monto (*) (US$) 3/"]
    num_cols = len(headers)
    table = sub.add_table(rows=1, cols=num_cols)
    table.style = 'TablaAnexo'
    
    # --- Anchos Personalizados ---
    table.autofit = False 
    table.allow_autofit = False
    
    widths = [
        Inches(1.9),  # Descripción (Amplia)
        Inches(0.7),  # Cantidad
        Inches(0.8),  # Precio
        Inches(0.7),  # Factor
        Inches(0.9),  # Monto S/
        Inches(0.9)   # Monto US$
    ]
    
    for i, width in enumerate(widths):
        table.columns[i].width = width
        table.rows[0].cells[i].width = width

    # Encabezados
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header_text); run.font.name = 'Arial'; run.font.size = Pt(8); run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(hdr_cells[i], top={"sz": 4, "val": "single"}, bottom={"sz": 4, "val": "single"})

    # Datos
    keys = ['descripcion', 'cantidad', 'precio_unitario', 'factor_ajuste', 'monto_soles', 'monto_dolares']
    
    for item in data:
        row_cells = table.add_row().cells
        for i, key in enumerate(keys):
            row_cells[i].width = widths[i]
            val = item.get(key)
            
            # Formato condicional
            if key == 'descripcion':
                cell_text = str(val)
                align = WD_ALIGN_PARAGRAPH.LEFT
            elif key == 'cantidad':
                cell_text = format_decimal_dinamico(val, 0)
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif key == 'precio_unitario':
                # Forzar 3 decimales fijos para el precio
                cell_text = f"S/ {float(val):,.3f}" if pd.notna(val) else "S/ 0.000"
                align = WD_ALIGN_PARAGRAPH.RIGHT
            elif key == 'monto_soles':
                cell_text = f"S/ {format_decimal_dinamico(val, 3)}"
                align = WD_ALIGN_PARAGRAPH.RIGHT
            else: # Factor
                cell_text = format_decimal_dinamico(val, 3)
                align = WD_ALIGN_PARAGRAPH.CENTER

            p = row_cells[i].paragraphs[0]
            run = p.add_run(cell_text); run.font.name = 'Arial'; run.font.size = Pt(8)
            p.alignment = align
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for edge in ['top', 'bottom']: set_cell_border(row_cells[i], **{edge: {"val": "nil"}})

    # Totales
    row_cells = table.add_row().cells
    p_total = row_cells[0].paragraphs[0]; run = p_total.add_run('Total'); run.bold = True; run.font.name = 'Arial'; run.font.size = Pt(8); p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sol = row_cells[4].paragraphs[0]; run = p_sol.add_run(f"S/ {total_soles:,.3f}"); run.bold = True; run.font.name = 'Arial'; run.font.size = Pt(8); p_sol.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_dol = row_cells[5].paragraphs[0]; run = p_dol.add_run(f"US$ {total_dolares:,.3f}"); run.bold = True; run.font.name = 'Arial'; run.font.size = Pt(8); p_dol.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    for cell in row_cells: set_cell_border(cell, top={"sz": 4, "val": "single"}, bottom={"sz": 4, "val": "single"})

    return sub


def create_ce2_lab_table_subdoc(doc_template, data, total_soles, total_dolares, nombre_matriz):
    """
    Tabla específica para CE2 Laboratorio:
    - Fila de título de Matriz al inicio.
    - Columnas detalladas: Parámetros, Puntos, Reportes, Precio Total, etc.
    """
    sub = doc_template.new_subdoc()
    
    headers = ["Parámetros 1/", "n.° de puntos", "n.° de reportes", "Precio Unitario (S/)", "Precio Total (S/)", "Factor de Ajuste 2/", "Monto (*) (S/)", "Monto (*) (US$) 3/"]
    num_cols = len(headers)
    table = sub.add_table(rows=1, cols=num_cols)
    table.style = 'TablaAnexo'
    
    # --- Anchos ---
    table.autofit = False; table.allow_autofit = False
    widths = [Inches(1.0), Inches(0.6), Inches(0.7), Inches(0.7), Inches(0.7), Inches(0.6), Inches(0.8), Inches(0.8)]
    for i, w in enumerate(widths): table.columns[i].width = w; table.rows[0].cells[i].width = w

    # Encabezados
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header_text); run.font.name = 'Arial'; run.font.size = Pt(8); run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(hdr_cells[i], top={"sz": 4, "val": "single"}, bottom={"sz": 4, "val": "single"})

    # --- Fila de MATRIZ (Aire, Ruido, etc.) ---
    if nombre_matriz:
        row_cells = table.add_row().cells
        merged_cell = row_cells[0].merge(row_cells[num_cols - 1])
        p = merged_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0)
        run = p.add_run(f"{nombre_matriz.upper()}")
        run.font.name = 'Arial'; run.font.size = Pt(8); run.bold = True
        for cell in row_cells: set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"})

    # Datos
    keys = ['descripcion', 'cantidad', 'reportes', 'precio_unitario', 'precio_total', 'factor_ajuste', 'monto_soles', 'monto_dolares']
    
    for item in data:
        row_cells = table.add_row().cells
        for i, key in enumerate(keys):
            row_cells[i].width = widths[i]
            val = item.get(key)
            
            if key == 'descripcion':
                cell_text = str(val)
                align = WD_ALIGN_PARAGRAPH.LEFT
                indent = Inches(0.1) # Sangría ligera para los parámetros
            elif key in ['cantidad', 'reportes']:
                cell_text = format_decimal_dinamico(val, 0)
                align = WD_ALIGN_PARAGRAPH.CENTER
                indent = Inches(0)
            elif key == 'factor_ajuste':
                cell_text = format_decimal_dinamico(val, 3)
                align = WD_ALIGN_PARAGRAPH.CENTER
                indent = Inches(0)
            elif key in ['precio_unitario', 'precio_total']:
                cell_text = f"S/ {float(val):,.3f}" if pd.notna(val) else "S/ 0.000"
                align = WD_ALIGN_PARAGRAPH.RIGHT
            elif key == 'monto_soles':
                cell_text = f"S/ {format_decimal_dinamico(val, 3)}"
                align = WD_ALIGN_PARAGRAPH.RIGHT
            elif key == 'monto_dolares':
                cell_text = f"US$ {format_decimal_dinamico(val, 3)}"
                align = WD_ALIGN_PARAGRAPH.RIGHT
                indent = Inches(0)

            p = row_cells[i].paragraphs[0]
            run = p.add_run(cell_text); run.font.name = 'Arial'; run.font.size = Pt(8)
            p.alignment = align
            p.paragraph_format.left_indent = indent
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for edge in ['top', 'bottom']: set_cell_border(row_cells[i], **{edge: {"val": "nil"}})

    # Totales
    row_cells = table.add_row().cells
    p_total = row_cells[0].paragraphs[0]; run = p_total.add_run('Total'); run.bold = True; run.font.name = 'Arial'; run.font.size = Pt(8); p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Monto Soles (Columna 6 -> índice 6)
    p_sol = row_cells[6].paragraphs[0]; run = p_sol.add_run(f"S/ {total_soles:,.3f}"); run.bold = True; run.font.name = 'Arial'; run.font.size = Pt(8); p_sol.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Monto Dólares (Columna 7 -> índice 7)
    p_dol = row_cells[7].paragraphs[0]; run = p_dol.add_run(f"US$ {total_dolares:,.3f}"); run.bold = True; run.font.name = 'Arial'; run.font.size = Pt(8); p_dol.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    for cell in row_cells: set_cell_border(cell, top={"sz": 4, "val": "single"}, bottom={"sz": 4, "val": "single"})

    return sub


def create_detailed_ce_table_subdoc(doc_template, data, total_soles, total_dolares, footnotes_data=None):
    """
    Crea una tabla de CE detallada con anchos de columna personalizados
    para maximizar el espacio de la Descripción.
    """
    sub = doc_template.new_subdoc()
    
    headers = ["Descripción", "Unidad", "Cantidad", "Precio unitario (S/)", "Factor de ajuste 5/", "Monto (S/)", "Monto (US$) 6/"]
    num_cols = len(headers)
    table = sub.add_table(rows=1, cols=num_cols)
    table.style = 'TablaAnexo'
    
    # --- CAMBIO: Control manual de anchos (Total aprox 7 pulgadas) ---
    table.autofit = False 
    table.allow_autofit = False
    
    # Definir anchos (ajusta según necesites)
    widths = [
        Inches(1.5),  # Descripción (Más ancha)
        Inches(0.7),  # Unidad (Más chica)
        Inches(0.7),  # Cantidad (Más chica)
        Inches(0.7),  # Precio
        Inches(0.5),  # Factor
        Inches(0.9),  # Monto S/
        Inches(0.9)   # Monto US$
    ]
    
    # Aplicar anchos a cada celda de la primera fila (necesario para que Word lo respete)
    for i, width in enumerate(widths):
        table.columns[i].width = width
        table.rows[0].cells[i].width = width

    # --- Fin cambio de anchos ---

    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header_text); run.font.name = 'Arial'; run.font.size = Pt(8); run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(hdr_cells[i], top={"sz": 4, "val": "single"}, bottom={"sz": 4, "val": "single"})

    current_grupo = None
    current_subgrupo = None
    grupo_counter = 0  # <--- AGREGAR CONTADOR

    for item in data:
        grupo = item.get('grupo')
        subgrupo = item.get('subgrupo')

        # 1. Encabezado de GRUPO
        if grupo and grupo != current_grupo:
            current_grupo = grupo
            current_subgrupo = None
            grupo_counter += 1  # <--- INCREMENTAR
            row_cells = table.add_row().cells
            merged_cell = row_cells[0].merge(row_cells[num_cols - 1])
            p = merged_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # USAR EL CONTADOR AQUÍ:
            run = p.add_run(f"{grupo} {grupo_counter}/") 
            run.font.name = 'Arial'; run.font.size = Pt(8); run.bold = True
            p.paragraph_format.left_indent = Inches(0)
            run = p.add_run(f"{grupo} 1/")
            run.font.name = 'Arial'; run.font.size = Pt(8); run.bold = True
            for cell in row_cells: set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"})

        # 2. Encabezado de SUBGRUPO
        if subgrupo and subgrupo != current_subgrupo:
            current_subgrupo = subgrupo
            row_cells = table.add_row().cells
            merged_cell = row_cells[0].merge(row_cells[num_cols - 1])
            p = merged_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Inches(0.1)
            run = p.add_run(f"{subgrupo}")
            run.font.name = 'Arial'; run.font.size = Pt(8); run.italic = True
            for cell in row_cells: set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"})
        
        # 3. Fila de ITEM
        row_cells = table.add_row().cells
        cell_map = {
            0: {'key': 'descripcion', 'align': WD_ALIGN_PARAGRAPH.LEFT, 'indent': Inches(0.2)},
            1: {'key': 'unidad', 'align': WD_ALIGN_PARAGRAPH.CENTER}, 
            2: {'key': 'cantidad', 'align': WD_ALIGN_PARAGRAPH.CENTER, 'format': "{:,.0f}"},
            3: {'key': 'precio_unitario', 'align': WD_ALIGN_PARAGRAPH.RIGHT, 'format': "S/ {:,.3f}"},
            4: {'key': 'factor_ajuste', 'align': WD_ALIGN_PARAGRAPH.CENTER, 'format': "{:,.3f}"},
            5: {'key': 'monto_soles', 'align': WD_ALIGN_PARAGRAPH.RIGHT, 'format': "S/ {:,.3f}"},
            6: {'key': 'monto_dolares', 'align': WD_ALIGN_PARAGRAPH.RIGHT, 'format': "US$ {:,.3f}"},
        }

        for i, info in cell_map.items():
            row_cells[i].width = widths[i] # Aplicar ancho a cada celda
            
            val = item.get(info['key'])
            
            # --- CAMBIO 2: Seguridad para la columna Unidad ---
            if info['key'] == 'unidad':
                # Si es la columna unidad, forzamos el texto tal cual viene (ej. "8 horas")
                # Esto evita que formatos numéricos intenten procesar el texto.
                cell_text = str(val) if val is not None else ''
            else:
                # Para el resto de columnas, usamos el formato definido
                cell_text = info.get('format', '{}').format(val) if pd.notna(val) and val != '' else str(val or '')
            # --- FIN CAMBIO 2 ---
            
            # Ajustes adicionales existentes (se mantienen igual)
            if info['key'] == 'cantidad':
                 cant_val = item.get('cantidad')
                 cell_text = format_decimal_dinamico(cant_val, 0) if cant_val == int(cant_val) else format_decimal_dinamico(cant_val, 3)
            
            if info['key'] == 'precio_unitario' and val == 0:
                cell_text = "S/ 0.000"

            p = row_cells[i].paragraphs[0]
            
            # --- INICIO: Soporte para Saltos de Línea ---
            if '\n' in cell_text:
                # Si el texto tiene saltos de línea, lo dividimos y agregamos breaks
                parts = cell_text.split('\n')
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)
                    
                    # Si no es la última línea, añadir un salto (Shift+Enter en Word)
                    if idx < len(parts) - 1:
                        run.add_break()
            else:
                # Texto normal sin saltos
                run = p.add_run(cell_text)
                run.font.name = 'Arial'
                run.font.size = Pt(8)
            # --- FIN: Soporte para Saltos de Línea ---
            p.alignment = info['align']
            if 'indent' in info: p.paragraph_format.left_indent = info['indent']
            
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for edge in ['top', 'bottom']: set_cell_border(row_cells[i], **{edge: {"val": "nil"}})

    # Fila de Totales
    row_cells = table.add_row().cells
    # Aplicar anchos a totales también
    for i, width in enumerate(widths): row_cells[i].width = width
    
    p_total_desc = row_cells[0].paragraphs[0]; run = p_total_desc.add_run('Total'); run.bold = True; run.font.name = 'Arial'; run.font.size = Pt(8); p_total_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_total_soles = row_cells[5].paragraphs[0]; run = p_total_soles.add_run(f"S/ {total_soles:,.3f}"); run.bold = True; run.font.name = 'Arial'; run.font.size = Pt(8); p_total_soles.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_total_dolares = row_cells[6].paragraphs[0]; run = p_total_dolares.add_run(f"US$ {total_dolares:,.3f}"); run.bold = True; run.font.name = 'Arial'; run.font.size = Pt(8); p_total_dolares.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    for cell in row_cells:
        set_cell_border(cell, top={"sz": 4, "val": "single"}, bottom={"sz": 4, "val": "single"})

    if table.rows and len(table.rows) > 1:
        last_data_row = table.rows[-2] 
        for cell in last_data_row.cells:
             set_cell_border(cell, bottom={"sz": 2, "val": "single"}) 
    
    if footnotes_data:
        sub.add_paragraph() 
        p_fuente = sub.add_paragraph(style=footnotes_data.get('style', 'FuenteTabla'))
        p_fuente.add_run("Fuente: ").bold = True
        for i, nota in enumerate(footnotes_data.get('list', [])):
            if i > 0: p_fuente.add_run("\n")
            p_fuente.add_run(nota)

    return sub


# En funciones.py

# In funciones.py

# --- Make sure these imports are at the top of your funciones.py file ---
from textos_manager import obtener_fuente_formateada # Needed for the helper function
# Add any other necessary imports like create_main_table_subdoc, etc.

# En funciones.py

# In funciones.py

def create_consolidated_bi_table_subdoc(doc_template, resultados_extremos, total_bi_uit, footnotes_data=None, map_texto_a_letra=None, map_clave_a_texto=None):
    """
    Creates the consolidated BI table, assigning superscripts correctly using the provided maps.
    (Version ensuring all helper calls include map arguments).
    """
    filas_para_tabla_final = []
    if not resultados_extremos:
        return create_main_table_subdoc(doc_template, ["Descripción", "Monto"], [], ['descripcion_texto', 'monto'])

    map_texto_a_letra = map_texto_a_letra if map_texto_a_letra else {}
    map_clave_a_texto = map_clave_a_texto if map_clave_a_texto else {}
    id_infraccion = None
    if resultados_extremos:
        if 'id_infraccion' in resultados_extremos[0]:
             id_infraccion = resultados_extremos[0].get('id_infraccion')
        elif 'footnote_data' in resultados_extremos[0] and 'id_infraccion' in resultados_extremos[0]['footnote_data']:
             id_infraccion = resultados_extremos[0]['footnote_data'].get('id_infraccion')

    # --- Helper function definition (remains the same as the previous correct version) ---
    def get_final_letter_for_row(original_row, extremo_index, local_map_clave_a_texto, local_map_texto_a_letra):
        # ... (The full corrected code for the helper function goes here) ...
        original_ref_letter = original_row.get('ref')
        original_mapping = resultados_extremos[extremo_index].get('footnote_mapping', {})

        clave_semantica = None
        if not original_ref_letter: return ""

        # 1. Find the semantic key ('ce_anexo', 'cok')
        if original_ref_letter in original_mapping:
            clave_semantica = original_mapping[original_ref_letter]
        elif original_ref_letter in original_mapping.keys(): # Fallback if ref was the key
             clave_semantica = original_ref_letter
        
        if not clave_semantica: return ""

        # 2. Find the specific text
        map_key_specific = (clave_semantica, extremo_index)
        texto_nota_encontrada = local_map_clave_a_texto.get(map_key_specific)

        # 3. Fallback for common notes
        if texto_nota_encontrada is None:
             for (k, i), txt in local_map_clave_a_texto.items():
                  if k == clave_semantica:
                      texto_nota_encontrada = txt
                      break
                      
        if not texto_nota_encontrada: return ""

        # 4. Find the final letter
        letra_final = local_map_texto_a_letra.get(texto_nota_encontrada, "")
        
        superindice = f"({letra_final})" if letra_final else ""
        return superindice
    # --- End Helper Function ---


    # 1. Extract common rows (no changes here)
    primer_resultado = resultados_extremos[0]
    cos_anual_row = next((row for row in primer_resultado.get('table_rows', []) if 'COS (anual)' in row.get('descripcion', '')), None)
    cosm_row = next((row for row in primer_resultado.get('table_rows', []) if 'COSm (mensual)' in row.get('descripcion', '')), None)
    tc_row = next((row for row in primer_resultado.get('table_rows', []) if 'Tipo de cambio' in row.get('descripcion', '')), None)
    uit_row = next((row for row in primer_resultado.get('table_rows', []) if 'Unidad Impositiva' in row.get('descripcion', '')), None)

    # 2. Assemble the final table row by row, ensuring maps are passed to helper
    # Add CE rows for each extreme
    for i, res in enumerate(resultados_extremos):
        if res.get('table_rows'):
            ce_row_orig = res['table_rows'][0]
            letra_super = get_final_letter_for_row(ce_row_orig, i, map_clave_a_texto, map_texto_a_letra) # <-- Pass maps
            filas_para_tabla_final.append({
                'descripcion_texto': f"{ce_row_orig.get('descripcion', '').split(' [Extremo')[0]} [Extremo {i+1}]",
                'descripcion_superindice': letra_super,
                'monto': ce_row_orig.get('monto', '')
            })

    # Add common rows (COS, COSm)
    if cos_anual_row:
        letra_super = get_final_letter_for_row(cos_anual_row, 0, map_clave_a_texto, map_texto_a_letra) # <-- Pass maps
        filas_para_tabla_final.append({'descripcion_texto': cos_anual_row.get('descripcion',''), 'descripcion_superindice': letra_super, 'monto': cos_anual_row.get('monto','')})
    if cosm_row:
         letra_super = get_final_letter_for_row(cosm_row, 0, map_clave_a_texto, map_texto_a_letra) # <-- Pass maps
         filas_para_tabla_final.append({'descripcion_texto': cosm_row.get('descripcion',''), 'descripcion_superindice': letra_super, 'monto': cosm_row.get('monto','')})

    # Add T rows for each extreme
    for i, res in enumerate(resultados_extremos):
        t_row_orig = next((row for row in res.get('table_rows', []) if 'T: meses' in row.get('descripcion', '')), None)
        if t_row_orig:
            letra_super = get_final_letter_for_row(t_row_orig, i, map_clave_a_texto, map_texto_a_letra) # <-- Pass maps
            filas_para_tabla_final.append({
                'descripcion_texto': f"{t_row_orig.get('descripcion', '').split(' [Extremo')[0]} [Extremo {i+1}]",
                'descripcion_superindice': letra_super,
                'monto': t_row_orig.get('monto', '')
            })

    # Add adjusted cost rows and calculate total
    costos_ajustados_soles = []
    for i, res in enumerate(resultados_extremos):
         ajustado_row_orig = next((row for row in res.get('table_rows', []) if 'Costo evitado ajustado' in row.get('descripcion', '')), None)
         if ajustado_row_orig:
            letra_super = get_final_letter_for_row(ajustado_row_orig, i, map_clave_a_texto, map_texto_a_letra) # <-- Pass maps
            monto_str = ajustado_row_orig.get('monto', '')
            filas_para_tabla_final.append({
                 'descripcion_texto': f"{ajustado_row_orig.get('descripcion', '').split(' [Extremo')[0]} [Extremo {i+1}]",
                 'descripcion_superindice': letra_super,
                 'monto': monto_str
            })
            try:
                numeric_part = monto_str.replace('S/','').replace('US$','').replace(',','').strip()
                costos_ajustados_soles.append(float(numeric_part))
            except (ValueError, AttributeError):
                 costos_ajustados_soles.append(0.0)

    total_ajustado_soles = sum(costos_ajustados_soles)
    filas_para_tabla_final.append({'descripcion_texto': 'Costo evitado ajustado total (S/)', 'descripcion_superindice': '', 'monto': f"S/ {total_ajustado_soles:,.3f}"})

    if tc_row:
        letra_super = get_final_letter_for_row(tc_row, 0, map_clave_a_texto, map_texto_a_letra) # <-- Pass maps
        filas_para_tabla_final.append({'descripcion_texto': tc_row.get('descripcion',''), 'descripcion_superindice': letra_super, 'monto': tc_row.get('monto','')})

    filas_para_tabla_final.append({'descripcion_texto': 'Beneficio ilícito (S/)', 'descripcion_superindice': letra_super, 'monto': f"S/ {total_ajustado_soles:,.3f}"})

    if uit_row:
        letra_super = get_final_letter_for_row(uit_row, 0, map_clave_a_texto, map_texto_a_letra) # <-- Pass maps
        filas_para_tabla_final.append({'descripcion_texto': uit_row.get('descripcion',''), 'descripcion_superindice': letra_super, 'monto': uit_row.get('monto','')})

    filas_para_tabla_final.append({'descripcion_texto': 'Beneficio Ilícito (UIT)', 'descripcion_superindice': '', 'monto': f"{total_bi_uit:,.3f} UIT"})

    # 3. Call create_main_table_subdoc
    return create_main_table_subdoc(
        doc_template,
        headers=["Descripción", "Monto"],
        data=filas_para_tabla_final,
        keys=['descripcion_texto', 'monto'],
        footnotes_data=footnotes_data,
        column_widths=(5, 1.5)
    )

def _replace_in_paragraph(p, pattern, numbering_manager):
    full_text = "".join(run.text for run in p.runs)
    if '###TABLE_TITLE(' in full_text:
        match = pattern.search(full_text)
        if match:
            title_text = match.group(1) # Extrae el texto del título
            new_title = numbering_manager.get_table_title(title_text)
            
            # Reemplaza el placeholder con el título numerado
            final_text = pattern.sub(new_title, full_text, 1)
            
            p.clear()
            run = p.add_run(final_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def post_process_numbering(doc, numbering_manager):
    """
    Busca placeholders (ej. ###TABLE_TITLE(Título)###) y los reemplaza.
    """
    # Expresión regular para encontrar ###TABLE_TITLE(...)###
    pattern = re.compile(r"###TABLE_TITLE\((.*?)\)###")

    for p in doc.paragraphs:
        _replace_in_paragraph(p, pattern, numbering_manager)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, pattern, numbering_manager)


class NumberingManager:
    """
    Gestiona la numeración secuencial de elementos como cuadros o tablas.
    """
    def __init__(self):
        self.table_count = 0

    def get_table_title(self, title_text):
        """
        Incrementa el contador y devuelve el título formateado para un cuadro.
        """
        self.table_count += 1
        return f"Cuadro n.° {self.table_count}: {title_text}"


class AcronymManager:
    """
    Gestiona el uso de acrónimos para definirlos una sola vez en plantillas de Word.
    """
    def __init__(self):
        self.defined_acronyms = set()

    def get(self, key, full_text, acronym):
        """
        Devuelve el texto completo la primera vez, y el acrónimo las siguientes.
        El resultado es un objeto RichText para formato.
        """
        rt = RichText()
        if key not in self.defined_acronyms:
            self.defined_acronyms.add(key)
            rt.add(f"{full_text} (en adelante, ")
            rt.add(acronym, bold=True, italic=True)
            rt.add(")")
        else:
            rt.add(acronym, bold=True, italic=True)
        return rt


def combinar_con_composer(path_plantilla_o_buffer, path_origen_o_buffer, path_final):
    """Combina dos documentos de Word usando docx-composer."""
    print(f"-> Combinando documentos con docx-composer...")
    try:
        plantilla = Document(path_plantilla_o_buffer)
        composer = Composer(plantilla)
        origen = Document(path_origen_o_buffer)
    except Exception as e:
        st.error(f"   ERROR abriendo documents para combinar: {e}")
        return False

    marcador = None
    marcador_idx = -1
    for i, p in enumerate(composer.doc.paragraphs):
        if '{{INSERTAR_CONTENIDO_AQUI}}' in p.text:
            marcador = p
            marcador_idx = i
            break
            
    if not marcador:
        st.error(f"   ERROR: No se encontró el marcador '{{INSERTAR_CONTENIDO_AQUI}}'.")
        return False
        
    composer.insert(marcador_idx, origen)
    for p in composer.doc.paragraphs:
        if '{{INSERTAR_CONTENIDO_AQUI}}' in p.text:
            p._element.getparent().remove(p._element)
            break
            
    composer.save(path_final)
    
    # Comprueba si 'path_final' es una ruta de archivo (string) o un buffer
    if isinstance(path_final, str):
        nombre_amigable = os.path.basename(path_final)
        print(f"-> Fusión guardada en '{nombre_amigable}'.")
    else:
        print(f"-> Fusión guardada en un buffer en memoria.")

    return True

# --- Funciones Auxiliares de Formato ---

def set_cell_border(cell, **kwargs):
    """
    Función auxiliar para definir los bordes de una celda.
    Uso: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            border = tcBorders.find(qn(tag))
            if border is None:
                border = OxmlElement(tag)
                tcBorders.append(border)
            for k, v in edge_data.items():
                border.set(qn('w:{}'.format(k)), str(v))


def set_cell_shading(cell, fill_color):
    """
    Función auxiliar para definir el color de fondo de una celda.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

# --- Funciones Constructoras de Contenido ---

# En funciones.py

def create_table_subdoc(doc_template, headers, data, keys, footnotes_data=None):
    """
    Crea una tabla con formato avanzado y ahora puede incluir un título y notas al pie.
    """
    sub = doc_template.new_subdoc()

    table = sub.add_table(rows=1, cols=len(headers))
    table.style = 'TablaAnexo'

    # --- Formato del Encabezado ---
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header_text)
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(hdr_cells[i], top={"sz": 4, "val": "single", "color": "000000"},
                        bottom={"sz": 4, "val": "single", "color": "000000"})

    # --- Formato de las Filas de Datos ---
    for item in data:
        row_cells = table.add_row().cells
        for i, key in enumerate(keys):
            cell_text = str(item.get(key, ''))
            p = row_cells[i].paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = 'Arial'
            run.font.size = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(row_cells[i], top={"val": "nil"}, bottom={"val": "nil"})

    # --- Formato de la Última Fila (Total) ---
    if data:
        last_row_cells = table.rows[-1].cells
        for cell in last_row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_border(cell, top={"sz": 4, "val": "single", "color": "000000"},
                            bottom={"sz": 4, "val": "single", "color": "000000"})

    # --- INICIO DE LA MODIFICACIÓN: AÑADIR NOTAS AL PIE ---
    if footnotes_data:
        footnotes_list = footnotes_data.get('list', [])
        style_name = footnotes_data.get('style', 'FuenteTabla')
        
        # Añadimos un párrafo en blanco para dar espacio
        sub.add_paragraph()
        
        p_fuente = sub.add_paragraph()
        run_fuente = p_fuente.add_run("Fuente:")
        run_fuente.bold = False
        p_fuente.style = style_name
        
        for nota in footnotes_list:
            sub.add_paragraph(nota, style=style_name)
    # --- FIN DE LA MODIFICACIÓN ---

    return sub

# En funciones.py

def create_main_table_subdoc(doc_template, headers, data, keys, texto_posterior=None, estilo_texto_posterior=None, footnotes_data=None, column_widths=None):
    """
    Crea una tabla principal con formato, manejando correctamente la primera columna
    para tablas con y sin superíndices (como BI y Multa).
    """
    sub = doc_template.new_subdoc()

    # Crear tabla y aplicar estilo base
    table = sub.add_table(rows=1, cols=len(headers))
    table.style = 'TablaCuerpo' # Asegúrate que este estilo exista en tu plantilla base
    SHADING_COLOR = "D9D9D9" # Color de sombreado para encabezado y última fila

    # Aplicar anchos de columna si se proporcionan
    if column_widths:
        if len(column_widths) == len(headers):
            for i, width in enumerate(column_widths):
                # Comprobar si la columna existe antes de intentar modificarla
                if i < len(table.columns):
                    table.columns[i].width = Inches(width)

    # --- Formatear Encabezado ---
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        # Comprobar si la celda existe
        if i < len(hdr_cells):
            cell = hdr_cells[i]
            p = cell.paragraphs[0]
            p.clear() # Limpiar contenido por defecto
            run = p.add_run(header_text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, SHADING_COLOR)

    # --- INICIO DE LA CORRECCIÓN LÓGICA ---
    # Formatear Filas de Datos
    for row_idx, item in enumerate(data):
        row_cells = table.add_row().cells
        is_last_row = (row_idx == len(data) - 1)

        for col_idx, key in enumerate(keys):
            if col_idx < len(row_cells):
                cell = row_cells[col_idx]
                p = cell.paragraphs[0]
                p.clear()

                # CASO 1: Es la columna 0 Y los datos tienen 'descripcion_texto' (TABLA BI)
                if col_idx == 0 and (item.get('descripcion_texto') is not None):
                    texto_principal = str(item.get('descripcion_texto', ''))
                    superindice = str(item.get('descripcion_superindice', ''))

                    run_texto = p.add_run(texto_principal)
                    run_texto.font.name = 'Arial'; run_texto.font.size = Pt(10)
                    if is_last_row: run_texto.bold = True

                    if superindice:
                        run_super = p.add_run(superindice)
                        run_super.font.superscript = True
                        run_super.font.name = 'Arial'; run_super.font.size = Pt(10)
                        if is_last_row: run_super.bold = True
                    
                    # Lógica de alineación movida DENTRO del if
                    if texto_principal.strip().startswith("CE:"):
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # CASO 2: Es la columna 0 pero NO tiene 'descripcion_texto' (TABLA MULTA)
                # O es cualquier otra columna (Monto)
                else:
                    cell_text = str(item.get(key, ''))
                    run = p.add_run(cell_text)
                    run.font.name = 'Arial'; run.font.size = Pt(10)
                    if is_last_row: run.bold = True
                    
                    # Lógica de alineación para el CASO 2
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT

                # --- FIN DE LA CORRECCIÓN LÓGICA ---
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if is_last_row:
                    set_cell_shading(cell, SHADING_COLOR)

    # --- Añadir texto opcional después de la tabla ---
    if texto_posterior:
        # Usar estilo proporcionado o 'Normal' como fallback
        style_to_use = estilo_texto_posterior if estilo_texto_posterior else 'Normal'
        try:
             # Verificar si el estilo existe en la plantilla base
             if style_to_use in doc_template.styles:
                 sub.add_paragraph(texto_posterior, style=style_to_use)
             else:
                 print(f"Advertencia: Estilo '{style_to_use}' no encontrado para texto_posterior. Usando estilo por defecto.")
                 p = sub.add_paragraph(texto_posterior) # Usar estilo por defecto si no se encuentra
        except Exception as e: # Captura más genérica por si acaso
             print(f"Error al aplicar estilo '{style_to_use}' para texto_posterior: {e}. Usando estilo por defecto.")
             p = sub.add_paragraph(texto_posterior)

    # --- Añadir notas al pie si se proporcionan ---
    if footnotes_data:
        footnotes_list = footnotes_data.get('list', [])
        elaboration_text = footnotes_data.get('elaboration', '')
        # Usar estilo 'FuenteTabla' por defecto o 'Normal' si 'FuenteTabla' no existe
        style_name = footnotes_data.get('style', 'FuenteTabla')
        footnote_style_exists = True
        try:
            if style_name not in doc_template.styles:
                print(f"Advertencia: Estilo de nota al pie '{style_name}' no encontrado. Usando estilo por defecto.")
                style_name = 'Normal' # Usar estilo por defecto como fallback
                footnote_style_exists = False # Marcar que el estilo original no se encontró
        except Exception as e:
             print(f"Error al verificar estilo '{style_name}': {e}. Usando estilo por defecto.")
             style_name = 'Normal'
             footnote_style_exists = False

        # Añadir prefijo "Fuente:" (solo si hay notas o texto de elaboración)
        if footnotes_list or elaboration_text:
            p_fuente = sub.add_paragraph()
            run_fuente = p_fuente.add_run("Fuente:")
    
            p_fuente.style = style_name # Aplicar estilo (original o fallback)

            # Añadir cada nota al pie
            for nota in footnotes_list:
                sub.add_paragraph(str(nota), style=style_name)

            # Añadir texto de elaboración si existe
            if elaboration_text:
                sub.add_paragraph(str(elaboration_text), style=style_name)

    return sub


# En funciones.py

def create_summary_table_subdoc(doc_template, headers, data, keys, texto_posterior=None, column_widths=None):
    """
    Crea la tabla de resumen final de la multa, con celdas combinadas y ahora con un título.
    """
    sub = doc_template.new_subdoc()

    table = sub.add_table(rows=1, cols=len(headers))
    table.style = 'TablaCuerpo'
    SHADING_COLOR = "D9D9D9"

    if column_widths:
        if len(column_widths) == len(headers):
            for i, width in enumerate(column_widths):
                table.columns[i].width = Inches(width)

    # --- Formato del Encabezado ---
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header_text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr_cells[i], SHADING_COLOR)

    # --- Fila de Datos Principal ---
    for item in data[:-1]:
        row_cells = table.add_row().cells
        for col_idx, key in enumerate(keys):
            cell_text = str(item.get(key, ''))
            p = row_cells[col_idx].paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- Fila de Total con Celdas Combinadas ---
    total_multa_str = data[-1]['Monto'] if data and 'Monto' in data[-1] else "0.000 UIT"
    total_cells = table.add_row().cells
    
    merged_cell = total_cells[0].merge(total_cells[1])
    p_total = merged_cell.paragraphs[0]
    run_total = p_total.add_run('Total')
    run_total.font.name = 'Arial'
    run_total.font.size = Pt(10)
    run_total.bold = True
    p_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(merged_cell, SHADING_COLOR)

    p_multa = total_cells[2].paragraphs[0]
    run_multa = p_multa.add_run(total_multa_str)
    run_multa.font.name = 'Arial'
    run_multa.font.size = Pt(10)
    run_multa.bold = True
    p_multa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(total_cells[2], SHADING_COLOR)
    set_cell_shading(total_cells[1], SHADING_COLOR)

    if texto_posterior:
        p = sub.add_paragraph()
        run = p.add_run(texto_posterior)
        p.style = 'FuenteTabla'

    return sub

def create_personal_table_subdoc(doc_template, headers, data, keys, texto_posterior=None, column_widths=None):
    """
    Crea la tabla de personal a capacitar, con formato similar al resumen y celdas combinadas.
    """
    sub = doc_template.new_subdoc()

    table = sub.add_table(rows=1, cols=len(headers))
    table.style = 'TablaCuerpo' # Usamos el mismo estilo para que se vea igual
    SHADING_COLOR = "D9D9D9"

    if column_widths:
        if len(column_widths) == len(headers):
            for i, width in enumerate(column_widths):
                table.columns[i].width = Inches(width)

    # --- Formato del Encabezado (sin cambios) ---
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header_text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr_cells[i], SHADING_COLOR)

    # --- Fila de Datos Principal (sin cambios) ---
    for item in data[:-1]: # Itera hasta la penúltima fila
        row_cells = table.add_row().cells
        for col_idx, key in enumerate(keys):
            cell_text = str(item.get(key, ''))
            p = row_cells[col_idx].paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            # --- INICIO DE LA CORRECCIÓN 1 ---
            # Col 0 (Perfil) y Col 2 (Cantidad) van centrados
            if col_idx == 0 or col_idx == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Col 1 (Descripción) va justificada
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # --- FIN DE LA CORRECCIÓN 1 ---
            row_cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- Fila de Total con Celdas Combinadas (AQUÍ ESTÁ EL CAMBIO) ---
    # 1. Ahora busca la clave 'Cantidad' en lugar de 'Monto'
    total_cantidad_str = str(data[-1]['Cantidad']) if data and 'Cantidad' in data[-1] else "0"
    total_cells = table.add_row().cells
    
    merged_cell = total_cells[0].merge(total_cells[1])
    p_total = merged_cell.paragraphs[0]
    run_total = p_total.add_run('Total')
    run_total.font.name = 'Arial'
    run_total.font.size = Pt(10)
    run_total.bold = True
    p_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(merged_cell, SHADING_COLOR)

    # 2. Inserta el valor de la cantidad total en la última celda
    p_cantidad = total_cells[2].paragraphs[0]
    run_cantidad = p_cantidad.add_run(total_cantidad_str)
    run_cantidad.font.name = 'Arial'
    run_cantidad.font.size = Pt(10)
    run_cantidad.bold = True
    p_cantidad.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(total_cells[2], SHADING_COLOR)
    set_cell_shading(total_cells[1], SHADING_COLOR) # Celda invisible pero necesaria

    # --- INICIO DE LA CORRECCIÓN 2 (Notas al pie) ---
    # Añadir la nota al pie (1)
    p_nota = sub.add_paragraph()
    run_nota = p_nota.add_run("(1) O a quien se designe del área en referencia.")
    p_nota.style = 'FuenteTabla' # Asume que el estilo 'FuenteTabla' es el correcto
    
    # Añadir el texto de Elaboración (que ya venía en texto_posterior)
    if texto_posterior:
        p = sub.add_paragraph()
        run = p.add_run(texto_posterior)
        p.style = 'FuenteTabla'
    # --- FIN DE LA CORRECCIÓN 2 ---

    return sub

def create_footnotes_subdoc(doc_template, footnotes_list, style_name='FuenteTabla', title_font_size=8):
    """
    Crea un subdocumento para el bloque de fuentes aplicando un estilo de párrafo específico.
    """
    sub = doc_template.new_subdoc()
    
    # --- INICIO DE LA MODIFICACIÓN ---
    # Añadimos cada fuente usando el ESTILO que nos pasen como parámetro
    for text in footnotes_list:
        sub.add_paragraph(text, style=style_name)
    # --- FIN DE LA MODIFICACIÓN ---

    return sub

def texto_con_numero(numero, genero='m'):
    """
    Formatea un número. Si es entero, lo devuelve como 'texto (número)'.
    Si tiene decimales, devuelve solo el número formateado hasta 3 decimales.
    """
    if numero is None:
        return "(N/A)"

    # --- INICIO DE LA MODIFICACIÓN ---
    # Comprobamos si el número puede ser tratado como un entero (ej: 15.0 o 15)
    if numero == int(numero):
        # --- LÓGICA PARA NÚMEROS ENTEROS ---
        num_entero = int(numero)
        
        # Convertir a texto
        texto_num = num2words(num_entero, lang='es')
        
        # Casos especiales para el número 1
        if num_entero == 1:
            if genero == 'm':
                texto_num = "un"
            elif genero == 'f':
                texto_num = "una"
        
        return f"{texto_num} ({num_entero})"
        
    else:
        # --- LÓGICA PARA NÚMEROS CON DECIMALES ---
        # Redondea a 3 decimales y elimina los ceros sobrantes al final
        num_str = f"{numero:.3f}".rstrip('0').rstrip('.')
        return num_str
    # --- FIN DE LA MODIFICACIÓN ---