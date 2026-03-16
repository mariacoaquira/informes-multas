from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Mm, RGBColor, Cm
from docx.opc.part import XmlPart
from docx.opc.packuri import PackURI
from docx.enum.table import WD_TABLE_ALIGNMENT
from funciones import formatear_lista_hechos

def configurar_pagina_a4(doc):
    """
    Fuerza el tamaño de página a A4 y establece los márgenes de tu formato oficial:
    Superior/Inferior: 2.5 cm
    Izquierdo/Derecho: 3.0 cm
    """
    seccion = doc.sections[0]
    
    # Tamaño A4 exacto (210 mm x 297 mm)
    seccion.page_width = Mm(210)
    seccion.page_height = Mm(297)
    
    # Márgenes exactos según tu captura
    seccion.top_margin = Cm(2.5)
    seccion.bottom_margin = Cm(2.5)
    seccion.left_margin = Cm(3.0)
    seccion.right_margin = Cm(3.0)

def configurar_estilo_base(doc):
    """
    Cambia la fuente por defecto a Arial 11, espaciado sencillo y TODO JUSTIFICADO.
    """
    # --- ESTILO NORMAL (Cuerpo del texto) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    p_format = style.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def crear_encabezado_ifi(doc, ruta_imagen, texto_debajo):
    """
    Crea el encabezado con una imagen ancha arriba y el texto institucional debajo,
    ambos perfectamente centrados, con un espacio de respiro al final.
    """
    header = doc.sections[0].header
    
    # 1. Párrafo para la imagen (PEGADO A LA DERECHA Y ANCHO EXACTO)
    p_img = header.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # <--- Alineado a la derecha
    run_img = p_img.add_run()
    
    try:
        # El ancho exacto del área de texto es 15.0 cm (21cm A4 - 3cm izq - 3cm der)
        run_img.add_picture(ruta_imagen, width=Cm(15.0)) 
    except FileNotFoundError:
        run_img.add_text("[IMAGEN NO ENCONTRADA: Asegúrate de tener un 'logo.png']")

    # 2. Párrafo para el texto (CENTRADO)
    p_txt = header.add_paragraph()
    p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    
    # --- NUEVO: Espacio de respiro debajo del encabezado ---
    p_txt.paragraph_format.space_after = Pt(10) 
    
    run_txt = p_txt.add_run(texto_debajo)
    run_txt.font.name = 'Arial'
    run_txt.font.size = Pt(8) 
    run_txt.bold = True

def crear_bloque_cabecera(doc, datos_cabecera):
    """
    Crea el bloque de A, DE, ASUNTO, etc. ocupando el 100% del ancho disponible.
    Si recibe una lista, crea párrafos independientes (evitando estiramientos de texto justificado)
    y aplica negrita según se indique.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT # Por si acaso
    
    tabla = doc.add_table(rows=len(datos_cabecera), cols=3)
    tabla.autofit = False
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    anchos = [Cm(3.5), Cm(0.5), Cm(11.0)] 
    
    for i, (etiqueta, valor) in enumerate(datos_cabecera):
        row = tabla.rows[i]
        
        row.cells[0].width = anchos[0]
        row.cells[1].width = anchos[1]
        row.cells[2].width = anchos[2]
        
        # 1. Celda de etiqueta
        p_etiq = row.cells[0].paragraphs[0]
        p_etiq.add_run(etiqueta).bold = True
        
        # 2. Celda de los ":"
        p_puntos = row.cells[1].paragraphs[0]
        p_puntos.add_run(":").bold = True
        p_puntos.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        
        # 3. Celda del valor (AQUÍ ESTÁ LA MAGIA)
        celda_valor = row.cells[2]
        es_ultima_fila = (i == len(datos_cabecera) - 1)
        espacio_abajo = Pt(0) if es_ultima_fila else Pt(14)

        # Ajustamos los espacios de la etiqueta y los dos puntos
        p_etiq.paragraph_format.space_before = Pt(0)
        p_etiq.paragraph_format.space_after = espacio_abajo
        p_puntos.paragraph_format.space_before = Pt(0)
        p_puntos.paragraph_format.space_after = espacio_abajo

        if isinstance(valor, list):
            # Si nos mandan una lista de tuplas (texto, es_negrita)
            for idx, (texto_linea, es_negrita) in enumerate(valor):
                if idx == 0:
                    p_val = celda_valor.paragraphs[0]
                else:
                    p_val = celda_valor.add_paragraph()
                
                p_val.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_val.paragraph_format.space_before = Pt(0)
                # Solo el último párrafo de la celda recibe el espacio grande de separación de filas
                p_val.paragraph_format.space_after = espacio_abajo if idx == len(valor) - 1 else Pt(0)
                
                if texto_linea: # Si no está vacío, lo escribimos
                    run = p_val.add_run(texto_linea)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.bold = es_negrita
        else:
            # Si es un texto simple normal (Como Asunto o Fecha)
            p_val = celda_valor.paragraphs[0]
            p_val.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_val.paragraph_format.space_before = Pt(0)
            p_val.paragraph_format.space_after = espacio_abajo
            run = p_val.add_run(valor)
            run.font.name = 'Arial'
            run.font.size = Pt(11)

def agregar_linea_horizontal(doc):
    """
    Inserta una línea horizontal sólida y delgada.
    """
    p = doc.add_paragraph()
    # Nos aseguramos de que este párrafo tampoco tenga espacios fantasma
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    
    # Cambiamos '12' por '6' para que sea una línea delgada y fina (0.75 pt)
    bottom.set(qn('w:sz'), '6') 
    
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000') 
    pBdr.append(bottom)
    pPr.append(pBdr)

def agregar_titulo_numerado(doc, numero, texto):
    """
    Crea un título numerado perfecto (ej: "1. Antecedentes") con sangría francesa 
    de 1 cm, tabulador exacto y un espacio automático por arriba.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-1.0)
    pf.tab_stops.clear_all()
    pf.tab_stops.add_tab_stop(Cm(1.0))
    
    # Le damos un respiro (espacio) por arriba automáticamente equivalente a un "Enter"
    pf.space_before = Pt(14) 
    
    run = p.add_run(f"{numero}.\t{texto}")
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True
    return p

def agregar_vineta_flecha(doc, texto, prefijo_negrita=""):
    """
    Agrega un párrafo simulando una viñeta de flecha (➤) alineada a 1 cm.
    Permite añadir un prefijo en negrita antes del texto regular.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    
    pf.left_indent = Cm(1.63)
    pf.first_line_indent = Cm(-0.63)
    pf.tab_stops.clear_all()
    pf.tab_stops.add_tab_stop(Cm(1.63))
    
    # Insertamos la flecha y el tabulador (\t)
    run_flecha = p.add_run("➤\t")
    run_flecha.font.name = 'Arial'
    run_flecha.font.size = Pt(11)
    
    # Si enviamos un prefijo (ej: "Hecho imputado n.° 1: "), lo pone en negrita
    if prefijo_negrita:
        run_prefijo = p.add_run(prefijo_negrita)
        run_prefijo.font.name = 'Arial'
        run_prefijo.font.size = Pt(11)
        run_prefijo.bold = True
        
    # Luego insertamos el texto normal del hecho
    run_texto = p.add_run(texto)
    run_texto.font.name = 'Arial'
    run_texto.font.size = Pt(11)
    
    pf.space_before = Pt(0)
    pf.space_after = Pt(6) 
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    return p

def agregar_subtitulo_numerado(doc, numero_texto, texto):
    """
    Crea un subtítulo (ej: "3.1.") alineado exactamente igual que los títulos principales,
    respetando el tabulador a 1.0 cm.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-1.0)
    pf.tab_stops.clear_all()
    pf.tab_stops.add_tab_stop(Cm(1.0))
    
    # Su espacio automático por arriba (equivalente a un Enter)
    pf.space_before = Pt(14) 
    
    # Concatenamos directamente sin añadir un punto extra
    run = p.add_run(f"{numero_texto}\t{texto}")
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True
    return p

def inicializar_notas_al_pie(doc):
    # 1. Verificar si ya existe la tubería de notas al pie
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            return

    # 2. El código XML estructural mínimo para las notas al pie
    footnotes_xml = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
    <w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:footnote w:type="separator" w:id="-1">
            <w:p><w:r><w:separator/></w:r></w:p>
        </w:footnote>
        <w:footnote w:type="continuationSeparator" w:id="0">
            <w:p><w:r><w:continuationSeparator/></w:r></w:p>
        </w:footnote>
    </w:footnotes>"""

    partname = PackURI('/word/footnotes.xml')
    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
    
    # 3. Cargamos el XML en memoria como una pieza oficial (XmlPart)
    footnotes_part = XmlPart.load(partname, content_type, footnotes_xml, doc.part.package)
    
    # 4. ¡LA CLAVE ESTÁ AQUÍ! 
    # Solo "relacionamos" el archivo. python-docx se encargará de guardarlo 
    # automáticamente en el ZIP final al momento de hacer doc.save()
    doc.part.relate_to(footnotes_part, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")

def agregar_nota_al_pie(parrafo, texto_nota, doc, espaciado_twips="567"):
    """
    Agrega una nota al pie dinámicamente convirtiendo los \n en PÁRRAFOS REALES
    para evitar que el texto justificado estire las líneas cortas.
    """
    inicializar_notas_al_pie(doc)
    
    footnotes_part = None
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            footnotes_part = rel.target_part
            break

    footnotes_element = footnotes_part.element
    
    current_ids = [
        int(f.get(qn('w:id'))) for f in footnotes_element.findall(qn('w:footnote')) 
        if f.get(qn('w:id')) is not None
    ]
    nuevo_id = str(max(current_ids) + 1) if current_ids else "1"

    footnote = OxmlElement('w:footnote')
    footnote.set(qn('w:id'), nuevo_id)

    def aplicar_arial_8(rPr_elemento):
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rFonts.set(qn('w:cs'), 'Arial')
        rPr_elemento.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '16')
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '16')
        rPr_elemento.append(sz)
        rPr_elemento.append(szCs)

# --- AQUÍ ESTÁ LA MAGIA: Dividimos por \n y creamos párrafos reales ---
    lineas = texto_nota.split('\n')
    
    for i, linea in enumerate(lineas):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'FootnoteText')
        pPr.append(pStyle)
        
        # Aseguramos que el párrafo sea Justificado
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        
        # --- NUEVO: Darle un "respiro" al final de la nota ---
        # Si es la última línea de esta nota, le agregamos un espacio por debajo
        if i == len(lineas) - 1:
            spacing = OxmlElement('w:spacing')
            # 160 twips equivalen a unos 8pt. Si quieres un Enter más grande, pon '240'
            spacing.set(qn('w:after'), '240') 
            pPr.append(spacing)
        
        # Sangría: El primer párrafo lleva Sangría Francesa (hanging) para el numerito.
        # Los siguientes solo llevan sangría izquierda para alinearse al bloque de texto.
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(espaciado_twips))
        if i == 0:
            ind.set(qn('w:hanging'), str(espaciado_twips))
        pPr.append(ind)
        p.append(pPr)

        if i == 0:
            # Insertar el superíndice numérico SOLO en el primer párrafo
            r_ref = OxmlElement('w:r')
            rPr_ref = OxmlElement('w:rPr')
            rStyle_ref = OxmlElement('w:rStyle')
            rStyle_ref.set(qn('w:val'), 'FootnoteReference')
            rPr_ref.append(rStyle_ref)
            aplicar_arial_8(rPr_ref)
            vertAlign = OxmlElement('w:vertAlign')
            vertAlign.set(qn('w:val'), 'superscript')
            rPr_ref.append(vertAlign)
            r_ref.append(rPr_ref)
            footnoteRef = OxmlElement('w:footnoteRef')
            r_ref.append(footnoteRef)
            p.append(r_ref)
            
            # Tabulador para separar el número del primer texto
            r_tab = OxmlElement('w:r')
            r_tab.append(OxmlElement('w:tab'))
            p.append(r_tab)

        # Insertar el texto de la línea
        r_text = OxmlElement('w:r')
        rPr_text = OxmlElement('w:rPr')
        aplicar_arial_8(rPr_text)
        r_text.append(rPr_text)
        
        partes_tab = linea.split('\t')
        for j, texto_parte in enumerate(partes_tab):
            if j > 0:
                r_text.append(OxmlElement('w:tab'))
            
            if texto_parte:
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = texto_parte
                r_text.append(t)
                
        p.append(r_text)
        footnote.append(p)

    footnotes_element.append(footnote)

    # --- 4. INYECTAR LA REFERENCIA EN EL PÁRRAFO VISIBLE ---
    run = parrafo.add_run()
    run.font.name = 'Arial' # Forzamos la misma fuente para evitar saltos
    r = run._r
    rPr = OxmlElement('w:rPr')
    
    # ¡MAGIA! Eliminamos la asignación del estilo "FootnoteReference".
    # Solo le aplicamos el formato de superíndice directamente.
    vertAlign2 = OxmlElement('w:vertAlign')
    vertAlign2.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign2)
    
    footnoteReference = OxmlElement('w:footnoteReference')
    footnoteReference.set(qn('w:id'), nuevo_id)
    r.append(rPr)
    r.append(footnoteReference)

def agregar_cuadro_formula(doc):
    """
    Crea un cuadro compacto, alineado correctamente, forzando texto a la 
    izquierda en su interior para evitar estiramientos de palabras.
    """
    # 1. TÍTULO DEL CUADRO
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(14)
    
    # --- LA LÍNEA MÁGICA ---
    # Le damos la misma sangría que al cuadro para que su "centro" se mueva a la derecha
    p_titulo.paragraph_format.left_indent = Cm(1.2) 
    
    run_titulo = p_titulo.add_run("Cuadro n.° 1: Fórmula para el cálculo de multa")
    run_titulo.font.name = 'Arial'
    run_titulo.font.size = Pt(11)
    run_titulo.bold = True

    # 2. CREAMOS LA CAJA (Tabla de 1x1)
    tabla = doc.add_table(rows=1, cols=1)
    tabla.style = 'Table Grid'
    
    # --- Devolvemos el cuadro a su posición perfecta (680 twips) ---
    tblPr = tabla._tblPr
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '680') 
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    celda = tabla.cell(0, 0)
    celda.width = Cm(14.0) 
    
    # -- A. La fórmula matemática centradita --
    p_eq = celda.paragraphs[0]
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq.paragraph_format.space_before = Pt(8) # Respiro superior reducido
    p_eq.paragraph_format.space_after = Pt(8)  
    
    run_eq = p_eq.add_run("Multa (M) = ( B / p ) · [F]")
    run_eq.font.name = 'Cambria Math'
    run_eq.font.size = Pt(10)
    
    # -- B. Variables "Donde:" --
    p_donde = celda.add_paragraph()
    p_donde.alignment = WD_ALIGN_PARAGRAPH.LEFT # Quitamos el justificado
    p_donde.paragraph_format.left_indent = Cm(0.2) # Menos margen interno
    p_donde.paragraph_format.right_indent = Cm(0.2)
    p_donde.paragraph_format.space_after = Pt(6) # El espacio debajo de "Donde:"
    run_donde = p_donde.add_run("Donde:")
    run_donde.font.name = 'Cambria Math'
    run_donde.font.size = Pt(10)

    def agregar_variable(var_letra, descripcion, es_ultima=False):
        p = celda.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT # Evita los huecos gigantes entre palabras
        p.paragraph_format.left_indent = Cm(0.2) 
        p.paragraph_format.right_indent = Cm(0.2) 
        
        p.paragraph_format.space_after = Pt(8) if es_ultima else Pt(6)
        
        run_var = p.add_run(var_letra)
        run_var.font.name = 'Cambria Math'
        run_var.font.size = Pt(10)
        
        run_desc = p.add_run(f" = {descripcion}")
        run_desc.font.name = 'Cambria Math'
        run_desc.font.size = Pt(10)

    agregar_variable("B", "Beneficio ilícito (obtenido por el administrado al incumplir la norma)")
    agregar_variable("p", "Probabilidad de detección")
    agregar_variable("F", "Suma de los factores para la graduación de sanciones (1+f1+f2+f3+f4+f5+f6+f7)", es_ultima=True)

    # 3. NOTA AL PIE DEL CUADRO (Elaboración)
    p_fuente = doc.add_paragraph()
    p_fuente.paragraph_format.space_before = Pt(0) 
    
    # Alineamos "Elaboración" a 1.2 cm para que empate exacto con el borde del cuadro
    p_fuente.paragraph_format.left_indent = Cm(1.0) 
    
    run_fuente = p_fuente.add_run("Elaboración: Subdirección de Sanción y Gestión de Incentivos (SSAG) – DFAI.")
    run_fuente.font.name = 'Arial' 
    run_fuente.font.size = Pt(8)

def agregar_subtitulo_letra(doc, letra_texto, texto):
    """
    Crea un subtítulo con letra (ej: "A.") donde la letra empieza a 1.0 cm 
    y el texto a 1.5 cm, dando un espacio más natural y compacto.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    
    # --- AJUSTE: Texto a 1.5 cm, la letra retrocede 0.5 cm (arranca en 1.0 cm) ---
    pf.left_indent = Cm(1.5)
    pf.first_line_indent = Cm(-0.5)
    pf.tab_stops.clear_all()
    pf.tab_stops.add_tab_stop(Cm(1.5))
    
    pf.space_before = Pt(14) 
    
    run = p.add_run(f"{letra_texto}\t{texto}")
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    
    run.italic = True
    run.underline = True
    
    return p

def dibujar_tabla_factores_graduacion(doc, datos_graduacion):
    """
    Dibuja la tabla de Factores de Graduación con proporción 3:1 (10.5 cm y 3.5 cm).
    """
    # Función local para colorear celdas
    def color_celda(celda, color_hex):
        tcPr = celda._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    # 1. TÍTULO DEL CUADRO
    if not hasattr(doc, 'contador_cuadros'):
        doc.contador_cuadros = 2
        
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.left_indent = Cm(1.0)
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run(f"Cuadro n.° {doc.contador_cuadros}: Factores para la Graduación de Sanciones")
    run_titulo.font.size = Pt(11)
    run_titulo.font.name = 'Arial'
    run_titulo.bold = True
    doc.contador_cuadros += 1

    # 2. CREACIÓN DE LA TABLA
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = 'Table Grid'
    tabla.autofit = False
    
    # Alineación y ancho total
    tabla.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = tabla._tblPr
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '680') 
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    ancho_desc = Cm(10.5)
    ancho_calif = Cm(3.5)
    
    # 3. ENCABEZADOS
    headers = ["Factores", "Calificación"]
    for i, header in enumerate(headers):
        celda = tabla.rows[0].cells[i]
        celda.width = ancho_desc if i == 0 else ancho_calif
        color_celda(celda, 'E7E6E6') # Gris claro
        
        # Centrado vertical
        tcPr = celda._tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
        
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    # 4. FILAS DE FACTORES (f1 a f7)
    nombres_factores = [
        "f1. Gravedad del daño al interés público y/o bien jurídico protegido",
        "f2. El perjuicio económico causado",
        "f3. Aspectos ambientales o fuentes de contaminación",
        "f4. Reincidencia en la comisión de la infracción",
        "f5. Corrección de la conducta infractora",
        "f6. Adopción de las medidas necesarias para revertir las consecuencias de la conducta infractora",
        "f7. Intencionalidad en la conducta del infractor"
    ]
    
    suma_porcentajes = 0.0
    
    for i, nombre in enumerate(nombres_factores):
        row = tabla.add_row()
        row.cells[0].width = ancho_desc
        row.cells[1].width = ancho_calif
        
        # Texto del factor
        p_desc = row.cells[0].paragraphs[0]
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_desc = p_desc.add_run(nombre)
        run_desc.font.name = 'Arial'
        run_desc.font.size = Pt(10)
        
        # Valor del factor (SIN EL SIGNO '+' FORZADO)
        subtotal_f = datos_graduacion.get(f"subtotal_f{i+1}", 0.0)
        suma_porcentajes += subtotal_f
        # Valor del factor: Python pone el '-' automáticamente, solo limpiamos los "-0%" fantasma
        subtotal_f = datos_graduacion.get(f"subtotal_f{i+1}", 0.0)
        suma_porcentajes += subtotal_f
        texto_porcentaje = f"{subtotal_f:.0%}".replace("-0%", "0%")
        
        p_val = row.cells[1].paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_val = p_val.add_run(texto_porcentaje)
        run_val.font.name = 'Arial'
        run_val.font.size = Pt(10)
        
        # Centrado vertical para la fila de datos
        for celda in row.cells:
            tcPr = celda._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

    # 5. FILA DE SUMATORIA
    row_sum = tabla.add_row()
    row_sum.cells[0].width = ancho_desc
    row_sum.cells[1].width = ancho_calif
    color_celda(row_sum.cells[0], 'E7E6E6')
    color_celda(row_sum.cells[1], 'E7E6E6')
    
    p_sum_desc = row_sum.cells[0].paragraphs[0]
    run_sum_desc = p_sum_desc.add_run("(f1+f2+f3+f4+f5+f6+f7)")
    run_sum_desc.font.name = 'Arial'
    run_sum_desc.font.size = Pt(10)
    run_sum_desc.bold = True
    
    p_sum_val = row_sum.cells[1].paragraphs[0]
    p_sum_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Igual aquí, mostramos el '-' natural pero limpiamos el cero negativo
    run_sum_val = p_sum_val.add_run(f"{suma_porcentajes:.0%}".replace("-0%", "0%"))
    run_sum_val.font.name = 'Arial'
    run_sum_val.font.size = Pt(10)
    run_sum_val.bold = True

    # 6. FILA DEL FACTOR F FINAL
    factor_f_final = datos_graduacion.get('factor_f_calculado', 1.0)
    
    row_f = tabla.add_row()
    row_f.cells[0].width = ancho_desc
    row_f.cells[1].width = ancho_calif
    color_celda(row_f.cells[0], 'E7E6E6')
    color_celda(row_f.cells[1], 'E7E6E6')
    
    p_f_desc = row_f.cells[0].paragraphs[0]
    run_f_desc = p_f_desc.add_run("Factores: F = (1+f1+f2+f3+f4+f5+f6+f7)")
    run_f_desc.font.name = 'Arial'
    run_f_desc.font.size = Pt(10)
    run_f_desc.bold = True
    
    p_f_val = row_f.cells[1].paragraphs[0]
    p_f_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f_val = p_f_val.add_run(f"{factor_f_final:.0%}")
    run_f_val.font.name = 'Arial'
    run_f_val.font.size = Pt(10)
    run_f_val.bold = True

    # --- CENTRADO VERTICAL PARA LAS FILAS TOTALES ---
    for row_final in [row_sum, row_f]:
        for celda in row_final.cells:
            tcPr = celda._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

    # 7. NOTA DE ELABORACIÓN
    p_fuente = doc.add_paragraph()
    p_fuente.paragraph_format.left_indent = Cm(1.0)
    run_fuente = p_fuente.add_run("Elaboración: Subdirección de Sanción y Gestión de Incentivos (SSAG) – DFAI.")
    run_fuente.font.name = 'Arial' 
    run_fuente.font.size = Pt(8)

def dibujar_tabla_multa_propuesta(doc, datos_hecho):
    """
    Dibuja la tabla de Multa Propuesta con proporción 3:1 (10.5 cm y 3.5 cm)
    y sombrea la última fila.
    """
    # Función local para colorear celdas
    def color_celda(celda, color_hex):
        tcPr = celda._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    # 1. Extracción de datos
    resultados = datos_hecho.get('resultados', {})
    totales_app = resultados.get('resultados_para_app', {}).get('totales', {})
    if not totales_app:
        totales_app = resultados.get('resultados_para_app', {})
        
    contexto_word = resultados.get('contexto_final_word', {})

    bi_uit = totales_app.get('beneficio_ilicito_uit', 0.0)
    p_det = totales_app.get('probabilidad', 1.0) 

    aplica_graduacion = datos_hecho.get('aplica_graduacion', 'No')
    if aplica_graduacion == 'Sí':
        factor_f = datos_hecho.get('graduacion', {}).get('factor_f_calculado', 1.0)
    else:
        factor_f = 1.0
    
    # --- LA CORRECCIÓN: Búsqueda exhaustiva de la multa ---
    multa_uit_cruda = contexto_word.get('multa_original_uit')
    if not multa_uit_cruda:
        multa_uit_cruda = totales_app.get('multa_final_uit', totales_app.get('multa_base_uit', 0.0))
        
    # Limpiamos el valor por si viene como texto (ej: "0.250 UIT")
    import re
    try:
        multa_uit = float(re.search(r'\d+(\.\d+)?', str(multa_uit_cruda).replace(',', '')).group())
    except Exception:
        multa_uit = 0.0

    # 2. Título de la tabla
    if not hasattr(doc, 'contador_cuadros'):
        doc.contador_cuadros = 2
        
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.left_indent = Cm(1.0)
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run(f"Cuadro n.° {doc.contador_cuadros}: Multa propuesta")
    run_titulo.font.size = Pt(11)
    run_titulo.font.name = 'Arial'
    run_titulo.bold = True
    doc.contador_cuadros += 1

    # 3. Creación de la tabla
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = 'Table Grid'
    tabla.autofit = False
    
    # Alineación a 1.0 cm simulada (680 dxa)
    tabla.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = tabla._tblPr
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '680') 
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    ancho_desc = Cm(10.5)
    ancho_val = Cm(3.5)

    # 4. Encabezados
    headers = ["Componentes", "Monto"]
    for i, header in enumerate(headers):
        celda = tabla.rows[0].cells[i]
        celda.width = ancho_desc if i == 0 else ancho_val
        color_celda(celda, 'E7E6E6')
        
        tcPr = celda._tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
        
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    # 5. Filas de datos
    filas = [
        ("Beneficio Ilícito (B)", f"{bi_uit:,.3f} UIT"),
        ("Probabilidad de detección (p)", f"{p_det:,.3f}"),
        ("Factores para la graduación de sanciones\nF=(1+f1+f2+f3+f4+f5+f6+f7)", f"{factor_f:,.2%}"),
        ("Multa en UIT (B/p)*(F)", f"{multa_uit:,.3f} UIT")
    ]

    for idx, (desc, val) in enumerate(filas):
        row = tabla.add_row()
        row.cells[0].width = ancho_desc
        row.cells[1].width = ancho_val
        
        es_ultima = (idx == len(filas) - 1)
        if es_ultima:
            # Sombreado gris para la última fila
            color_celda(row.cells[0], 'E7E6E6')
            color_celda(row.cells[1], 'E7E6E6')

        p_desc = row.cells[0].paragraphs[0]
        p_desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_d = p_desc.add_run(desc)
        run_d.font.name = 'Arial'
        run_d.font.size = Pt(10)
        if es_ultima: run_d.bold = True

        p_val = row.cells[1].paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_v = p_val.add_run(val)
        run_v.font.name = 'Arial'
        run_v.font.size = Pt(10)
        if es_ultima: run_v.bold = True
        
        # Centrado vertical para todas las celdas
        for celda in row.cells:
            tcPr = celda._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

    # 6. Nota de elaboración
    p_elab = doc.add_paragraph()
    p_elab.paragraph_format.left_indent = Cm(1.0)
    run_elab = p_elab.add_run("Elaboración: Subdirección de Sanción y Gestión de Incentivos (SSAG) – DFAI.")
    run_elab.font.name = 'Arial'
    run_elab.font.size = Pt(8)

def dibujar_tabla_resumen_multas(doc, imputaciones_data, multa_total_uit_str):
    """
    Dibuja la tabla final de Resumen de Multas (Cuadro 8 en tu ejemplo),
    fusionando las celdas de la última fila para el "Total".
    """
    def color_celda(celda, color_hex):
        tcPr = celda._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    # 1. TÍTULO DEL CUADRO
    if not hasattr(doc, 'contador_cuadros'):
        doc.contador_cuadros = 2
        
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.left_indent = Cm(1.0)
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- LÓGICA PLURAL/SINGULAR PARA EL TÍTULO ---
    cantidad_hechos = len(imputaciones_data)
    texto_titulo_resumen = "Resumen de multa" if cantidad_hechos == 1 else "Resumen de multas"
    
    run_titulo = p_titulo.add_run(f"Cuadro n.° {doc.contador_cuadros}: {texto_titulo_resumen}")
    run_titulo.font.size = Pt(11)
    run_titulo.font.name = 'Arial'
    run_titulo.bold = True
    doc.contador_cuadros += 1

    # 2. CREACIÓN DE LA TABLA
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = 'Table Grid'
    tabla.autofit = False
    
    # 1. Alineación idéntica a la tabla de Multa Propuesta
    tabla.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = tabla._tblPr
    
    # ¡LA LÍNEA MÁGICA!: Fuerza a Word a respetar medidas fijas y no auto-expandirse
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '680') 
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # 2. Proporción 1:2:1 para un total exacto de 14.0 cm (3.5 + 7.0 + 3.5)
    anchos = [Cm(3.5), Cm(7.0), Cm(3.5)]
    
    # 3. ENCABEZADOS
    headers = ["Numeral", "Infracción", "Multa"]
    for i, header in enumerate(headers):
        celda = tabla.rows[0].cells[i]
        celda.width = anchos[i]
        color_celda(celda, 'E7E6E6')
        
        tcPr = celda._tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
        
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    # 4. FILAS DE DATOS
    import re
    for i, datos_hecho in enumerate(imputaciones_data):
        row = tabla.add_row()
        row.cells[0].width = anchos[0]
        row.cells[1].width = anchos[1]
        row.cells[2].width = anchos[2]
        
        resultados = datos_hecho.get('resultados', {})
        totales_app = resultados.get('resultados_para_app', {}).get('totales', {})
        if not totales_app: totales_app = resultados.get('resultados_para_app', {})
        contexto_word = resultados.get('contexto_final_word', {})

        if datos_hecho.get('aplica_reduccion', 'No') == 'Sí':
            val_crudo = totales_app.get('multa_con_reduccion_uit', 0.0)
        else:
            val_crudo = contexto_word.get('multa_original_uit')
            if not val_crudo:
                val_crudo = totales_app.get('multa_final_uit', totales_app.get('multa_base_uit', 0.0))

        try:
            val_num = float(re.search(r'\d+(\.\d+)?', str(val_crudo).replace(',', '')).group())
        except Exception:
            val_num = 0.0
            
        numeral = f"4.{i+2}"
        infraccion = f"Hecho imputado n.° {i+1}"
        multa_str = f"{val_num:,.3f} UIT"
        
        valores = [numeral, infraccion, multa_str]
        for j, val in enumerate(valores):
            celda = row.cells[j]
            tcPr = celda._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
            
            p = celda.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.name = 'Arial'
            run.font.size = Pt(10)

    # 5. FILA TOTAL (Celdas 0 y 1 combinadas)
    row_total = tabla.add_row()
    
    # Asignamos anchos individuales ANTES de fusionar para que Word no se pierda
    row_total.cells[0].width = anchos[0]
    row_total.cells[1].width = anchos[1]
    row_total.cells[2].width = anchos[2]
    
    # Fusionamos
    row_total.cells[0].merge(row_total.cells[1])
    celda_total_lbl = row_total.cells[0]
    celda_total_val = row_total.cells[2]
    
    # REFORZAMOS el ancho explícito: 3.5 + 7.0 = 10.5 cm
    celda_total_lbl.width = Cm(10.5)
    celda_total_val.width = Cm(3.5)
    
    color_celda(celda_total_lbl, 'E7E6E6')
    color_celda(celda_total_val, 'E7E6E6')
    
    # Texto "Total"
    tcPr1 = celda_total_lbl._tc.get_or_add_tcPr()
    vAlign1 = OxmlElement('w:vAlign')
    vAlign1.set(qn('w:val'), 'center')
    tcPr1.append(vAlign1)
    p_tlbl = celda_total_lbl.paragraphs[0]
    p_tlbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tlbl = p_tlbl.add_run("Total")
    run_tlbl.font.name = 'Arial'
    run_tlbl.font.size = Pt(10)
    run_tlbl.bold = True
    
    # Texto Valor Final
    tcPr2 = celda_total_val._tc.get_or_add_tcPr()
    vAlign2 = OxmlElement('w:vAlign')
    vAlign2.set(qn('w:val'), 'center')
    tcPr2.append(vAlign2)
    p_tval = celda_total_val.paragraphs[0]
    p_tval.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tval = p_tval.add_run(f"{multa_total_uit_str} UIT")
    run_tval.font.name = 'Arial'
    run_tval.font.size = Pt(10)
    run_tval.bold = True

    # 6. NOTA AL PIE DE TABLA
    p_elab = doc.add_paragraph()
    p_elab.paragraph_format.left_indent = Cm(1.0)
    run_elab = p_elab.add_run("Elaboración: Subdirección de Sanción y Gestión de Incentivos (SSAG) – DFAI.")
    run_elab.font.name = 'Arial'
    run_elab.font.size = Pt(8)

def dibujar_tabla_anexo_graduacion(doc, datos_hecho, num_hecho_idx):
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    num_hecho = num_hecho_idx + 1

    # 1. Subtítulo: Hecho imputado n.° X
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"Hecho imputado n.° {num_hecho}")
    run_sub.bold = True
    run_sub.underline = True
    run_sub.font.size = Pt(11)
    run_sub.font.name = 'Arial'

    doc.add_paragraph()

    # 2. Título de tabla
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tit = p_tit.add_run("Tabla n.° 1: Factores para la Graduación de las Sanciones")
    run_tit.bold = True
    run_tit.font.size = Pt(11)
    run_tit.font.name = 'Arial'

    # 3. Creación de tabla y XML Mágico
    tabla = doc.add_table(rows=2, cols=4)
    tabla.style = 'Table Grid'

    tblPr = tabla._tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    # --- AJUSTE: ÍTEM a 1.0 cm, CRITERIOS a 9.5 cm ---
    anchos = [Cm(1.2), Cm(9.5), Cm(2.5), Cm(1.8)]

    for row in tabla.rows:
        for i, cell in enumerate(row.cells):
            cell.width = anchos[i]

    # 4. Construcción de Encabezados con celdas combinadas
    tabla.rows[0].cells[0].merge(tabla.rows[1].cells[0])
    tabla.rows[0].cells[1].merge(tabla.rows[1].cells[1])
    tabla.rows[0].cells[3].merge(tabla.rows[1].cells[3])

    def format_header_cell(cell, text):
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
        
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E7E6E6')
        tcPr.append(shd)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Arial'

    format_header_cell(tabla.rows[0].cells[0], "ÍTEM")
    format_header_cell(tabla.rows[0].cells[1], "CRITERIOS")
    format_header_cell(tabla.rows[0].cells[2], "CALIFICACIÓN")
    format_header_cell(tabla.rows[1].cells[2], "DAÑO\nPOTENCIAL")
    format_header_cell(tabla.rows[0].cells[3], "SUBTOTAL")

    # 5. Diccionario estructural oficial de las opciones
    blocks = [
        {"type": "header", "item": "f1", "titulo": "GRAVEDAD DEL DAÑO AL AMBIENTE"},
        {"type": "sub", "item": "1.1", "titulo_italica": "El daño involucra uno o más de los siguientes Componentes Ambientales: a) Agua, b) Suelo, c) Aire, d) Flora y e) Fauna.", "opciones": [("El daño afecta a un (01) componente ambiental.", "10%"), ("El daño afecta a dos (02) componentes ambientales.", "20%"), ("El daño afecta a tres (03) componentes ambientales.", "30%"), ("El daño afecta a cuatro (04) componentes ambientales.", "40%"), ("El daño afecta a cinco (05) componentes ambientales.", "50%")], "key_val": "1.1 Componentes Ambientales", "factor": "f1"},
        {"type": "sub", "item": "1.2", "titulo_italica": "Grado de incidencia en la calidad del ambiente.", "opciones": [("Impacto mínimo.", "6%"), ("Impacto regular.", "12%"), ("Impacto alto.", "18%"), ("Impacto total.", "24%")], "key_val": "1.2 Incidencia en la calidad", "factor": "f1"},
        {"type": "sub", "item": "1.3", "titulo_italica": "Según la extensión geográfica.", "opciones": [("El impacto está localizado en el área de influencia directa.", "10%"), ("El impacto está localizado en el área de influencia indirecta.", "20%")], "key_val": "1.3 Extensión geográfica", "factor": "f1"},
        {"type": "sub", "item": "1.4", "titulo_italica": "Sobre la reversibilidad/recuperabilidad.", "opciones": [("Reversible en el corto plazo.", "6%"), ("Recuperable en el corto plazo.", "12%"), ("Recuperable en el mediano plazo.", "18%"), ("Recuperable en el largo plazo o irrecuperable.", "24%")], "key_val": "1.4 Reversibilidad/Recuperabilidad", "factor": "f1"},
        {"type": "sub", "item": "1.5", "titulo_italica": "Afectación sobre recursos naturales, área natural protegida o zona de amortiguamiento.", "opciones": [("No existe afectación o esta es indeterminable con la información disponible.", "0%"), ("El impacto se ha producido en un área natural protegida, zona de amortiguamiento o ha afectado recursos naturales declarados en alguna categoría de amenaza o en peligro de extinción, o sobre los cuales exista veda, restricción o prohibición de su aprovechamiento.", "40%")], "key_val": "1.5 Afectación recursos/áreas protegidas", "factor": "f1"},
        {"type": "sub", "item": "1.6", "titulo_italica": "Afectación a comunidades nativas o campesinas.", "opciones": [("No afecta a comunidades nativas o campesinas.", "0%"), ("Afecta a una comunidad nativa o campesina.", "15%"), ("Afecta a más de una comunidad nativa o campesina.", "30%")], "key_val": "1.6 Afectación comunidades", "factor": "f1"},
        {"type": "sub", "item": "1.7", "titulo_italica": "Afectación a la salud de las personas", "opciones": [("No afecta a la salud de las personas o no se puede determinar con la información disponible.", "0%"), ("Afecta la salud de las personas.", "60%")], "key_val": "1.7 Afectación salud", "factor": "f1"},
        {"type": "sub", "item": "f2.", "titulo_bold": "PERJUICIO ECONÓMICO CAUSADO: El perjuicio económico causado es mayor en una población más desprotegida, lo que se refleja en la incidencia de pobreza total.", "titulo_italica": "Incidencia de pobreza total", "opciones": [("El impacto ocurre en una zona con incidencia de pobreza total hasta 19,6%.", "4%"), ("El impacto ocurre en una zona con incidencia de pobreza total mayor a 19,6% hasta 39,1%.", "8%"), ("El impacto ocurre en una zona con incidencia de pobreza total mayor a 39,1% hasta 58,7%.", "12%"), ("El impacto ocurre en una zona con incidencia de pobreza total mayor a 58,7% hasta 78,2%.", "16%"), ("El impacto ocurre en una zona con incidencia de pobreza total mayor a 78,2%.", "20%")], "key_val": "Incidencia de pobreza total", "factor": "f2"},
        {"type": "sub", "item": "f3.", "titulo_bold": "ASPECTOS AMBIENTALES O FUENTES DE CONTAMINACIÓN: efluentes, residuos sólidos, efluentes atmosféricos, ruido, radiaciones no ionizantes, u otras.", "opciones": [("El impacto involucra un (01) aspecto ambiental o fuente de contaminación.", "6%"), ("El impacto involucra dos (02) aspectos ambientales o fuentes de contaminación.", "12%"), ("El impacto involucra tres (03) aspectos ambientales o fuentes de contaminación.", "18%"), ("El impacto involucra cuatro (04) aspectos ambientales o fuentes de contaminación.", "24%"), ("El impacto involucra cinco (05) aspectos ambientales o fuentes de contaminación.", "30%")], "key_val": "Cantidad de aspectos", "factor": "f3"},
        {"type": "sub", "item": "f4.", "titulo_bold": "REINCIDENCIA EN LA COMISIÓN DE LA INFRACCIÓN:", "opciones": [("Por la comisión de actos u omisiones que constituyan la misma infracción dentro del plazo de un (01) año desde que quedó firme la resolución de la sanciona la primera infracción", "20%")], "key_val": "Reincidencia", "factor": "f4"},
        {"type": "sub", "item": "f5.", "titulo_bold": "CORRECCIÓN DE LA CONDUCTA INFRACTORA:", "opciones": [("El administrado subsana el acto u omisión imputada como constitutivo de infracción administrativa de manera voluntaria, antes del inicio del procedimiento administrativo sancionador.", "Eximente"), ("El administrado, a requerimiento de la autoridad, corrige el acto u omisión imputada como constitutivo de infracción administrativa, calificada como incumplimiento leve, antes del inicio del procedimiento administrativo sancionador. Dicha corrección debe estar adecuadamente acreditada", "Eximente"), ("El administrado, a requerimiento de la autoridad, corrige el acto u omisión imputada como constitutivo de infracción administrativa, calificada como incumplimiento trascendente, antes del inicio del procedimiento administrativo sancionador. Dicha corrección debe estar adecuadamente acreditada", "-40%"), ("El administrado, a requerimiento de la autoridad, corrige el acto u omisión imputada como constitutivo de infracción administrativa, luego del inicio del procedimiento administrativo sancionador, antes de la resolución final de primera instancia. Dicha corrección debe estar adecuadamente acreditada", "-20%")], "key_val": "Subsanación/Corrección", "factor": "f5"},
        {"type": "sub", "item": "f6.", "titulo_bold": "ADOPCIÓN DE LAS MEDIDAS NECESARIAS PARA REVERTIR LAS CONSECUENCIAS DE LA CONDUCTA INFRACTORA", "opciones": [("No ejecutó ninguna medida.", "30%"), ("Ejecutó medidas tardías.", "20%"), ("Ejecutó medidas parciales.", "10%"), ("Ejecutó medidas necesarias e inmediatas para remediar los efectos de la conducta infractora.", "-10%")], "key_val": "Medidas adoptadas", "factor": "f6"},
        {"type": "sub", "item": "f7.", "titulo_bold": "INTENCIONALIDAD EN LA CONDUCTA DEL INFRACTOR:", "opciones": [("Cuando se acredita o verifica la intencionalidad.", "72%")], "key_val": "Intencionalidad", "factor": "f7"}
    ]

    grad_data = datos_hecho.get('graduacion', {})

    # --- NUEVAS FUNCIONES DE DIBUJO ---
    def color_celda(celda, color_hex):
        tcPr = celda._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def add_cell_text(cell, text, bold=False, italic=False, align='LEFT'):
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center') # Centrado vertical activado
        tcPr.append(vAlign)
        p = cell.paragraphs[0]
        
        # Soporte para texto justificado
        if align == 'CENTER': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'RIGHT': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'JUSTIFY': p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.name = 'Arial'
        if bold: run.bold = True
        if italic: run.italic = True

# 6. Rellenar las filas iterando sobre la estructura
    for block in blocks:
        if block['type'] == 'header':
            row = tabla.add_row()
            for i in range(4): 
                row.cells[i].width = anchos[i]
                color_celda(row.cells[i], 'E7E6E6') # Fondo gris a toda la fila
                
            add_cell_text(row.cells[0], block['item'], bold=True, align='CENTER')
            add_cell_text(row.cells[1], block['titulo'], bold=True, align='JUSTIFY') # Justificado
        
        elif block['type'] == 'sub':
            # Extraer el valor guardado
            key = f"grad_{num_hecho_idx}_{block['factor']}_{block['key_val']}_valor"
            val = grad_data.get(key, 0.0)
            if val == "Eximente": subtotal_str = "Eximente"
            else: subtotal_str = f"{float(val):+.0%}".replace("+0%", "0%")
            
            # Fila de Título del Sub-criterio
            row_tit = tabla.add_row()
            for i in range(4): 
                row_tit.cells[i].width = anchos[i]
                # APLICAMOS GRIS A TODA LA FILA PARA EVITAR EL "CORTE"
                color_celda(row_tit.cells[i], 'E7E6E6')
            
            add_cell_text(row_tit.cells[0], block['item'], bold=True, align='CENTER')
            
            tcPr_tit = row_tit.cells[1]._tc.get_or_add_tcPr()
            vAlign_tit = OxmlElement('w:vAlign')
            vAlign_tit.set(qn('w:val'), 'center')
            tcPr_tit.append(vAlign_tit)
            
            p_tit = row_tit.cells[1].paragraphs[0]
            p_tit.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
            
            if 'titulo_bold' in block:
                r_b = p_tit.add_run(block['titulo_bold'] + " ")
                r_b.bold = True; r_b.font.size = Pt(8); r_b.font.name = 'Arial'
            if 'titulo_italica' in block:
                r_i = p_tit.add_run(block['titulo_italica'])
                r_i.italic = True; r_i.bold = True; r_i.font.size = Pt(8); r_i.font.name = 'Arial'
            
            # Dejamos la celda de Subtotal vacía en la fila gris para que la línea sea continua
            add_cell_text(row_tit.cells[3], "")
            
            # Filas de Opciones
            start_row_idx = len(tabla.rows) # Índice de la primera opción
            for idx_op, (op_text, op_pct) in enumerate(block['opciones']):
                r_op = tabla.add_row()
                for i in range(4): r_op.cells[i].width = anchos[i]
                add_cell_text(r_op.cells[1], op_text, align='JUSTIFY')
                add_cell_text(r_op.cells[2], op_pct, align='CENTER')
                
                # Inyectamos el Subtotal SOLO en la primera celda blanca de opción
                if idx_op == 0:
                    add_cell_text(r_op.cells[3], subtotal_str, bold=True, align='CENTER')
            
            end_row_idx = len(tabla.rows) - 1
            
            # Combinamos la celda "Subtotal" SOLO a lo largo de las opciones (sin tocar el título gris)
            if end_row_idx > start_row_idx:
                tabla.rows[start_row_idx].cells[3].merge(tabla.rows[end_row_idx].cells[3])

    # 7. Fila Final (Totales)
    row_tot = tabla.add_row()
    for i in range(4): row_tot.cells[i].width = anchos[i]
    row_tot.cells[0].merge(row_tot.cells[2])
    
    p_tot = row_tot.cells[0].paragraphs[0]
    p_tot.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_tot = p_tot.add_run("Total Factores para la Graduación de Sanciones: F=(1 + f1+f2+f3+f4+f5+f6+f7)")
    r_tot.bold = True; r_tot.font.size = Pt(8); r_tot.font.name = 'Arial'
    
    factor_f = grad_data.get('factor_f_calculado', 1.0)
    add_cell_text(row_tot.cells[3], f"{factor_f:.0%}", bold=True, align='CENTER')
    
    # Fondo Gris a la fila final
    shd1 = OxmlElement('w:shd')
    shd1.set(qn('w:val'), 'clear'); shd1.set(qn('w:color'), 'auto'); shd1.set(qn('w:fill'), 'E7E6E6')
    row_tot.cells[0]._tc.get_or_add_tcPr().append(shd1)

    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'), 'clear'); shd2.set(qn('w:color'), 'auto'); shd2.set(qn('w:fill'), 'E7E6E6')
    row_tot.cells[3]._tc.get_or_add_tcPr().append(shd2)

    # 8. Nota al pie de tabla
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.left_indent = Cm(0)
    r_foot = p_foot.add_run("Elaboración: Subdirección de Sanción y Gestión de Incentivos (SSAG)-DFAI.")
    r_foot.font.size = Pt(8)
    r_foot.font.name = 'Arial'

# =========================================================
# FUNCIÓN MAESTRA PARA GENERAR EL DOCUMENTO DESDE STREAMLIT
# =========================================================
def generar_documento_ifi(info_expediente, context_data, imputaciones_data, conf_data):
    from datetime import date
    
    # --- FUNCIÓN AUXILIAR: Capitalizar nombres inteligentemente en español ---
    def capitalizar_nombre_espanol(nombre):
        if not nombre: return ""
        palabras_menores = ["de", "la", "el", "los", "las", "y", "en", "por", "del", "a", "las"]
        palabras = nombre.title().split()
        resultado = []
        for i, p in enumerate(palabras):
            # Mantenemos las palabras menores en minúscula, salvo que sea la primera palabra
            if p.lower() in palabras_menores and i > 0:
                resultado.append(p.lower())
            else:
                resultado.append(p)
        return " ".join(resultado)

    # --- 1. PREPARAMOS LAS VARIABLES GENERALES DESDE STREAMLIT ---
    ht_val = info_expediente.get("HT", "[HT NO ENCONTRADO]")
    num_informe_full = context_data.get("num_informe_multa_ifi", "XXXXX")
    num_informe = num_informe_full.split('-')[0].replace('N° ', '').replace('n.°', '').strip() if num_informe_full else "XXXXX"
    
    # 1.1 Administrado (Con mayúsculas y minúsculas corregidas)
    administrado_crudo = context_data.get("administrado", "[ADMINISTRADO]")
    administrado = capitalizar_nombre_espanol(administrado_crudo)
    
    # 1.2 Referencia (Ahora usa el número de Expediente en lugar de la RSD)
    num_expediente = context_data.get("expediente", "[N.° DE EXPEDIENTE]")
    referencia = f"Expediente n.° {num_expediente}"
    
    fecha_txt = context_data.get("fecha_hoy", "[FECHA]")

    # ========================================================
    # --- 2. CONSTRUCCIÓN DE LOS BLOQUES "A" Y "DE" ---
    # ========================================================
    
    nombre_sub1 = context_data.get("nombre_encargado_sub1", "[NOMBRE A]")
    cargo_sub1 = context_data.get("cargo_encargado_sub1", "[CARGO A]")
    bloque_a = [(nombre_sub1, True), (cargo_sub1, False)]

    titulo_sub2 = context_data.get("titulo_encargado_sub2", "")
    nombre_sub2 = context_data.get("nombre_encargado_sub2", "[NOMBRE SUBDIRECTOR]")
    cargo_sub2 = context_data.get("cargo_encargado_sub2", "[CARGO SUBDIRECTOR]")
    col_sub2 = context_data.get("colegiatura_encargado_sub2", "")
    
    titulo_rev = context_data.get("titulo_revisor", "")
    nombre_rev = context_data.get("nombre_revisor", "[NOMBRE REVISOR]")
    cargo_rev = context_data.get("cargo_revisor", "[CARGO REVISOR]")
    col_rev = context_data.get("colegiatura_revisor", "")
    
    titulo_ana = context_data.get("titulo_analista", "")
    nombre_ana = context_data.get("nombre_analista", "[NOMBRE ANALISTA]")
    cargo_ana = context_data.get("cargo_analista", "[CARGO ANALISTA]")

    def armar_bloque_persona(titulo, nombre, cargo, colegiatura=""):
        linea_nom = f"{titulo} {nombre}".strip() if titulo else nombre
        lineas = [(linea_nom, True), (cargo, False)]
        if colegiatura:
            lineas.append((colegiatura, False))
        return lineas

    bloque_sub2 = armar_bloque_persona(titulo_sub2, nombre_sub2, cargo_sub2, col_sub2)
    bloque_rev = armar_bloque_persona(titulo_rev, nombre_rev, cargo_rev, col_rev)
    bloque_ana = armar_bloque_persona(titulo_ana, nombre_ana, cargo_ana)
    bloque_de = bloque_sub2 + [("", False)] + bloque_rev + [("", False)] + bloque_ana

    # --- 3. VARIABLES PARA HECHOS Y ANEXOS ---
    hechos_imputados = [hecho.get('texto_hecho', '') for hecho in imputaciones_data]
    hay_factores_graduacion = any(h.get('aplica_graduacion') == 'Sí' for h in imputaciones_data)
    
    # 3.1 Asunto (Dinámico según la cantidad de hechos)
    cantidad_hechos = len(hechos_imputados)
    texto_asunto = "Propuesta de cálculo de multa" if cantidad_hechos == 1 else "Propuesta de cálculo de multas"

    # ========================================================
    # --- 4. INICIO DE LA CONSTRUCCIÓN DEL WORD DESDE CERO ---
    # ========================================================
    doc = Document()
    
    configurar_pagina_a4(doc)
    configurar_estilo_base(doc)
    
    texto_encabezado = "Decenio de la Igualdad de Oportunidades para Mujeres y Hombres\nAño de la Esperanza y el Fortalecimiento de la Democracia"
    crear_encabezado_ifi(doc, "logo.png", texto_encabezado)
    
    p_ht = doc.add_paragraph()
    p_ht.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_ht = p_ht.add_run(ht_val)
    run_ht.font.name = 'Arial'
    run_ht.font.size = Pt(8)
    run_ht.bold = True

    doc.add_paragraph() 
    
    anio_actual = date.today().year
    texto_titulo = f"INFORME n.° {num_informe}-{anio_actual}-OEFA/DFAI-SSAG"
    
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run(texto_titulo)
    run_titulo.font.name = 'Arial'
    run_titulo.font.size = Pt(11)
    run_titulo.bold = True

    doc.add_paragraph() 
    
    # --- BLOQUE DE CABECERA DINÁMICO ---
    datos_para_cabecera = [
        ("A", bloque_a),
        ("DE", bloque_de),
        ("ASUNTO", texto_asunto), # <--- Aquí inyectamos el asunto plural o singular
        ("REFERENCIA", referencia), # <--- Aquí inyectamos el expediente
        ("ADMINISTRADO", administrado), # <--- Aquí inyectamos el nombre capitalizado
        ("FECHA", f"Lima, {fecha_txt}")
    ]
    
    crear_bloque_cabecera(doc, datos_para_cabecera)
    agregar_linea_horizontal(doc)
    
    # ==========================================
    # 1. ANTECEDENTES
    # ==========================================
    agregar_titulo_numerado(doc, 1, "Antecedentes")
    doc.add_paragraph()
    
    # 1. Extraemos los datos del context_data
    numero_rsd = context_data.get('numero_rsd', '[N.° RSD]') 
    fecha_rsd_texto = context_data.get('fecha_rsd_texto', '[FECHA RSD]')
    subdireccion = context_data.get('subdireccion', '[SUBDIRECCIÓN]')
    id_subdireccion = context_data.get('id_subdireccion', '[SIGLAS]')
    
    # 2. Lógica dinámica real para CUALQUIER cantidad de infracciones (usando num2words)
    from num2words import num2words
    
    if cantidad_hechos == 1:
        texto_infracciones = "una (1) presunta infracción administrativa"
    else:
        # Convertimos el número a texto en español (ej. 121 -> "ciento veintiuno")
        texto_num = num2words(cantidad_hechos, lang='es')
        # Cambiamos "uno" por "una" para que concuerde en femenino con "infracciones" (ej. "ciento veintiuna")
        if texto_num.endswith("uno"):
            texto_num = texto_num[:-3] + "una"
            
        texto_infracciones = f"{texto_num} ({cantidad_hechos}) presuntas infracciones administrativas"

    # 3. Construimos el primer párrafo por "tramos" (runs) para aplicar estilos
    p_texto = doc.add_paragraph()
    p_texto.paragraph_format.left_indent = Cm(1.0)

    # Función auxiliar para añadir pedacitos de texto con o sin formato
    def agregar_tramo(parrafo, texto, destacado=False):
        run = parrafo.add_run(texto)
        if destacado:
            run.bold = True
            run.italic = True

    # Armamos el Párrafo 1 como si fuera un rompecabezas
    agregar_tramo(p_texto, f"Mediante la Resolución Subdirectoral n.° {numero_rsd} (en adelante, ")
    agregar_tramo(p_texto, "la Resolución Subdirectoral", destacado=True)
    agregar_tramo(p_texto, f"), notificada el {fecha_rsd_texto}, la {subdireccion} (en adelante, ")
    agregar_tramo(p_texto, f"la {id_subdireccion}", destacado=True)
    agregar_tramo(p_texto, ") de la Dirección de Fiscalización y Aplicación de Incentivos (en adelante, ")
    agregar_tramo(p_texto, "la DFAI", destacado=True)
    agregar_tramo(p_texto, ") del Organismo de Evaluación y Fiscalización Ambiental (en adelante, ")
    agregar_tramo(p_texto, "el OEFA", destacado=True)
    agregar_tramo(p_texto, "), inició el procedimiento administrativo sancionador (en adelante, ")
    agregar_tramo(p_texto, "el PAS", destacado=True)
    agregar_tramo(p_texto, f") a {administrado} (en adelante, ")
    agregar_tramo(p_texto, "el administrado", destacado=True)
    agregar_tramo(p_texto, f") por la comisión de {texto_infracciones}.")

    doc.add_paragraph() # Enter de separación
    
    # 4. Lógica singular/plural para el segundo párrafo
    if cantidad_hechos == 1:
        texto_calculo = "cálculo de multa"
        texto_referidos = "del hecho imputado referido"
    else:
        texto_calculo = "cálculo de multas"
        texto_referidos = "de los hechos imputados referidos"
    
    # 5. Armamos el Párrafo 2 con la misma lógica de tramos
    p_texto_2 = doc.add_paragraph()
    p_texto_2.paragraph_format.left_indent = Cm(1.0)
    
    agregar_tramo(p_texto_2, f"En ese sentido, y en base a la información que obra en el Expediente n.° {num_expediente}, la Subdirección de Sanción y Gestión de Incentivos (en adelante, ")
    agregar_tramo(p_texto_2, "la SSAG", destacado=True)
    
    # Aquí inyectamos nuestras variables calculadas inteligentemente
    agregar_tramo(p_texto_2, f"), a través del presente informe, efectuará la propuesta de {texto_calculo} {texto_referidos} en la Resolución Subdirectoral:")

    doc.add_paragraph() # Enter de separación antes de la lista de hechos
    
    # Inyectamos la lista dinámica con su prefijo en negrita
    for i, hecho in enumerate(hechos_imputados):
        prefijo = f"Hecho imputado n.° {i + 1}: "
        agregar_vineta_flecha(doc, hecho, prefijo_negrita=prefijo)

    # ==========================================
    # --- 2. OBJETO ---
    # ==========================================
    agregar_titulo_numerado(doc, 2, "Objeto")
    doc.add_paragraph()
    
    # Lógica singular/plural para el objeto
    # (Reutilizamos la variable `texto_calculo` que ya vale "cálculo de multa" o "cálculo de multas")
    if cantidad_hechos == 1:
        texto_mencionados = "al hecho imputado mencionado"
    else:
        texto_mencionados = "a los hechos imputados mencionados"
        
    p_objeto = doc.add_paragraph(
        f"El presente informe tiene como objeto realizar la propuesta de {texto_calculo} "
        f"correspondiente {texto_mencionados} en el numeral precedente."
    )
    p_objeto.paragraph_format.left_indent = Cm(1.0)

    # ==========================================
    # --- 3. FÓRMULA PARA EL CÁLCULO DE MULTA ---
    # ==========================================
    agregar_titulo_numerado(doc, 3, "Fórmula para el cálculo de multa")
    agregar_subtitulo_numerado(doc, "3.1.", "Fórmula")
    doc.add_paragraph()

    p_razonabilidad = doc.add_paragraph()
    p_razonabilidad.paragraph_format.left_indent = Cm(1.0)

    texto_principal = (
        "La multa se calcula al amparo del principio de razonabilidad que rige la potestad "
        "sancionadora de la administración, de acuerdo con lo establecido en el numeral 3 del "
        "artículo 248° del Texto Único Ordenado de la Ley del Procedimiento Administrativo General – TUO de la LPAG"
    )
    p_razonabilidad.add_run(texto_principal)

    texto_nota = (
        "Decreto Supremo n.° 004-2019-JUS, que aprueba el Texto Único Ordenado de la Ley n.° 27444 - Ley del "
        "Procedimiento Administrativo General. Procedimiento Sancionador\n"
        "Artículo 248°. - Principios de la potestad sancionadora administrativa\n"
        "La potestad sancionadora de todas las entidades está regida adicionalmente por los siguientes principios especiales: (…)\n"
        "3. Razonabilidad. - Las autoridades deben prever que la comisión de la conducta sancionable no resulte más ventajosa "
        "para el infractor que cumplir las normas infringidas o asumir la sanción. Sin embargo, las sanciones a ser aplicadas "
        "deberán ser proporcionales al incumplimiento calificado como infracción, observando los siguientes criterios que se señalan "
        "a efectos de su graduación:\n"
        "a) El beneficio ilícito resultante por la comisión de la infracción;\n"
        "b) La probabilidad de detección de la infracción;\n"
        "c) La gravedad del daño al interés público y/o bien jurídico protegido;\n"
        "d) EI perjuicio económico causado;\n"
        "e) La reincidencia, por la comisión de la misma infracción dentro del plazo de un año desde que quedó firme la resolución que sancionó la primera infracción.\n"
        "f) Las circunstancias de la comisión de la infracción; y\n"
        "g) La existencia o no de intencionalidad en la conducta del infractor. (…)"
    )
    agregar_nota_al_pie(p_razonabilidad, texto_nota, doc)
    p_razonabilidad.add_run(".")

    doc.add_paragraph()

    p_explicacion = doc.add_paragraph()
    p_explicacion.paragraph_format.left_indent = Cm(1.0)
    p_explicacion.add_run(
        "La fórmula para el cálculo de la multa a ser aplicada considera el beneficio ilícito (B), "
        "dividido entre la probabilidad de detección (p); este resultado es multiplicado por un factor F, "
        "cuyo valor considera los factores para la graduación de sanciones establecidos en la "
        "metodología de cálculo de multas del OEFA"
    )

    texto_nota_mcm = (
        "La Metodología para el cálculo de las multas base y la aplicación de los factores "
        "de gradualidad a utilizar en la graduación de sanciones fue aprobada mediante "
        "Resolución de Presidencia del Consejo Directivo n.° 035-2013-OEFA/PCD y "
        "modificada por Resolución de Consejo Directivo n.° 024-2017-OEFA/CD."
    )
    agregar_nota_al_pie(p_explicacion, texto_nota_mcm, doc)

    p_explicacion.add_run(" (en adelante, ")
    run_mcm = p_explicacion.add_run("la MCM")
    run_mcm.bold = True
    run_mcm.italic = True
    p_explicacion.add_run("). La fórmula es la siguiente:")

    agregar_cuadro_formula(doc)

    # --- 3.2. CRITERIOS ---
    agregar_subtitulo_numerado(doc, "3.2.", "Criterios")
    doc.add_paragraph()
    
    p_criterios_1 = doc.add_paragraph(
        "Mediante la Resolución de Presidencia del Consejo Directivo n.° 00083-2022-OEFA/PCD, "
        "se aprueba el Manual de aplicación de criterios objetivos de la metodología para el cálculo "
        "de las multas base y la aplicación de los factores para la graduación de sanciones en el OEFA."
    )
    p_criterios_1.paragraph_format.left_indent = Cm(1.0)
    
    doc.add_paragraph()
    
    p_criterios_2 = doc.add_paragraph(
        "Asimismo, los conceptos o criterios contenidos en el Manual Explicativo de la Metodología "
        "del Cálculo de Multas del OEFA aprobado por el artículo 3° de la Resolución de Presidencia de Consejo "
        "Directivo del OEFA n.° 035-2013-OEFA/PCD (actualmente derogado), son utilizados en el presente "
        "análisis, de manera referencial, y, en tanto no se opongan a los criterios de graduación de multas "
        "vigentes, aprobados por la Ley n.° 27444, Ley del Procedimiento Administrativo General, en concordancia "
        "con la Resolución de Presidencia de Consejo Directivo del OEFA n.° 035-2013-OEFA/PCD, modificada con "
        "Resolución de Consejo Directivo n.° 024-2017-OEFA/CD."
    )
    p_criterios_2.paragraph_format.left_indent = Cm(1.0)

    # --- 4. DETERMINACIÓN DE LA SANCIÓN ---
    agregar_titulo_numerado(doc, 4, "Determinación de la sanción")
    agregar_subtitulo_numerado(doc, "4.1.", "Consideraciones generales en los cálculos de multa")
    agregar_subtitulo_letra(doc, "A.", "Sobre los costos de mercado en la determinación del beneficio ilícito")

    doc.add_paragraph()
    
    p_a_1 = doc.add_paragraph(
        "Desde un punto de vista económico, ante una multa, el administrado infractor y la ciudadanía en general "
        "deberían estar convencidos de que dicha multa posiciona al infractor en una situación desfavorable frente "
        "a aquellos administrados que cumplieron diligentemente sus obligaciones. Asimismo, lo opuesto ocurriría si "
        "se permitiera que el administrado infractor obtenga un beneficio como resultado del no cumplimiento y de la "
        "información imperfecta existente producto de las asimetrías entre los administrados y la autoridad (problema "
        "del principal-agente), posicionando a los administrados diligentes en una desventaja competitiva y creando "
        "un desincentivo al cumplimiento."
    )
    p_a_1.paragraph_format.left_indent = Cm(1.0)

    doc.add_paragraph()
    
    p_a_2 = doc.add_paragraph(
        "Al respecto, cabe recordar que esta subdirección resuelve el cálculo de multas en un contexto de información "
        "asimétrica y para ello, se aproxima a los costos de mercado, cuyas fuentes y procesos de cálculos satisfacen "
        "un estándar de fundamentación superior al de cualquier otro regulador y se encuentran a disposición del "
        "administrado, observando el debido procedimiento (notificando al administrado los informes de multas, incluyendo "
        "el detalle de los componentes de la metodología correspondiente), dotando de razonabilidad (con el uso de costos "
        "de mercado), celeridad (ejecutando los cálculos de multas expeditivamente), con participación del administrado "
        "(requiriendo comprobantes de pago asociados a realidad y actividad económica); así como la simplicidad (desarrollando "
        "un proceso técnico que permite al administrado conocer de qué forma se arribó a la multa)."
    )
    p_a_2.paragraph_format.left_indent = Cm(1.0)

    doc.add_paragraph()

    p_a_3 = doc.add_paragraph(
        "De otro lado, frente a circunstancias ajenas al genuino espíritu de esta subdirección, como, por ejemplo, la no "
        "apertura de un enlace web o la omisión involuntaria de una captura de pantalla de una fuente; el administrado –o la "
        "autoridad correspondiente– podría corroborar fácilmente, a través de la abundante información web, que el costo "
        "imputado no escapa a los rangos de costos de mercado; lo cual, de ninguna manera, deberían invalidar los cálculos "
        "efectuados."
    )
    p_a_3.paragraph_format.left_indent = Cm(1.0)

    doc.add_paragraph()
    
    p_a_4 = doc.add_paragraph()
    p_a_4.paragraph_format.left_indent = Cm(1.0)
    p_a_4.add_run(
        "Así, en la búsqueda de la disuasión y la maximización del bienestar social, el cual comprende no solo a la empresa "
        "(administrado) sino también a los demás agentes que componen la sociedad, y en línea con lo dispuesto en la "
        "Resolución de Consejo Directivo n.° 001-2024-OEFA/CD"
    )
    
    texto_nota_precedente = (
        "Resolución de consejo directivo n.° 00001-2024-OEFA/CD, publicado el 6 de febrero de 2024:\n"
        "(…) Artículo 1°. - Disponer la publicación de los precedentes administrativos de observancia obligatoria "
        "contenidos en las Resoluciones N.os 543-2023-OEFA/TFA-SE y 551-2023-OEFA/TFA-SE del 21 de noviembre de 2023, "
        "en el diario oficial El Peruano, en el Portal de Transparencia Estándar y en la sede digital del Organismo "
        "de Evaluación y Fiscalización Ambiental-OEFA (https://www.gob.pe/OEFA) en el plazo de dos (2) días hábiles "
        "contado desde su emisión. (…)"
    )
    agregar_nota_al_pie(p_a_4, texto_nota_precedente, doc)
    
    p_a_4.add_run(
        "; que declara precedente administrativo de observancia obligatoria la Resolución n.° "
        "543-2023-OEFA/TFA-SE, para acreditar el costo evitado el administrado podría encontrarse en dos situaciones diferenciadas:"
    )

    p_esc1 = doc.add_paragraph()
    p_esc1.paragraph_format.left_indent = Cm(1.5)
    p_esc1.paragraph_format.first_line_indent = Cm(-0.5)
    p_esc1.paragraph_format.tab_stops.clear_all()
    p_esc1.paragraph_format.tab_stops.add_tab_stop(Cm(1.5))
    p_esc1.paragraph_format.space_before = Pt(14)
    p_esc1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_esc1.add_run("a)\tEscenario 1: Previo a la fecha de cálculo de multa, el administrado ")
    r1_2 = p_esc1.add_run("no ha realizado actividades iguales o semejantes")
    r1_2.bold = True
    r1_2.italic = True
    r1_2.underline = True
    p_esc1.add_run(
        " al costo evitado asociadas a la obligación incumplida. Motivo por el cual resulta pertinente "
        "tomar en cuenta cotizaciones o presupuestos presentados por el administrado para acreditar el costo evitado."
    )

    p_esc2 = doc.add_paragraph()
    p_esc2.paragraph_format.left_indent = Cm(1.5)
    p_esc2.paragraph_format.first_line_indent = Cm(-0.5)
    p_esc2.paragraph_format.tab_stops.clear_all()
    p_esc2.paragraph_format.tab_stops.add_tab_stop(Cm(1.5))
    p_esc2.paragraph_format.space_before = Pt(14)
    p_esc2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_esc2.add_run("b)\tEscenario 2: Previo a la fecha de cálculo de multa, el administrado ")
    r2_2 = p_esc2.add_run("ha realizado actividades iguales o semejantes")
    r2_2.bold = True
    r2_2.italic = True
    r2_2.underline = True
    p_esc2.add_run(
        " al costo evitado asociado a la obligación incumplida. Motivo por el cual resulta razonable "
        "asumir que el administrado cuenta con comprobantes de pago debidamente sustentados y por lo "
        "tanto es pertinente que presente dichos documentos contables a fin de acreditar el costo evitado"
    )
    
    texto_nota_esc2 = (
        "Los comprobantes de pago que se presenten, junto con los documentos vinculados a estos, "
        "deben acreditar que su emisor puede ejecutar las actividades que contemplan y resultan "
        "específicos para el caso concreto."
    )
    agregar_nota_al_pie(p_esc2, texto_nota_esc2, doc)
    p_esc2.add_run(".")

    doc.add_paragraph()

    # =======================================================
    # LÓGICA DINÁMICA DE ESCENARIOS (Periódicos vs No Periódicos)
    # =======================================================
    
    periodicas_ids = ['INF001', 'INF002', 'INF005', 'INF007', 'INF008', 'INF004']
    
    hechos_periodicos = []
    hechos_no_periodicos = []

    # 1. Clasificar todos los hechos (guardamos el número como string)
    for i, datos_hecho in enumerate(imputaciones_data):
        id_infraccion = datos_hecho.get('id_infraccion', '')
        num_hecho_solo = f"{i + 1}"
        
        if any(pid in id_infraccion for pid in periodicas_ids):
            hechos_periodicos.append(num_hecho_solo)
        else:
            hechos_no_periodicos.append(num_hecho_solo)

    # 2. Párrafo para HECHOS PERIÓDICOS (Escenario 2)
    if hechos_periodicos:
        lista_con_tag = [f"n.° {n}" for n in hechos_periodicos]
        hechos_listos = formatear_lista_hechos(
            lista_con_tag, 
            singular_prefix="al hecho imputado", 
            plural_prefix="a los hechos imputados"
        )
        verbo = "corresponden" if len(hechos_periodicos) > 1 else "corresponde"
        
        p_esc_periodicos = doc.add_paragraph()
        p_esc_periodicos.paragraph_format.left_indent = Cm(1.0)
        p_esc_periodicos.add_run(
            f"Sobre ello, respecto {hechos_listos}, de la revisión de los documentos "
            f"obrantes en el presente PAS, se advierte que, el administrado se encontraría en "
            f"un escenario del tipo 2, toda vez que habría realizado actividades iguales o "
            f"semejantes a los costos evitados asociados a las obligaciones incumplidas, dado "
            f"que {verbo} a incumplimientos formales y/u obligaciones periódicas. No "
            f"obstante, hasta la emisión del presente informe, el administrado no ha presentado "
            f"ningún comprobante de pago ni factura para poder ser evaluada."
        )

    # 3. Párrafo para HECHOS NO PERIÓDICOS (Incierto)
    if hechos_no_periodicos:
        if hechos_periodicos:
            doc.add_paragraph() # Enter de separación si hubo párrafo previo
            
        lista_con_tag_np = [f"n.° {n}" for n in hechos_no_periodicos]
        hechos_listos_np = formatear_lista_hechos(
            lista_con_tag_np, 
            singular_prefix="al hecho imputado", 
            plural_prefix="a los hechos imputados"
        )
        
        p_esc_no_periodicos = doc.add_paragraph()
        p_esc_no_periodicos.paragraph_format.left_indent = Cm(1.0)
        p_esc_no_periodicos.add_run(
            f"Además, de la revisión de los documentos obrantes en el presente PAS, en relación "
            f"{hechos_listos_np}, no se tiene información suficiente para determinar en qué "
            f"escenario se encontraría el administrado, toda vez que, hasta la emisión del presente informe, "
            f"no ha presentado ningún comprobante de pago, ni factura ni boletas para poder ser evaluadas."
        )

    doc.add_paragraph()

    p_a_6 = doc.add_paragraph(
        "Finalmente, esta subdirección considera que la introducción de costos no asociados a comprobantes de pago "
        "por parte del administrado, refuerza la información asimétrica, toda vez que este último no revela su "
        "propia información de costos incurridos y, a su vez, redunda en una incorrecta señal de disuasión frente "
        "a los demás administrados, lo que refleja un escenario no razonable de búsqueda de costos más económicos a "
        "favor del administrado infractor, sin que este haya destinado efectivamente un presupuesto para tal fin; "
        "configurándose un posible incentivo perverso en el uso de cotizaciones de menor costo con el fin de reducir "
        "la sanción."
    )
    p_a_6.paragraph_format.left_indent = Cm(1.0)

    agregar_subtitulo_letra(doc, "B.", "Sobre los insumos para el cálculo de multas")
    doc.add_paragraph()

    p_b_1 = doc.add_paragraph(
        "Para la elaboración del presente informe, se considera el MAPRO PM5, en lo referido a las solicitudes "
        "de multa, aprobado mediante Resolución de Presidencia de Consejo Directivo n.° 00061-2022-OEFA/PCD del "
        "4 de noviembre de 2022."
    )
    p_b_1.paragraph_format.left_indent = Cm(1.0)

    doc.add_paragraph()

    p_b_2 = doc.add_paragraph()
    p_b_2.paragraph_format.left_indent = Cm(1.0)
    p_b_2.add_run(
        "Asimismo, las estimaciones económicas asociadas al expediente bajo análisis se encuentran motivadas a partir "
        "de los insumos provistos por parte los equipos técnicos, en lo referido a las actividades asociadas al costo "
        "evitado y a los factores de graduación f1, f3, f5, f6; y legales, en lo referido a los factores de graduación "
        "f4 y f7; quienes, a partir de los medios probatorios que obran en el presente expediente y la "
    )
    run_expertise = p_b_2.add_run("expertise")
    run_expertise.italic = True
    
    # 1. Lógica gramatical rápida para singular/plural
    if cantidad_hechos == 1:
        texto_conductas = "la conducta infractora"
        texto_la_infraccion = "la infracción"
    else:
        texto_conductas = "las conductas infractoras"
        texto_la_infraccion = "las infracciones"

    # 2. Inyectamos la variable dinámica 'texto_conductas'
    p_b_2.add_run(
        " profesional correspondiente, considerando las asimetrías de información, efectúan una aproximación de los "
        f"aspectos mínimos indispensables requeridos para el cálculo de la sanción de {texto_conductas} bajo análisis."
    )

    doc.add_paragraph()

    # 3. Cambiamos el texto base y le inyectamos la variable 'texto_la_infraccion'
    p_b_3 = doc.add_paragraph(
        f"Bajo las consideraciones antes mencionadas, se procede a la estimación de la multa para {texto_la_infraccion} bajo análisis."
    )
    p_b_3.paragraph_format.left_indent = Cm(1.0)

# =======================================================
    # --- BUCLE DE ANÁLISIS DE HECHOS DINÁMICOS ---
    # =======================================================
    for i, datos_hecho in enumerate(imputaciones_data):
        hecho_texto = datos_hecho.get('texto_hecho', '')
        id_infraccion = datos_hecho.get('id_infraccion', '')
        resultados = datos_hecho.get('resultados', {})
        
        num_subseccion = f"4.{i + 2}."  
        num_hecho = i + 1               
        
        # Imprimimos el título del hecho
        p_hecho = doc.add_paragraph()
        pf = p_hecho.paragraph_format
        pf.left_indent = Cm(1.0)
        pf.first_line_indent = Cm(-1.0)
        pf.tab_stops.clear_all()
        pf.tab_stops.add_tab_stop(Cm(1.0))
        pf.space_before = Pt(14)
        
        run_num = p_hecho.add_run(f"{num_subseccion}\tHecho imputado n.° {num_hecho}: ")
        run_num.font.name = 'Arial'
        run_num.font.size = Pt(11)
        run_num.bold = True
        
        run_texto = p_hecho.add_run(hecho_texto)
        run_texto.font.name = 'Arial'
        run_texto.font.size = Pt(11)
        
        doc.add_paragraph()

        # ---------------------------------------------------
        # i) BENEFICIO ILÍCITO (B) 
        # ---------------------------------------------------
        p_bi = doc.add_paragraph()
        p_bi.paragraph_format.left_indent = Cm(1.5)
        p_bi.paragraph_format.first_line_indent = Cm(-0.5) 
        p_bi.paragraph_format.tab_stops.clear_all()
        p_bi.paragraph_format.tab_stops.add_tab_stop(Cm(1.5))
        
        run_bi = p_bi.add_run("i)\tBeneficio Ilícito (B)")
        run_bi.font.name = 'Arial'
        run_bi.font.size = Pt(11)
        run_bi.bold = True

        doc.add_paragraph()

        import importlib
        try:
            modulo_redactor = importlib.import_module(f"textos_infracciones.{id_infraccion}")
        except ImportError:
            modulo_redactor = importlib.import_module("textos_infracciones.generico")

        # LLAMADA DINÁMICA: Dibuja Beneficio Ilícito
        if hasattr(modulo_redactor, 'redactar_beneficio_ilicito'):
            modulo_redactor.redactar_beneficio_ilicito(doc, datos_hecho, num_hecho)
        else:
            p_marcador = doc.add_paragraph(f"[Falta programar Beneficio Ilícito para {id_infraccion}]")
            p_marcador.runs[0].font.color.rgb = RGBColor(255, 0, 0)

        # ---------------------------------------------------
        # ii) PROBABILIDAD DE DETECCIÓN (p)
        # ---------------------------------------------------
        # LLAMADA DINÁMICA: Dibuja Probabilidad de Detección
        if hasattr(modulo_redactor, 'redactar_probabilidad'):
            modulo_redactor.redactar_probabilidad(doc, datos_hecho, num_hecho)
        else:
            p_marcador = doc.add_paragraph(f"[Falta programar Probabilidad para {id_infraccion}]")
            p_marcador.runs[0].font.color.rgb = RGBColor(255, 0, 0)

        # ---------------------------------------------------
        # iii) FACTORES PARA LA GRADUACIÓN DE SANCIONES (F) - (ESTÁNDAR DE IFI.PY)
        # ---------------------------------------------------
        p_factores = doc.add_paragraph()
        p_factores.paragraph_format.left_indent = Cm(1.5)
        p_factores.paragraph_format.first_line_indent = Cm(-0.5)
        p_factores.paragraph_format.tab_stops.clear_all()
        p_factores.paragraph_format.tab_stops.add_tab_stop(Cm(1.5))
        
        run_factores = p_factores.add_run("iii)\tFactores para la Graduación de Sanciones (F)")
        run_factores.font.name = 'Arial'
        run_factores.font.size = Pt(11)
        run_factores.bold = True

        doc.add_paragraph()

        # --- LÓGICA CONDICIONAL DE GRADUACIÓN ---
        aplica_graduacion = datos_hecho.get('aplica_graduacion', 'No')
        
        if aplica_graduacion == 'Sí':
            datos_graduacion = datos_hecho.get('graduacion', {})
            
            # Analizamos qué factores están activos (distintos de 0)
            factores_activos = []
            factores_inactivos = []
            for num in range(1, 8):
                if datos_graduacion.get(f'subtotal_f{num}', 0.0) != 0:
                    factores_activos.append(f"f{num}")
                else:
                    factores_inactivos.append(f"f{num}")
            
            from funciones import texto_con_numero
            
            # Formateo de los activos (ej: "dos (2)", "f1 y f2")
            cantidad_activos_txt = texto_con_numero(len(factores_activos), genero='m')
            
            if len(factores_activos) == 1:
                lista_activos_txt = factores_activos[0]
            elif len(factores_activos) > 1:
                lista_activos_txt = ", ".join(factores_activos[:-1]) + " y " + factores_activos[-1]
            else:
                lista_activos_txt = "ninguno"
            
            # Formateo de los inactivos
            if len(factores_inactivos) == 1:
                lista_inactivos_txt = f"el factor {factores_inactivos[0]} no tiene"
            elif len(factores_inactivos) > 1:
                lista_inactivos_txt = "los factores " + ", ".join(factores_inactivos[:-1]) + " y " + factores_inactivos[-1] + " no tienen"
            else:
                lista_inactivos_txt = "ningún factor tiene"

            # 1. Párrafo Introductorio
            p_desc_f = doc.add_paragraph(
                f"La determinación de los factores para la graduación de sanciones sigue lo establecido "
                f"en el MCM del OEFA; por ello, de acuerdo con la información disponible en el presente "
                f"expediente y el análisis del equipo técnico, se ha estimado pertinente aplicar {cantidad_activos_txt} "
                f"de los siete (7) factores para la graduación de sanciones: {lista_activos_txt}. El detalle y la motivación es el siguiente:"
            )
            p_desc_f.paragraph_format.left_indent = Cm(1.0)
            p_desc_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph()
            
            # 2. Marcador rojo para edición manual
            p_rojo = doc.add_paragraph()
            p_rojo.paragraph_format.left_indent = Cm(1.0)
            run_rojo = p_rojo.add_run("[Ingresar los detalles de los factores de graduación]")
            run_rojo.font.color.rgb = RGBColor(255, 0, 0)
            run_rojo.bold = True
            doc.add_paragraph()
            
            # 3. Subtítulo: Otros factores
            p_otros = doc.add_paragraph()
            p_otros.paragraph_format.left_indent = Cm(1.0)
            run_otros = p_otros.add_run("Otros factores")
            run_otros.bold = True
            run_otros.underline = True
            doc.add_paragraph()
            
            p_inactivos = doc.add_paragraph(
                f"De otro lado, de la revisión del expediente se advierte que, con la información disponible "
                f"{lista_inactivos_txt} una calificación de 0%."
            )
            p_inactivos.paragraph_format.left_indent = Cm(1.0)
            p_inactivos.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph()
            
            # 4. Subtítulo: Total de factores
            p_totalf = doc.add_paragraph()
            p_totalf.paragraph_format.left_indent = Cm(1.0)
            run_totalf = p_totalf.add_run("Total de factores")
            run_totalf.bold = True
            run_totalf.underline = True
            doc.add_paragraph()
            
            factor_f_calculado = datos_graduacion.get('factor_f_calculado', 1.0)
            texto_f_final = f"{factor_f_calculado:.2f} ({factor_f_calculado:.0%})"
            
            p_resumenf = doc.add_paragraph()
            p_resumenf.paragraph_format.left_indent = Cm(1.0)
            p_resumenf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_resumenf.add_run(f"En total, los factores para la graduación de sanciones resultan en {texto_f_final}")
            
            # Insertamos nota al pie
            agregar_nota_al_pie(p_resumenf, "Para mayor detalle, ver Anexo n.° 2.", doc)
            
            p_resumenf.add_run(". El detalle es el siguiente:")
            doc.add_paragraph()
            
            # 5. Dibujar la tabla
            dibujar_tabla_factores_graduacion(doc, datos_graduacion)
            doc.add_paragraph()

        else:
            # Texto estándar si NO aplican factores
            p_desc_factores = doc.add_paragraph()
            p_desc_factores.paragraph_format.left_indent = Cm(1.0)
            p_desc_factores.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            p_desc_factores.add_run(
                "Para el presente caso, dada la naturaleza de la infracción bajo análisis y para efectos del cálculo "
                "de la multa, no se identifica la existencia de factores para la graduación de la sanción. En tal sentido, "
                "en la fórmula de la multa se consigna una calificación de 1.0 (100%)."
            )
            doc.add_paragraph()

        # ---------------------------------------------------
        # iv) MULTA PROPUESTA - (ESTÁNDAR DE IFI.PY)
        # ---------------------------------------------------
        p_multa = doc.add_paragraph()
        p_multa.paragraph_format.left_indent = Cm(1.5)
        p_multa.paragraph_format.first_line_indent = Cm(-0.5)
        p_multa.paragraph_format.tab_stops.clear_all()
        p_multa.paragraph_format.tab_stops.add_tab_stop(Cm(1.5))
        
        run_multa = p_multa.add_run("iv)\tMulta Propuesta")
        run_multa.font.name = 'Arial'
        run_multa.font.size = Pt(11)
        run_multa.bold = True

        doc.add_paragraph()

        # Extraemos la multa final para redactar el párrafo
        totales_app = resultados.get('resultados_para_app', {}).get('totales', {})
        if not totales_app:
            totales_app = resultados.get('resultados_para_app', {})
            
        contexto_word = resultados.get('contexto_final_word', {})
        
        # --- LA CORRECCIÓN: Búsqueda exhaustiva de la multa ---
        monto_uit_crudo = contexto_word.get('multa_original_uit')
        if not monto_uit_crudo:
            monto_uit_crudo = totales_app.get('multa_final_uit', totales_app.get('multa_base_uit', 0.0))
        
        # Limpiamos el valor por si viene como texto (ej: "0.250 UIT")
        import re
        try:
            monto_uit_val = float(re.search(r'\d+(\.\d+)?', str(monto_uit_crudo).replace(',', '')).group())
        except Exception:
            monto_uit_val = 0.0
            
        monto_uit_str = f"{monto_uit_val:,.3f} UIT"
        
        # Párrafo introductorio con la variable en negrita
        p_desc_multa = doc.add_paragraph()
        p_desc_multa.paragraph_format.left_indent = Cm(1.0)
        p_desc_multa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        p_desc_multa.add_run("Luego de aplicar la fórmula para el cálculo de la multa, se identificó que la misma asciende a ")
        
        run_monto_bold = p_desc_multa.add_run(monto_uit_str)
        run_monto_bold.bold = True
        
        p_desc_multa.add_run(". El resumen de la multa y sus componentes se presenta en el siguiente cuadro.")
        
        doc.add_paragraph()

        # Llamamos a la función constructora de la tabla
        dibujar_tabla_multa_propuesta(doc, datos_hecho)
        
        doc.add_paragraph()

        # ---------------------------------------------------
        # NUEVA SECCIÓN: REDUCCIÓN DE LA MULTA (CONDICIONAL)
        # ---------------------------------------------------
        aplica_reduccion = datos_hecho.get('aplica_reduccion', 'No')
        numeracion_principios = "v)" # Por defecto, si no hay reducción, Principios es v)
        monto_final_para_principios = monto_uit_str # Por defecto, la multa no tiene descuento
        
        if aplica_reduccion == 'Sí':
            numeracion_principios = "vi)" # Empujamos Principios al numeral vi)
            
            p_red = doc.add_paragraph()
            p_red.paragraph_format.left_indent = Cm(1.5)
            p_red.paragraph_format.first_line_indent = Cm(-0.5)
            p_red.paragraph_format.tab_stops.clear_all()
            p_red.paragraph_format.tab_stops.add_tab_stop(Cm(1.5))
            
            run_red = p_red.add_run("v)\tReducción de la multa por Reconocimiento de Responsabilidad, en aplicación del RPAS")
            run_red.font.name = 'Arial'
            run_red.font.size = Pt(11)
            run_red.bold = True
            
            doc.add_paragraph()
            
            from babel.dates import format_date
            
            # 1. Extracción y formateo de variables
            memo_num = datos_hecho.get('memo_num', '[N.° MEMO]')
            memo_fecha_obj = datos_hecho.get('memo_fecha')
            memo_fecha = format_date(memo_fecha_obj, "d 'de' MMMM 'de' yyyy", locale='es') if memo_fecha_obj else '[FECHA MEMO]'
            
            escrito_num = datos_hecho.get('escrito_num', '[N.° ESCRITO]')
            escrito_fecha_obj = datos_hecho.get('escrito_fecha')
            escrito_fecha = format_date(escrito_fecha_obj, "d 'de' MMMM 'de' yyyy", locale='es') if escrito_fecha_obj else '[FECHA ESCRITO]'
            
            id_subdireccion = context_data.get('id_subdireccion', '[SUBDIRECCIÓN]')
            texto_reduccion = datos_hecho.get('texto_reduccion', '[TEXTO REDUCCIÓN]')
            porcentaje_reduccion = datos_hecho.get('porcentaje_reduccion', '[%]')
            
            # Extraemos la multa final reducida (para imprimirla aquí y mandarla a Principios)
            multa_con_reduccion = totales_app.get('multa_con_reduccion_uit', 0.0)
            multa_red_str = f"{multa_con_reduccion:,.3f} UIT"
            monto_final_para_principios = multa_red_str # Actualizamos la multa final para el siguiente párrafo
            
            # 2. Párrafo Introductorio
            p_red1 = doc.add_paragraph()
            p_red1.paragraph_format.left_indent = Cm(1.0)
            p_red1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_red1.add_run(
                f"De acuerdo con el Memorando n.° {memo_num}, del {memo_fecha}, la {id_subdireccion} "
                f"informó a esta subdirección que mediante escrito con registro n.° {escrito_num}, "
                f"del {escrito_fecha}, el administrado reconoció su responsabilidad administrativa por la "
                f"comisión de la infracción bajo análisis. Asimismo, indicó que dicho reconocimiento fue realizado "
                f"{texto_reduccion}."
            )
            
            doc.add_paragraph()
            
            # 3. Párrafo de Aplicación con Memoria de RPAS
            p_red2 = doc.add_paragraph()
            p_red2.paragraph_format.left_indent = Cm(1.0)
            p_red2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            p_red2.add_run("En tal sentido, de acuerdo con el artículo 13° del ")
            
            if not hasattr(doc, 'acronimos_usados'): doc.acronimos_usados = set()
                
            if 'RPAS' not in doc.acronimos_usados:
                p_red2.add_run("Reglamento de Procedimiento Administrativo Sancionador (en adelante, ")
                run_rpas = p_red2.add_run("RPAS")
                run_rpas.bold = True
                run_rpas.italic = True
                p_red2.add_run(")")
                doc.acronimos_usados.add('RPAS')
            else:
                run_rpas = p_red2.add_run("RPAS")
                run_rpas.bold = True
                run_rpas.italic = True
                
            # Construcción de la nota al pie con tabulaciones integradas
            texto_nota_rpas = (
                "Reglamento del Procedimiento Administrativo Sancionador del OEFA, aprobado por Resolución de Consejo Directivo n.º 027-2017-OEFA/CD: Artículo 13°. - Reducción de la multa por reconocimiento de responsabilidad (…)\n"
                "13.3 El porcentaje de reducción de la multa se otorgará de acuerdo a un criterio de oportunidad en la formulación del reconocimiento de responsabilidad, según el siguiente cuadro:\n"
                "50%\tDesde el inicio del procedimiento administrativo sancionador hasta la presentación de los descargos a la imputación de cargos.\n"
                "30%\tLuego de presentados los descargos a la imputación de cargos hasta antes de la emisión de la Resolución Final."
            )
            agregar_nota_al_pie(p_red2, texto_nota_rpas, doc)
            
            p_red2.add_run(
                f", corresponde la aplicación del descuento del {porcentaje_reduccion} de la multa calculada respecto "
                f"al hecho imputado. Por lo tanto, la multa pasa de {monto_uit_str} a {multa_red_str} por dicho reconocimiento."
            )
            
            doc.add_paragraph()
            
        # Inyectamos en memoria el número de sección que le tocará a Principios
        datos_hecho['numeracion_principios'] = numeracion_principios

        # ---------------------------------------------------
        # v o vi) APLICACIÓN DE PRINCIPIOS
        # ---------------------------------------------------
        # LLAMADA DINÁMICA: Dibuja Principios con la multa final (Original o con Descuento)
        if hasattr(modulo_redactor, 'redactar_principios'):
            modulo_redactor.redactar_principios(doc, datos_hecho, num_hecho, monto_final_para_principios)
        else:
            p_marcador = doc.add_paragraph(f"[Falta programar Principios para {id_infraccion}]")
            p_marcador.runs[0].font.color.rgb = RGBColor(255, 0, 0)

    # --- 5. ANÁLISIS DE NO CONFISCATORIEDAD ---
    agregar_titulo_numerado(doc, 5, "Análisis de no confiscatoriedad")
    doc.add_paragraph()

    # 1. Extraemos las variables pluralizadas del contexto
    plural_infraccion_analizada = context_data.get('plural_infraccion_analizada', 'la infracción analizada')
    plural_cada_infraccion = context_data.get('plural_cada_infraccion', 'la infracción')

    # 2. Buscamos el ID del sector (Rubro) desde los resultados guardados
    id_rubro = ""
    if imputaciones_data and 'resultados' in imputaciones_data[0]:
        resultados_app = imputaciones_data[0]['resultados'].get('resultados_para_app', {})
        id_rubro = resultados_app.get('id_rubro', '')

    p_conf = doc.add_paragraph()
    p_conf.paragraph_format.left_indent = Cm(1.0)
    p_conf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_conf.add_run("De acuerdo con lo establecido en el numeral 12.2 del artículo 12° del ")

    # 3. Lógica de memoria para RPAS
    if not hasattr(doc, 'acronimos_usados'): doc.acronimos_usados = set()
    
    if 'RPAS' not in doc.acronimos_usados:
        p_conf.add_run("Reglamento de Procedimiento Administrativo Sancionador (en adelante, ")
        run_rpas_conf = p_conf.add_run("RPAS")
        run_rpas_conf.bold = True
        run_rpas_conf.italic = True
        p_conf.add_run(")")
        doc.acronimos_usados.add('RPAS')
    else:
        run_rpas_conf = p_conf.add_run("RPAS")
        run_rpas_conf.bold = True
        run_rpas_conf.italic = True

    # 4. Nota al pie 1 (Base legal)
    texto_nota_rpas_1 = (
        "Reglamento del Procedimiento Administrativo Sancionador del OEFA, aprobado por Resolución de Presidencia de Consejo Directivo n.° 027-2017-OEFA/CD.\n"
        "“Artículo 12°. - Determinación de las multas (…)\n"
        "12.2 La multa a ser impuesta no puede ser mayor al diez por ciento (10%) del ingreso bruto anual percibido por el infractor el año anterior a la fecha en que ha cometido la infracción.”"
    )
    agregar_nota_al_pie(p_conf, texto_nota_rpas_1, doc)

    p_conf.add_run(
        f", la multa a imponerse por {plural_infraccion_analizada}, no puede ser mayor al diez por ciento (10 %) "
        f"del ingreso bruto anual percibido por el infractor en el año anterior a la fecha en que ha cometido "
        f"{plural_cada_infraccion}. Asimismo, los ingresos deberán ser debidamente acreditados por el administrado"
    )

    # 5. Nota al pie 2 (Condicional: Se oculta si el sector es R019)
    if id_rubro != "R019":
        texto_nota_trib = (
            "Cabe señalar que de acuerdo con el literal b) del artículo 180° del Código Tributario para el caso "
            "de los contribuyentes que se encuentren en el Régimen General, se considerará como ingreso a la "
            "información contenida en los campos o casillas de la Declaración Jurada Anual en las que se consignen "
            "los conceptos de Ventas Netas y/o Ingresos por Servicios y otros ingresos gravables y no gravables "
            "de acuerdo con la Ley del Impuesto a la Renta."
        )
        agregar_nota_al_pie(p_conf, texto_nota_trib, doc)

    p_conf.add_run(".")
    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # 6. Párrafo de solicitud de información y respuesta del administrado
    # -------------------------------------------------------------------------
    # A) Calculamos los años de los ingresos (año de infracción - 1) dinámicamente
    anios_ingresos = set()
    for datos_hecho in imputaciones_data:
        extremos_del_hecho = datos_hecho.get('extremos', [])
        # Fallback si no hay 'extremos' pero sí fechas sueltas
        if not extremos_del_hecho:
            fecha_inc = datos_hecho.get('fecha_incumplimiento') or datos_hecho.get('fecha_incumplimiento_extremo')
            if fecha_inc:
                try: anios_ingresos.add(fecha_inc.year - 1)
                except AttributeError: pass
        else:
            for extremo in extremos_del_hecho:
                fecha_inc = extremo.get('fecha_incumplimiento') or extremo.get('fecha_incumplimiento_extremo')
                if fecha_inc:
                    try: anios_ingresos.add(fecha_inc.year - 1)
                    except AttributeError: pass

    anios_ordenados = sorted(list(anios_ingresos))
    if len(anios_ordenados) == 1:
        texto_anios = str(anios_ordenados[0])
    elif len(anios_ordenados) > 1:
        texto_anios = ", ".join(str(y) for y in anios_ordenados[:-1]) + " y " + str(anios_ordenados[-1])
    else:
        texto_anios = "[AÑOS]"

    # B) Obtenemos unidad fiscalizable
    unidad_fiscalizable = "[UNIDAD FISCALIZABLE]"
    if imputaciones_data and 'resultados' in imputaciones_data[0]:
        resultados_app = imputaciones_data[0]['resultados'].get('resultados_para_app', {})
        unidad_fiscalizable = resultados_app.get('administrado', '[UNIDAD FISCALIZABLE]')

    p_conf2 = doc.add_paragraph()
    p_conf2.paragraph_format.left_indent = Cm(1.0)
    p_conf2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    id_subdireccion_conf = context_data.get('id_subdireccion', '[SUBDIRECCIÓN]')
    administrado_nombre = context_data.get('administrado', '')

    # C) Primera oración: Solicitud de la Autoridad
    p_conf2.add_run(
        f"Para tal efecto, mediante la Resolución Subdirectoral, la {id_subdireccion_conf} del OEFA, "
        f"solicitó al administrado la remisión de sus ingresos brutos correspondientes al {texto_anios}, "
        f"a fin de proceder con el análisis de no confiscatoriedad."
    )

    # D) Condicional Especial para Municipalidades
    if 'municipalidad' in administrado_nombre.lower():
        p_conf2.add_run(
            f" Es importante precisar que, dicha información debió estar debidamente acreditada y vinculada "
            f"a la actividad que el administrado desarrolló en la unidad fiscalizable, es decir, en el "
            f"\"{unidad_fiscalizable}\", según lo señalado en el Informe de Supervisión del presente expediente."
        )

    # E) Condicional de Respuesta del Administrado (Aplica / No Aplica)
    aplica_conf = conf_data.get('aplica', 'No')
    
    if aplica_conf == 'Sí':
        from babel.dates import format_date
        escrito_num = conf_data.get('escrito_num_conf', '[N.° ESCRITO]')
        escrito_fecha_obj = conf_data.get('escrito_fecha_conf')
        escrito_fecha = format_date(escrito_fecha_obj, "d 'de' MMMM 'de' yyyy", locale='es') if escrito_fecha_obj else '[FECHA ESCRITO]'
        
        p_conf2.add_run(
            f" En atención a ello, el administrado presentó la información requerida mediante el escrito "
            f"de descargo n.° {escrito_num}, del {escrito_fecha}, en el cual se incluye el detalle de los "
            f"ingresos brutos solicitados."
        )
    else:
        p_conf2.add_run(
            " Sin embargo, el administrado no atendió el requerimiento de información. Por lo tanto, "
            "no se ha podido realizar el análisis de no confiscatoriedad."
        )

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # 7. Evaluación matemática de Confiscatoriedad (Múltiples Años)
    # -------------------------------------------------------------------------
    if aplica_conf == 'Sí':
        import re
        import streamlit as st # Importamos st localmente para acceder a la base de datos de UIT
        
        # Obtenemos la tabla de UIT directamente de la memoria de la aplicación
        df_uit = st.session_state.datos_calculo.get('df_uit')
        datos_por_anio = conf_data.get('datos_por_anio', {})
        escrito_num_global = conf_data.get('escrito_num_conf', '[N.° ESCRITO]')
        
        # A) Agrupar multas totales por año de incumplimiento
        multas_por_anio = {}
        for datos_hecho in imputaciones_data:
            # Extraemos la multa final del hecho (reducida o base)
            res = datos_hecho.get('resultados', {})
            tot_app = res.get('resultados_para_app', {}).get('totales', {})
            if not tot_app: tot_app = res.get('resultados_para_app', {})
            ctx_word = res.get('contexto_final_word', {})
            
            if datos_hecho.get('aplica_reduccion', 'No') == 'Sí':
                val_c = tot_app.get('multa_con_reduccion_uit', 0.0)
            else:
                val_c = ctx_word.get('multa_original_uit') or tot_app.get('multa_final_uit', tot_app.get('multa_base_uit', 0.0))
                
            try: m_val = float(re.search(r'\d+(\.\d+)?', str(val_c).replace(',', '')).group())
            except Exception: m_val = 0.0
            
            # Buscar el año principal de este hecho
            anio_asig = None
            exts = datos_hecho.get('extremos', [])
            if exts:
                fi = exts[0].get('fecha_incumplimiento') or exts[0].get('fecha_incumplimiento_extremo')
                if fi: anio_asig = fi.year
            else:
                fi = datos_hecho.get('fecha_incumplimiento') or datos_hecho.get('fecha_incumplimiento_extremo')
                if fi: anio_asig = fi.year
                
            if anio_asig:
                multas_por_anio[anio_asig] = multas_por_anio.get(anio_asig, 0.0) + m_val
                
        # B) Generar un párrafo matemático por cada año de ingreso analizado
        anios_ordenados_conf = sorted(list(datos_por_anio.keys()))
        
        for anio_inc in anios_ordenados_conf:
            datos_ingresos = datos_por_anio[anio_inc]
            anio_ing = anio_inc - 1
            anio_ref_uit = int(datos_ingresos.get('anio_ingresos', anio_ing))
            ingreso_bruto_soles = float(datos_ingresos.get('ingreso_total_soles', 0.0))
            
            # Extraer valor UIT histórico de la base de datos
            uit_val = 4600.0 # Valor por defecto si falla
            if df_uit is not None and not df_uit.empty:
                fila_u = df_uit[df_uit['Anio'] == anio_ref_uit]
                if not fila_u.empty:
                    uit_val = float(fila_u.iloc[0]['Valor_UIT'])
                    
            # Cálculos matemáticos
            ingreso_bruto_uit = ingreso_bruto_soles / uit_val if uit_val > 0 else 0.0
            tope_10_porciento = ingreso_bruto_uit * 0.10
            multa_total_del_anio = multas_por_anio.get(anio_inc, 0.0)
            es_confiscatoria = multa_total_del_anio > tope_10_porciento
            
            # Formateos (Con comas de miles y decimales)
            ing_uit_str = f"{ingreso_bruto_uit:,.3f}"
            tope_str = f"{tope_10_porciento:,.3f}"
            multa_anio_str = f"{multa_total_del_anio:,.3f}"
            ing_soles_str = f"{ingreso_bruto_soles:,.2f}"
            uit_val_str = f"{uit_val:,.2f}"
            
            # Dibujar Párrafo
            p_item = doc.add_paragraph()
            p_item.paragraph_format.left_indent = Cm(1.0)
            p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            p_item.add_run(f"En el {anio_ing}, los ingresos brutos obtenidos por el administrado ascendieron a {ing_uit_str} UIT")
            
            # Insertar nota al pie detallada (Excluyendo al sector R019)
            if id_rubro != "R019":
                texto_nota_item = (
                    f"Los ingresos totales del {anio_ing} expresados en UIT, se obtuvieron dividiendo los ingresos consignados "
                    f"en la cuenta Netas y/o Ingresos por Servicios y otros ingresos gravables y no gravables del {anio_ing}, "
                    f"los cuales ascendieron a S/ {ing_soles_str} entre la UIT del {anio_ref_uit}, equivalente a S/ {uit_val_str}, "
                    f"lo cual da como resultado {ing_uit_str}.\n"
                    f"Fuente:\n"
                    f"-\tEstado de Resultados del 01 de enero al 31 de diciembre del {anio_ing}, a través del escrito con registro n.° {escrito_num_global}.\n"
                    f"-\tUIT: SUNAT - Índices y tasas. http://www.sunat.gob.pe/indicestasas/uit.html"
                )
                agregar_nota_al_pie(p_item, texto_nota_item, doc)
                
            p_item.add_run(f". En consecuencia, considerando que la multa total estimada para el incumplimiento del {anio_inc} asciende a {multa_anio_str} UIT, se concluye que esta ")
            
            if es_confiscatoria:
                p_item.add_run(f"resulta confiscatoria, por lo que se topa al diez por ciento (10 %) de los ingresos brutos anuales percibidos en {anio_ing}, equivalente a {tope_str} UIT.")
            else:
                p_item.add_run(f"no resulta confiscatoria, al encontrarse por debajo del diez por ciento (10 %) de los ingresos brutos anuales percibidos en {anio_ing}, equivalente a {tope_str} UIT.")

    # --- 6. CONCLUSIONES ---
    agregar_titulo_numerado(doc, 6, "Conclusiones")
    doc.add_paragraph()

    # --- CÁLCULO DINÁMICO DE LA MULTA TOTAL ---
    import re
    multa_total_uit_val = 0.0
    
    # Verificamos si hubo alguna reducción global para la redacción
    aplica_reduccion_global = False
    porcentaje_red = "50%" # Default, pero lo buscaremos
    
    for datos_hecho in imputaciones_data:
        resultados = datos_hecho.get('resultados', {})
        totales_app = resultados.get('resultados_para_app', {}).get('totales', {})
        if not totales_app: totales_app = resultados.get('resultados_para_app', {})
        contexto_word = resultados.get('contexto_final_word', {})

        if datos_hecho.get('aplica_reduccion', 'No') == 'Sí':
            aplica_reduccion_global = True
            porcentaje_red = datos_hecho.get('porcentaje_reduccion', '50%')
            val_crudo = totales_app.get('multa_con_reduccion_uit', 0.0)
        else:
            val_crudo = contexto_word.get('multa_original_uit')
            if not val_crudo:
                val_crudo = totales_app.get('multa_final_uit', totales_app.get('multa_base_uit', 0.0))

        try:
            val_num = float(re.search(r'\d+(\.\d+)?', str(val_crudo).replace(',', '')).group())
        except Exception:
            val_num = 0.0

        multa_total_uit_val += val_num

    multa_total_uit = f"{multa_total_uit_val:,.3f}"
    
    # Variables de entorno
    aplica_confiscatoriedad = conf_data.get('aplica', 'No') == 'Sí'
    cantidad_hechos = len(imputaciones_data)
    texto_incumplimiento = "el presunto incumplimiento" if cantidad_hechos == 1 else "los presuntos incumplimientos"

    # --- PÁRRAFO CONDICIONAL DE CONCLUSIONES ---
    p_conclusiones = doc.add_paragraph()
    p_conclusiones.paragraph_format.left_indent = Cm(1.0)
    p_conclusiones.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_conclusiones.add_run("En base al principio de razonabilidad que rige la potestad sancionadora del OEFA, luego de aplicar la metodología para el cálculo de multas y sus criterios objetivos;")

    if aplica_reduccion_global:
        p_conclusiones.add_run(f" el descuento del {porcentaje_red} de la multa por reconocimiento de responsabilidad, en aplicación del RPAS")

    # Lógica de conjunción ("y" vs ";")
    if not aplica_reduccion_global and not aplica_confiscatoriedad:
        p_conclusiones.add_run(" y")
    else:
        p_conclusiones.add_run(";")

    p_conclusiones.add_run(" el análisis de tope de multas por tipificación de infracciones")

    if aplica_confiscatoriedad:
        p_conclusiones.add_run(" y el análisis de no confiscatoriedad")

    p_conclusiones.add_run("; se propone una sanción de ")
    
    # Multa en negrita
    run_mt = p_conclusiones.add_run(f"{multa_total_uit} UIT")
    run_mt.bold = True
    
    p_conclusiones.add_run(f" para {texto_incumplimiento} materia de análisis, de acuerdo con el siguiente detalle:")
    
    doc.add_paragraph()

    # --- DIBUJO DE LA TABLA RESUMEN ---
    dibujar_tabla_resumen_multas(doc, imputaciones_data, multa_total_uit)
    
    # --- 7. RECOMENDACIONES ---
    agregar_titulo_numerado(doc, 7, "Recomendaciones")
    doc.add_paragraph()

    id_subdireccion = context_data.get('id_subdireccion', '[SUBDIRECCIÓN]')
    p_recom_intro = doc.add_paragraph(f"Se recomienda que la {id_subdireccion} del OEFA solicite al administrado:")
    p_recom_intro.paragraph_format.left_indent = Cm(1.0)
    doc.add_paragraph()
    
    # Extraemos variables necesarias para las condiciones
    administrado_nombre = context_data.get('administrado', '')
    aplica_confiscatoriedad = conf_data.get('aplica', 'No') == 'Sí'
    ph_multa_propuesta = context_data.get('ph_multa_propuesta', 'la multa propuesta resulta')
    ph_la_infraccion = context_data.get('ph_la_infraccion', 'la infracción')
    
    # Viñeta 1 (Condicional: Solo si NO hay confiscatoriedad)
    if not aplica_confiscatoriedad:
        p_v1 = doc.add_paragraph()
        p_v1.paragraph_format.left_indent = Cm(1.5)
        p_v1.paragraph_format.first_line_indent = Cm(-0.5)
        p_v1.paragraph_format.tab_stops.clear_all()
        p_v1.paragraph_format.tab_stops.add_tab_stop(Cm(1.5))
        
        if 'municipalidad' in administrado_nombre.lower():
            texto_v1 = (f"Información sobre sus ingresos brutos correspondientes al {texto_anios}, "
                        f"debidamente acreditados y vinculados directamente a la actividad que desarrolla "
                        f"en la unidad fiscalizable, a fin de verificar si {ph_multa_propuesta} no confiscatoria.")
        else:
            texto_v1 = (f"Sus ingresos brutos correspondientes al {texto_anios}, presentados ante la SUNAT; "
                        f"a fin de verificar si {ph_multa_propuesta} no confiscatoria.")
        
        run_v1 = p_v1.add_run(f"●\t{texto_v1}")
        run_v1.font.name = 'Arial'
        run_v1.font.size = Pt(11)
        
        doc.add_paragraph()

    # Viñeta 2 (Siempre visible)
    p_v2 = doc.add_paragraph()
    p_v2.paragraph_format.left_indent = Cm(1.5)
    p_v2.paragraph_format.first_line_indent = Cm(-0.5)
    p_v2.paragraph_format.tab_stops.clear_all()
    p_v2.paragraph_format.tab_stops.add_tab_stop(Cm(1.5))
    
    texto_v2 = (f"Recibos por honorarios, facturas o boletas que se encuentren directamente asociados a "
                f"{ph_la_infraccion} bajo análisis, para que puedan ser evaluados y, de corresponder, "
                f"considerados en el cálculo de la sanción.")
    run_v2 = p_v2.add_run(f"●\t{texto_v2}")
    run_v2.font.name = 'Arial'
    run_v2.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # --- FIRMAS E INICIALES ---
    import unicodedata
    
    # Función mágica para limpiar acentos y tildes (ej: "García" -> "Garcia")
    def quitar_tildes(texto):
        if not texto: return ""
        return ''.join(c for c in unicodedata.normalize('NFD', str(texto)) if unicodedata.category(c) != 'Mn')
        
    iniciales_corchetes = quitar_tildes(context_data.get('ssag_iniciales_corchetes', '[INICIALES]'))
    iniciales_linea = quitar_tildes(context_data.get('ssag_iniciales_linea', 'AAA/bbb/ccc'))
    
    p_firmas = doc.add_paragraph()
    p_firmas.paragraph_format.left_indent = Cm(1.0)
    p_firmas.paragraph_format.space_before = Pt(60) # Da el espacio de respiro para firmar
    
    run_corchetes = p_firmas.add_run(iniciales_corchetes)
    run_corchetes.font.name = 'Arial'
    run_corchetes.font.size = Pt(11)
    run_corchetes.bold = True
    
    doc.add_paragraph()

    p_linea = doc.add_paragraph()
    p_linea.paragraph_format.left_indent = Cm(1.0)
    run_linea = p_linea.add_run(iniciales_linea)
    run_linea.font.name = 'Arial'
    run_linea.font.size = Pt(8)

    # --- SECCIÓN DE ANEXOS ---
    doc.add_page_break()

    p_anexo1 = doc.add_paragraph()
    p_anexo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. "Anexo n.° 1" en negrita y subrayado
    run_a1 = p_anexo1.add_run("Anexo n.° 1")
    run_a1.font.name = 'Arial'
    run_a1.font.size = Pt(11)
    run_a1.bold = True
    run_a1.underline = True

    # =========================================================
    # BUCLE PARA DIBUJAR LOS ANEXOS DE COSTO EVITADO
    # =========================================================
    import importlib
    
    for i, datos_hecho in enumerate(imputaciones_data):
        id_infraccion = datos_hecho.get('id_infraccion', '')
        
        try:
            modulo_redactor = importlib.import_module(f"textos_infracciones.{id_infraccion}")
        except ImportError:
            modulo_redactor = importlib.import_module("textos_infracciones.generico")
            
        # Comprobamos si el script tiene la función para dibujar anexos
        if hasattr(modulo_redactor, 'redactar_anexo_ce'):
            modulo_redactor.redactar_anexo_ce(doc, datos_hecho, i + 1)
        else:
            # Si no la tiene (como en el genérico), ponemos el texto rojo por defecto
            p_marc_a1 = doc.add_paragraph(f"[Aquí se insertarán las tablas de Costo Evitado del Hecho {i + 1}]")
            p_marc_a1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_marc_a1.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    siguiente_numero_anexo = 2

    if hay_factores_graduacion:
        doc.add_page_break()
        p_anexo2 = doc.add_paragraph()
        p_anexo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_a2 = p_anexo2.add_run(f"Anexo n.° {siguiente_numero_anexo}")
        run_a2.font.name = 'Arial'
        run_a2.font.size = Pt(11)
        run_a2.bold = True
        run_a2.underline = True
        
        doc.add_paragraph()
        
        # BUCLE INTELIGENTE: Dibuja la tabla SOLAMENTE para los hechos donde se activó
        for i, datos_hecho in enumerate(imputaciones_data):
            if datos_hecho.get('aplica_graduacion', 'No') == 'Sí':
                dibujar_tabla_anexo_graduacion(doc, datos_hecho, i)
                doc.add_paragraph()
                
        siguiente_numero_anexo += 1

    doc.add_page_break()
    p_anexo_final = doc.add_paragraph()
    p_anexo_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_af = p_anexo_final.add_run(f"Anexo n.° {siguiente_numero_anexo}")
    run_af.font.name = 'Arial'
    run_af.font.size = Pt(11)
    run_af.bold = True
    run_af.underline = True

    doc.add_paragraph()
    
    p_anexo_sub = doc.add_paragraph()
    p_anexo_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_asub = p_anexo_sub.add_run("(Precios consultados y cotizaciones)")
    run_asub.font.name = 'Arial'
    run_asub.font.size = Pt(11)
    run_asub.bold = True
    run_asub.underline = True

    doc.add_paragraph()

    # --- GUARDADO INICIAL DEL DOCUMENTO MAESTRO ---
    nombre_archivo = f"IFI_{administrado}.docx".replace(" ", "_")
    doc.save(nombre_archivo)
    
# =========================================================
    # FUSIÓN DE DOCUMENTOS EXTERNOS (DOCXCOMPOSE)
    # =========================================================
    try:
        from docxcompose.composer import Composer
        from docx import Document as DocxDocument
        from sheets import descargar_archivo_drive # Importamos la llave maestra de Drive
        
        # 1. Abrimos el Word principal que acabamos de generar
        master_doc = DocxDocument(nombre_archivo)
        composer = Composer(master_doc)
        hubo_anexos = False
        
        # Usamos un Set para no pegar el mismo Word de sustento dos veces
        ids_ya_procesados = set()
        
        # 2. Recorremos los hechos buscando los IDs de las cotizaciones y salarios
        for datos_hecho in imputaciones_data:
            resultados = datos_hecho.get('resultados', {})
            ids_anexos_drive = resultados.get('ids_anexos', [])
            
            for id_drive in ids_anexos_drive:
                # Verificamos que sea un ID válido y no se haya pegado antes
                if id_drive and str(id_drive).strip() not in ['None', ''] and id_drive not in ids_ya_procesados:
                    try:
                        # Descargamos el Word del Drive directamente en la memoria RAM
                        anexo_io = descargar_archivo_drive(id_drive.strip())
                        if anexo_io:
                            anexo_io.seek(0) 
                            doc_anexo = DocxDocument(anexo_io)
                            
                            # --- MAGIA DEL SALTO DE PÁGINA ---
                            # Si YA pegamos un anexo antes, insertamos un salto de página 
                            # en el documento maestro antes de pegar el siguiente.
                            if hubo_anexos:
                                master_doc.add_page_break()
                            
                            # Se anexa al final con todos sus formatos e imágenes
                            composer.append(doc_anexo) 
                            
                            ids_ya_procesados.add(id_drive)
                            hubo_anexos = True
                    except Exception as e:
                        print(f"Error al descargar o fusionar anexo de sustento {id_drive}: {e}")
                        
        # 3. Si pegamos al menos un anexo, sobrescribimos el archivo final
        if hubo_anexos:
            composer.save(nombre_archivo)
            
    except ImportError:
        print("Advertencia: docxcompose no está instalado o falló. No se anexaron los docs de sustento.")

    return nombre_archivo