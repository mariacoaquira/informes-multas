# textos_infracciones/INF004.py
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from ifi import agregar_nota_al_pie
from textos_cok import DICCIONARIO_COK

# =======================================================
# FUNCIONES AUXILIARES PARA FORMATO DE TABLAS Y TEXTOS
# =======================================================
def dar_color_celda(celda, color_hex):
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def centrar_verticalmente_celda(celda):
    tcPr = celda._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def formatear_celda_descripcion(parrafo, texto, es_negrita=False):
    font_name = 'Arial'
    font_size = Pt(10)
    
    def add_r(t, bold=es_negrita, sub=False, sup=False):
        run = parrafo.add_run(t)
        run.font.name = font_name
        run.font.size = font_size
        run.bold = bold
        if sub: run.font.subscript = True
        if sup: run.font.superscript = True

    texto = texto.replace("UIT 2026", "UIT2026")
    # Normalizamos por si acaso en el futuro app.py decide enviarlo con la T
    texto = texto.replace("CE*(1+COSm)T", "CE*(1+COSm)") 
    
    # --- LA CORRECCIÓN ESTÁ AQUÍ ---
    if "CE*(1+COSm)" in texto:
        partes = texto.split("CE*(1+COSm)")
        for i, parte in enumerate(partes):
            if parte: add_r(parte)
            if i < len(partes) - 1:
                add_r("CE*(1+COS")
                add_r("m", sub=True)
                add_r(")")
                add_r("T", sup=True) # Inyectamos la T en superíndice forzosamente
                
    elif "UIT2026" in texto:
        partes = texto.split("UIT2026")
        for i, parte in enumerate(partes):
            if parte: add_r(parte)
            if i < len(partes) - 1:
                add_r("UIT")
                add_r("2026", sub=True)
                
    elif "COSm" in texto:
        partes = texto.split("COSm")
        for i, parte in enumerate(partes):
            if parte: add_r(parte)
            if i < len(partes) - 1:
                add_r("COS")
                add_r("m", sub=True)
                
    else:
        add_r(texto)

def generar_titulo_cuadro(doc, titulo_base):
    if not hasattr(doc, 'contador_cuadros'):
        doc.contador_cuadros = 2
        
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.left_indent = Cm(1.0)
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_titulo = p_titulo.add_run(f"Cuadro n.° {doc.contador_cuadros}: {titulo_base}")
    run_titulo.font.size = Pt(11)
    run_titulo.font.name = 'Arial'
    run_titulo.bold = True
    
    doc.contador_cuadros += 1

def arreglar_letras(datos_tabla, diccionario_notas):
    """Fuerza la sincronía reordenando la tabla y las notas de forma independiente."""
    alfabeto = "abcdefghijklmnopqrstuvwxyz"
    
    # ==========================================
    # 1. ARREGLAMOS LA TABLA (ej: a, b, c, e -> a, b, c, d)
    # ==========================================
    letras_tabla = set()
    for item in datos_tabla:
        sup = str(item.get("ref", item.get("descripcion_superindice", "")))
        if sup and sup not in ["None", ""]:
            letras_tabla.add(sup)
            
    letras_tabla_ordenadas = sorted(list(letras_tabla))
    mapa_tabla = {vieja: alfabeto[i] for i, vieja in enumerate(letras_tabla_ordenadas)}
    
    for item in datos_tabla:
        sup = str(item.get("ref", item.get("descripcion_superindice", "")))
        if sup in mapa_tabla:
            item['ref'] = mapa_tabla[sup]
            item['descripcion_superindice'] = mapa_tabla[sup]
            
    # ==========================================
    # 2. ARREGLAMOS LAS NOTAS (ej: a, b, c, d -> a, b, c, d)
    # ==========================================
    letras_notas_ordenadas = sorted(list(diccionario_notas.keys()))
    mapa_notas = {vieja: alfabeto[i] for i, vieja in enumerate(letras_notas_ordenadas)}
    
    nuevas_notas = {}
    for vieja, texto in diccionario_notas.items():
        nuevas_notas[mapa_notas[vieja]] = texto
        
    return datos_tabla, nuevas_notas

def dibujar_notas_cuadro(doc, diccionario_notas, datos_hecho, extremo_data=None, bi_data=None):
    """Dibuja Fuente y Elaboración. Rescata los textos del COK y formatea la nota al pie."""
    if not diccionario_notas: return
    
    p_fuente = doc.add_paragraph()
    p_fuente.paragraph_format.left_indent = Cm(1.0)
    p_fuente.add_run("Fuente:").font.size = Pt(8)
    
    resultados_app = datos_hecho.get('resultados', {}).get('resultados_para_app', {})
    id_rubro = resultados_app.get('id_rubro', 'R016')
    unidad_fisc = resultados_app.get('administrado', '[UNIDAD FISCALIZABLE]') 
    
    letra_cok = None
    if bi_data:
        for item in bi_data:
            if "COS" in str(item.get("descripcion", "")):
                letra_cok = str(item.get("ref", ""))
                if letra_cok: break

    fecha_inc = "[FECHA_INC]"
    fecha_ext = "[FECHA_EXT]"
    if extremo_data:
        fecha_inc = extremo_data.get('fecha_incumplimiento_texto', datos_hecho.get('fecha_incumplimiento_texto', fecha_inc))
        fecha_ext = extremo_data.get('fecha_extemporanea_texto', datos_hecho.get('fecha_extemporanea_texto', fecha_ext))
    
    contexto = datos_hecho.get('context_data', {})
    fecha_hoy = contexto.get('fecha_hoy', '[FECHA_HOY]')
    
    for letra, texto_nota in sorted(diccionario_notas.items()):
        p_nota = doc.add_paragraph()
        p_nota.paragraph_format.left_indent = Cm(1.5)
        p_nota.paragraph_format.first_line_indent = Cm(-0.5)
        p_nota.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        p_nota.add_run(f"{letra})\t").font.size = Pt(8)
        texto_str = str(texto_nota)
        
        texto_str = texto_str.replace("{fecha_incumplimiento_texto}", str(fecha_inc))
        texto_str = texto_str.replace("{fecha_extemporanea_texto}", str(fecha_ext))
        texto_str = texto_str.replace("{fecha_hoy_texto}", str(fecha_hoy))
        
        if letra == letra_cok or "{fuente_cos}" in texto_str or "Costo de Oportunidad" in texto_str:
            datos_cok = DICCIONARIO_COK.get(id_rubro)
            if not datos_cok: 
                datos_cok = DICCIONARIO_COK.get("R016") 
                
            if datos_cok:
                texto_str = datos_cok.get("texto", texto_str)
                if "[NOTA_SECTOR]" in texto_str:
                    partes_cok = texto_str.split("[NOTA_SECTOR]")
                    
                    p_nota.add_run(partes_cok[0]).font.size = Pt(8)
                    
                    nota_pie_texto = datos_cok.get("nota_pie", "").replace("{unidad_fiscalizable}", unidad_fisc)
                    agregar_nota_al_pie(p_nota, nota_pie_texto, doc)
                    
                    if p_nota.runs:
                        p_nota.runs[-1].font.size = Pt(8)
                    
                    texto_str = partes_cok[1] 
        
        if texto_str:
            p_nota.add_run(texto_str).font.size = Pt(8)

    p_elab = doc.add_paragraph()
    p_elab.paragraph_format.left_indent = Cm(1.0)
    p_elab.add_run("Elaboración: Subdirección de Sanción y Gestión de Incentivos (SSAG) – DFAI.").font.size = Pt(8)

# =======================================================
# 1. FUNCIÓN PARA EL CUERPO DEL INFORME (Solo BI)
# =======================================================
def redactar_beneficio_ilicito(doc, datos_hecho, num_hecho):
    # 1. Primer párrafo
    p_desc_bi = doc.add_paragraph(
        "El beneficio ilícito proviene del costo evitado de no remitir la información "
        "requerida en el plazo legal establecido. En este caso, el administrado incumplió lo precisado "
        "en el párrafo anterior."
    )
    p_desc_bi.paragraph_format.left_indent = Cm(1.0)
    p_desc_bi.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()

    # --- LA MAGIA: MEMORIA DE ACRÓNIMOS ---
    if not hasattr(doc, 'acronimos_usados'):
        doc.acronimos_usados = set()
        
    p_escenario = doc.add_paragraph()
    p_escenario.paragraph_format.left_indent = Cm(1.0)
    p_escenario.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_escenario.add_run(
        "En el escenario de cumplimiento, el administrado debería llevar a cabo las inversiones necesarias "
        "para cumplir con sus compromisos ambientales. En tal sentido, para el cálculo del "
    )
    
    # Aplicamos Negrita y Cursiva al acrónimo CE
    if 'CE' not in doc.acronimos_usados:
        p_escenario.add_run("costo evitado total (en adelante, ")
        run_ce = p_escenario.add_run("CE")
        run_ce.bold = True
        run_ce.italic = True
        p_escenario.add_run(")")
        doc.acronimos_usados.add('CE') 
    else:
        run_ce = p_escenario.add_run("CE")
        run_ce.bold = True
        run_ce.italic = True
        
    p_escenario.add_run(" se ha considerado como mínimo indispensable, el desarrollo de la siguiente actividad")
    
    # Insertamos la nota al pie del anexo
    texto_nota_ce = "Para mayor detalle, ver Anexo n.° 1 del presente informe."
    agregar_nota_al_pie(p_escenario, texto_nota_ce, doc)
    p_escenario.add_run(":")
    
    doc.add_paragraph()

    # =======================================================
    # VIÑETA DE FLECHA PARA LA ACTIVIDAD (Buscando en contexto_final_word)
    # =======================================================
    # Extraemos los resultados matemáticos
    resultados = datos_hecho.get('resultados', {})
    
    # ¡AQUÍ ESTÁ LA CLAVE! Buscamos en el diccionario antiguo de las plantillas
    contexto_word = resultados.get('contexto_final_word', {})
    
    # Extraemos las horas exactamente como se llamaban en tu Jinja
    horas_trabajo = contexto_word.get('plazo_final_horas_extremo')
    
    # Fallback por si acaso
    if not horas_trabajo:
        horas_trabajo = datos_hecho.get('plazo_final_horas_extremo')
    if not horas_trabajo:
        horas_trabajo = "[HORAS]" 

    p_vineta = doc.add_paragraph()
    pf = p_vineta.paragraph_format
    pf.left_indent = Cm(1.63)
    pf.first_line_indent = Cm(-0.63)
    pf.tab_stops.clear_all()
    pf.tab_stops.add_tab_stop(Cm(1.63))
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run_bullet = p_vineta.add_run("➤\t")
    run_bullet.font.name = 'Arial'
    run_bullet.font.size = Pt(11)
    
    # Texto subrayado y en negrita
    run_title = p_vineta.add_run("CE: Sistematización y remisión de la información")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(11)
    run_title.bold = True
    run_title.underline = True
    
    # Texto normal con la variable de horas inyectada
    run_text = p_vineta.add_run(
        " a presentar, en este caso, la información requerida por la Autoridad de Supervisión. "
        "Para dicha actividad se considera, como mínimo indispensable a un (1) profesional encargado de "
        "recopilar, revisar, validar, enviar y dar seguimiento a la información a presentar, asegurando "
        "el cumplimiento de los requisitos de forma, modo y plazo establecidos por la normativa, por un "
        f"periodo de trabajo de {horas_trabajo}"
    )
    run_text.font.name = 'Arial'
    run_text.font.size = Pt(11)
    
    # Insertamos la nota al pie de las horas
    texto_nota_horas = "Se consideran ocho (8) horas de jornada laboral por día."
    agregar_nota_al_pie(p_vineta, texto_nota_horas, doc)
    p_vineta.add_run(".")
    
    doc.add_paragraph()

    # =======================================================
    # NUEVO PÁRRAFO: JUSTIFICACIÓN DE PLAZOS Y TFA
    # =======================================================
    p_tfa = doc.add_paragraph()
    # Alineamos a 1.63 cm para que calce exacto con el texto de la viñeta
    p_tfa.paragraph_format.left_indent = Cm(1.63)
    p_tfa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_tfa.add_run("Respecto a los días destinados para esta actividad, y en concordancia con lo señalado por el ")
    
    # Lógica de memoria para TFA
    if 'TFA' not in doc.acronimos_usados:
        p_tfa.add_run("Tribunal de Fiscalización Ambiental (en adelante, ")
        run_tfa = p_tfa.add_run("TFA")
        run_tfa.bold = True
        run_tfa.italic = True
        p_tfa.add_run(")")
        doc.acronimos_usados.add('TFA')
    else:
        run_tfa = p_tfa.add_run("TFA")
        run_tfa.bold = True
        run_tfa.italic = True
        
    p_tfa.add_run(
        " mediante las Resoluciones n.° 740-2024-OEFA/TFA-SE y n.° 419-2023-OEFA/TFA-SE, se "
        "entiende que los plazos establecidos para la remisión de reportes se ajustan al período "
        "en el cual el recurrente no pudo cumplir con la entrega de la información requerida."
    )
    
    doc.add_paragraph()

    # =======================================================
    # PÁRRAFO DEL PLAZO OTORGADO CON NOTA AL PIE CONDICIONAL
    # =======================================================
    # 1. Extraemos las variables desde el contexto de la infracción
    resultados = datos_hecho.get('resultados', {})
    contexto_word = resultados.get('contexto_final_word', {})
    
    plazo_total_dias = contexto_word.get('plazo_total_dias', '[DÍAS]')
    plazo_total_horas = contexto_word.get('plazo_total_horas', '[HORAS]')
    doc_req_num = contexto_word.get('doc_req_num', '[DOC_REQ]')
    fecha_req = contexto_word.get('fecha_requerimiento', '[FECHA_REQ]')
    dias_orig = contexto_word.get('dias_habiles_orig', '[DÍAS_ORIG]')
    
    aplica_amp = contexto_word.get('aplica_ampliacion', False)
    doc_amp_num = contexto_word.get('doc_amp_num', '[DOC_AMP]')
    doc_amp_fecha = contexto_word.get('doc_amp_fecha', '[FECHA_AMP]')
    dias_amp = contexto_word.get('dias_habiles_amp', '[DÍAS_AMP]')
    
    fecha_max = contexto_word.get('fecha_max_presentacion', '[FECHA_MAX]')
    es_extemporaneo = contexto_word.get('es_extemporaneo', False)
    fecha_ext = contexto_word.get('fecha_extemporanea', '[FECHA_EXT]')
    
    # Aquí puedes extraer el tipo de incumplimiento si necesitaras variar el texto base
    tipo_incumplimiento = datos_hecho.get('tipo_seleccionado', '')

    # 2. Armamos la nota al pie con lógica condicional de Python
    texto_nota_plazo = (
        f"Mediante {doc_req_num}, del {fecha_req}, la Autoridad de Supervisión requirió al administrado "
        f"información y documentación a fin de verificar el cumplimiento de sus obligaciones ambientales, "
        f"otorgándole un plazo de {dias_orig}. "
    )
    
    # Condicional 1: Si aplica ampliación
    if aplica_amp:
        texto_nota_plazo += f"Posteriormente, mediante {doc_amp_num}, del {doc_amp_fecha}, se le concedió un plazo adicional de {dias_amp}. "
        
    texto_nota_plazo += (
        f"En consecuencia, el administrado contó con un plazo total de {plazo_total_dias}, equivalentes a {plazo_total_horas} "
        f"de trabajo para la presentación de la información requerida en su totalidad. Al respecto, cabe precisar que el plazo "
        f"para la presentación venció el {fecha_max}"
    )
    
    # Condicional 2: Si es extemporáneo
    if es_extemporaneo:
        texto_nota_plazo += f"; sin embargo, la información fue remitida extemporáneamente el {fecha_ext}."
    else:
        texto_nota_plazo += "."

    # 3. Dibujamos el párrafo a 1.63 cm
    p_plazo = doc.add_paragraph()
    p_plazo.paragraph_format.left_indent = Cm(1.63)
    p_plazo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Puedes aplicar ifs aquí también para el texto del párrafo si varía según el tipo
    # if tipo_incumplimiento == "No remitió": ...
    
    p_plazo.add_run(
        f"En esa línea, se considera razonable tomar en cuenta el plazo otorgado para la presentación "
        f"de la información, es decir, {plazo_total_dias}, equivalentes a {plazo_total_horas}"
    )
    
    # Insertamos la nota al pie viva
    agregar_nota_al_pie(p_plazo, texto_nota_plazo, doc)
    
    p_plazo.add_run(
        " de trabajo, el cual ofrece un margen adecuado para que el personal responsable lleve a cabo "
        "las tareas de recolección, validación y sistematización de los datos, asegurando que la "
        "información remitida cumpla con los criterios exigidos, sin comprometer su calidad ni precisión."
    )
    
    doc.add_paragraph()

    # =======================================================
    # PÁRRAFO CONDICIONAL: PROPORCIONALIDAD (Solo si es incompleto)
    # =======================================================
    
    # 1. Extraemos los números PUROS directamente de los inputs de la aplicación (100% seguro)
    total_items_num = int(datos_hecho.get('num_items_solicitados', 1))
    
    extremos_crudos = datos_hecho.get('extremos', [])
    items_extremo_num = int(extremos_crudos[0].get('cantidad_items', 1)) if extremos_crudos else total_items_num

    # ¡LA CONDICIÓN MAGISTRAL! Solo se ejecuta si los ítems del extremo son menores al total
    if items_extremo_num < total_items_num:
        
        import re
        from funciones import texto_con_numero
        
        def extraer_numero(valor, default=1):
            if isinstance(valor, (int, float)): return float(valor)
            if not valor: return default
            match = re.search(r'\d+(\.\d+)?', str(valor))
            return float(match.group()) if match else default

        # Extraemos los textos bonitos y formateados desde el contexto para redactar
        total_items_txt = contexto_word.get('total_items_requeridos', str(total_items_num))
        items_extremo_txt = contexto_word.get('items_extremo_actual', str(items_extremo_num))
        
        # Extraemos valores en crudo para pasarlos por el formateador
        horas_por_item_raw = contexto_word.get('horas_por_item_unitario', 1)
        plazo_final_dias_raw = contexto_word.get('plazo_final_dias_extremo', 1)
        
        h_item_num = extraer_numero(horas_por_item_raw, 1)
        dias_fin_num = extraer_numero(plazo_final_dias_raw, 1)
        
        # Formateamos inteligentemente (Singular o Plural)
        horas_por_item_txt = f"{texto_con_numero(h_item_num, genero='f')} {'hora' if h_item_num == 1 else 'horas'}"
        plazo_final_dias_txt = f"{texto_con_numero(dias_fin_num, genero='m')} {'día hábil' if dias_fin_num == 1 else 'días hábiles'}"

        p_proporcion = doc.add_paragraph()
        p_proporcion.paragraph_format.left_indent = Cm(1.63) # A la misma altura
        p_proporcion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Condicional inline para el texto
        texto_estado = "fue remitida fuera de plazo" if es_extemporaneo else "no fue remitida"
        
        p_proporcion.add_run(
            f"No obstante, dichas {plazo_total_horas} fueron establecidas para la presentación de la totalidad "
            f"de la información requerida. En ese sentido, se considera razonable tomar en cuenta únicamente la "
            f"proporción del plazo correspondiente a la parte de la información que {texto_estado}. "
            f"Por lo tanto, el tiempo a considerarse en el presente caso asciende a {horas_trabajo}"
        )
        
        # Armamos la nota al pie usando los textos dinámicos
        texto_nota_proporcion = (
            f"El requerimiento comprendía {total_items_txt}, para cuya atención se otorgó un plazo total de "
            f"{plazo_total_dias}, equivalentes a {plazo_total_horas} de trabajo, lo que representa una estimación de "
            f"{horas_por_item_txt} por ítem. En tal sentido, aplicando una estimación proporcional en función de "
            f"{items_extremo_txt} de remisión, se determina un periodo de {horas_trabajo} de trabajo (equivalentes a "
            f"{plazo_final_dias_txt}), el cual se considera para la estimación del costo evitado asociado al profesional "
            f"responsable y al uso del equipo informático correspondiente."
        )
        
        # Insertamos la nota al pie viva
        agregar_nota_al_pie(p_proporcion, texto_nota_proporcion, doc)
        
        p_proporcion.add_run(" de trabajo.")
        
        doc.add_paragraph()

    # =======================================================
    # PÁRRAFO DE ALQUILER DE EQUIPO (LAPTOP)
    # =======================================================
    p_laptop = doc.add_paragraph()
    p_laptop.paragraph_format.left_indent = Cm(1.63) # A la misma altura
    p_laptop.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_laptop.add_run(
        "Además, se contempla el alquiler de una laptop para la sistematización y remisión "
        "de la información por los días de trabajo"
    )
    
    # Armamos e insertamos la nota al pie
    texto_nota_laptop = (
        "Se contempla la misma cantidad de días de alquiler de laptop que el tiempo de contratación del profesional, "
        "dado que éste hará uso del equipo para la realización de sus actividades. Se consideran ocho (8) horas de "
        "jornada laboral por día."
    )
    agregar_nota_al_pie(p_laptop, texto_nota_laptop, doc)
    
    p_laptop.add_run(".")
    
    doc.add_paragraph()

    # =======================================================
    # PÁRRAFO FINAL: CAPITALIZACIÓN Y ACTUALIZACIÓN
    # =======================================================
    # Extraemos la variable para saber si el cálculo es en dólares
    bi_moneda_es_dolares = contexto_word.get('bi_moneda_es_dolares', False)

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.left_indent = Cm(1.0) # Regresamos a 1.0 cm
    p_cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Separamos en tramos para aplicarle estilo al CE
    p_cap.add_run("Una vez estimado el ")
    run_ce_final = p_cap.add_run("CE")
    run_ce_final.bold = True
    run_ce_final.italic = True
    
    p_cap.add_run("; este monto es capitalizado aplicando el ")
    
    # 1. Lógica de memoria para COS
    if 'COS' not in doc.acronimos_usados:
        p_cap.add_run("Costo de Oportunidad Sectorial (en adelante, ")
        run_cos = p_cap.add_run("COS")
        run_cos.bold = True
        run_cos.italic = True
        p_cap.add_run(")")
        doc.acronimos_usados.add('COS')
    else:
        run_cos = p_cap.add_run("COS")
        run_cos.bold = True
        run_cos.italic = True

    # 2. Condicional principal (Extemporáneo vs Normal)
    if es_extemporaneo:
        p_cap.add_run(" desde la fecha de inicio del presunto incumplimiento hasta la fecha de cumplimiento extemporáneo. hasta la fecha del cálculo de la multa. ")
        
        if bi_moneda_es_dolares:
            p_cap.add_run("Luego el valor obtenido es transformado a moneda nacional. ")
            
        p_cap.add_run("Finalmente, el resultado es actualizado hasta la fecha de emisión del presente informe mediante un ajuste inflacionario y expresado en la ")
    else:
        p_cap.add_run(" desde la fecha de inicio del presunto incumplimiento hasta la fecha del cálculo de la multa. Finalmente, el resultado es ")
        
        if bi_moneda_es_dolares:
            p_cap.add_run("transformado a moneda nacional y ")
            
        p_cap.add_run("expresado en la ")

    # 3. Lógica de memoria para UIT
    if 'UIT' not in doc.acronimos_usados:
        p_cap.add_run("Unidad Impositiva Tributaria (en adelante, ")
        run_uit = p_cap.add_run("UIT")
        run_uit.bold = True
        run_uit.italic = True
        p_cap.add_run(")")
        doc.acronimos_usados.add('UIT')
    else:
        run_uit = p_cap.add_run("UIT")
        run_uit.bold = True
        run_uit.italic = True

    p_cap.add_run(" vigente. El detalle del beneficio ilícito se presenta en el siguiente cuadro:")
    
    doc.add_paragraph()

    # =======================================================
    # DIBUJO DE TABLAS Y NOTAS
    # =======================================================
    resultados_app = datos_hecho.get('resultados', {}).get('resultados_para_app', {})
    extremos = resultados_app.get('extremos', [])

    if extremos and isinstance(extremos, list):
        for j, extremo_data in enumerate(extremos):
            bi_data = extremo_data.get('bi_data', [])
            notas_bi = extremo_data.get('diccionario_notas', {}) 
            
            # 1. Sincronizamos letras ANTES de dibujar
            bi_data_arreglado, notas_bi_arreglado = arreglar_letras(bi_data, notas_bi)
            
            # 2. Dibujamos la tabla y las notas con los datos arreglados
            dibujar_tabla_bi(doc, bi_data_arreglado, "Beneficio Ilícito")
            dibujar_notas_cuadro(doc, notas_bi_arreglado, datos_hecho, extremo_data, bi_data_arreglado)
            doc.add_paragraph()
    else:
        totales_finales = resultados_app.get('totales', resultados_app)
        bi_data_raw = totales_finales.get('bi_data_raw', [])
        notas_bi = extremos[0].get('diccionario_notas', {}) if extremos else {}
        
        # 1. Sincronizamos letras ANTES de dibujar
        bi_data_arreglado, notas_bi_arreglado = arreglar_letras(bi_data_raw, notas_bi)
        
        # 2. Dibujamos la tabla y las notas con los datos arreglados
        dibujar_tabla_bi(doc, bi_data_arreglado, "Beneficio Ilícito")
        dibujar_notas_cuadro(doc, notas_bi_arreglado, datos_hecho, bi_data=bi_data_arreglado)
        doc.add_paragraph()

    # =======================================================
    # PÁRRAFO FINAL DEL BENEFICIO ILÍCITO
    # =======================================================
    # Obtenemos el total del Beneficio Ilícito de los resultados
    totales_finales = resultados_app.get('totales', resultados_app)
    bi_uit_val = totales_finales.get('beneficio_ilicito_uit', 0.0)
    
    # Lo formateamos a 3 decimales y le añadimos " UIT" si no lo tiene
    if isinstance(bi_uit_val, (int, float)):
        bi_uit_texto = f"{bi_uit_val:,.3f} UIT"
    else:
        bi_uit_texto = f"{bi_uit_val}" if "UIT" in str(bi_uit_val) else f"{bi_uit_val} UIT"
        
    p_bi_final = doc.add_paragraph()
    p_bi_final.paragraph_format.left_indent = Cm(1.0)
    p_bi_final.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Lo dividimos en tramos para darle negrita solo a la variable
    p_bi_final.add_run("De acuerdo con lo anterior, el beneficio ilícito estimado para esta presunta infracción asciende a ")
    
    run_bi_val = p_bi_final.add_run(bi_uit_texto)
    run_bi_val.bold = True
    
    p_bi_final.add_run(".")
    
    doc.add_paragraph()

# =======================================================
# 4. FUNCIÓN PARA PROBABILIDAD DE DETECCIÓN
# =======================================================
def redactar_probabilidad(doc, datos_hecho, num_hecho):
    p_prob = doc.add_paragraph()
    p_prob.paragraph_format.left_indent = Cm(1.5)
    p_prob.paragraph_format.first_line_indent = Cm(-0.5) 
    p_prob.paragraph_format.tab_stops.clear_all()
    p_prob.paragraph_format.tab_stops.add_tab_stop(Cm(1.5))
    
    run_prob = p_prob.add_run("ii)\tProbabilidad de Detección (p)")
    run_prob.font.name = 'Arial'
    run_prob.font.size = Pt(11)
    run_prob.bold = True

    doc.add_paragraph()

    # Párrafo principal a 1.0 cm de sangría
    p_desc_prob = doc.add_paragraph()
    p_desc_prob.paragraph_format.left_indent = Cm(1.0)
    p_desc_prob.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_desc_prob.add_run("En el presente caso se ha considerado una probabilidad de detección muy alta (1.0)")
    
    # Inserción de la nota al pie viva
    texto_nota_prob = (
        "Conforme con la tabla n.° 1 del Anexo II de la Metodología para el cálculo de las multas base y la "
        "aplicación de los factores de gradualidad a utilizar en la graduación de sanciones, aprobada mediante "
        "Resolución de Presidencia del Consejo Directivo n.° 035-2013-OEFA/PCD y modificada por Resolución de "
        "Consejo Directivo n.° 024-2017-OEFA/CD."
    )
    agregar_nota_al_pie(p_desc_prob, texto_nota_prob, doc)
    
    # Continuación de la oración
    p_desc_prob.add_run(" porque la autoridad pudo conocer e identificar esta infracción con facilidad, toda vez que ésta se encuentra sujeta a un plazo específico, cuya fecha de vencimiento es conocida.")
    
    doc.add_paragraph()

# =======================================================
# 5. FUNCIÓN PARA PRINCIPIOS DE TIPIFICACIÓN
# =======================================================
def redactar_principios(doc, datos_hecho, num_hecho, monto_uit):
    p_princ = doc.add_paragraph()
    p_princ.paragraph_format.left_indent = Cm(1.5)
    p_princ.paragraph_format.first_line_indent = Cm(-0.5)
    p_princ.paragraph_format.tab_stops.clear_all()
    p_princ.paragraph_format.tab_stops.add_tab_stop(Cm(1.5))
    
    # --- AQUÍ ESTÁ EL CAMBIO: Numeración Dinámica ---
    numeracion = datos_hecho.get('numeracion_principios', 'v)')
    
    run_princ = p_princ.add_run(f"{numeracion}\tAplicación de los Principios: Tipificación y Razonabilidad")
    run_princ.font.name = 'Arial'
    run_princ.font.size = Pt(11)
    run_princ.bold = True

    doc.add_paragraph()

    # --- PRIMER PÁRRAFO: Tipificación ---
    p_desc_princ1 = doc.add_paragraph()
    p_desc_princ1.paragraph_format.left_indent = Cm(1.0)
    p_desc_princ1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_desc_princ1.add_run(
        "En aplicación a lo previsto en el numeral 1.2 del Cuadro de Tipificación de Infracciones y Escala "
        "de Sanciones vinculadas con la eficacia de la Fiscalización Ambiental, aprobada por Resolución de "
        "Consejo Directivo n.° 042-2013-OEFA/CD y sus modificatorias; se dispuso que el monto aplicable para "
        "una infracción de este tipo está en el rango de "
    )
    
    run_rango = p_desc_princ1.add_run("hasta 100 UIT")
    run_rango.bold = True
    p_desc_princ1.add_run(".")
    
    doc.add_paragraph()

    # --- SEGUNDO PÁRRAFO: Razonabilidad con nota al pie ---
    p_desc_princ2 = doc.add_paragraph()
    p_desc_princ2.paragraph_format.left_indent = Cm(1.0)
    p_desc_princ2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_desc_princ2.add_run("Con relación al principio de razonabilidad, en línea con lo aprobado mediante Resolución de Consejo Directivo n.° 001-2020-OEFA/CD")
    
    texto_nota_princ = (
        "El OEFA dispuso que la multa determinada con la metodología de cálculo de multas base y la aplicación de los factores "
        "para la graduación de sanciones, constituye la sanción monetaria correspondiente, prevaleciendo este monto sobre el "
        "valor del tope mínimo previsto para el respectivo tipo infractor."
    )
    agregar_nota_al_pie(p_desc_princ2, texto_nota_princ, doc)
    
    p_desc_princ2.add_run(
        ", se verifica que, al encontrarse la multa calculada en el rango normativo vigente, "
        "corresponde sancionar al administrado con dicho monto, el cual asciende a "
    )
    
    # Formateamos la variable para asegurarnos que diga "UIT"
    texto_monto_final = f"{monto_uit}" if "UIT" in str(monto_uit) else f"{monto_uit} UIT"
    
    run_monto = p_desc_princ2.add_run(texto_monto_final)
    run_monto.bold = True
    p_desc_princ2.add_run(".")
    
    doc.add_paragraph()

# =======================================================
# 2. FUNCIÓN PARA LA SECCIÓN DE ANEXOS (Solo CE)
# =======================================================
def dibujar_fuente_anexo_inf004(doc, ctx_word):
    # 1. Extraemos TODAS las variables almacenadas en el contexto_final_word
    fuente_salario = ctx_word.get('fuente_salario', '[FUENTE_SALARIO]')
    pdf_salario = ctx_word.get('pdf_salario', '[ENLACE_SALARIO]')
    sustento = ctx_word.get('sustento_item_profesional', '')
    anexo_num = ctx_word.get('ph_anexo_ce_num', '[N_ANEXO]')
    fi_mes = ctx_word.get('fi_mes', '[MES_FI]')
    fi_ipc = ctx_word.get('fi_ipc', '[IPC]')
    ref_ipc = ctx_word.get('ref_ipc_salario', '[REF_IPC]')
    fi_tc = ctx_word.get('fi_tc', '[TC]')
    fecha_hoy = ctx_word.get('fecha_hoy', '[FECHA_HOY]')
    
    # Convertimos los valores a texto usando sus decimales exactos (sin redondeos forzados)
    fi_ipc_str = str(fi_ipc)
    fi_tc_str = str(fi_tc)

    # Función interna para pintar rápidamente los párrafos
    def add_p(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0) # Alineado al borde izquierdo del anexo
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.name = 'Arial'
        if bold: run.bold = True

    # 2. Redactamos bloque por bloque exactamente con el formato solicitado
    add_p("Fuente:")
    add_p(f"1/ {fuente_salario}")
    add_p("Disponible en:")
    add_p(f"{pdf_salario}")
    
    if sustento and sustento != 'None':
        add_p(f"{sustento}")
        
    add_p("Nota 1: De acuerdo con el Artículo 8° del Texto Único Ordenado del Decreto Legislativo n.° 728, Ley de Productividad y Competitividad Laboral, la remuneración se establece mensualmente, salvo que se pacte expresamente por semana, día u hora. Para determinar la remuneración por hora, el monto mensual se divide entre treinta (30) días, independientemente del número de días del mes, y luego entre ocho (8), considerando una jornada laboral estándar de 8 horas diarias. De este modo, se obtiene la remuneración por hora de cada trabajador. Este método garantiza la proporcionalidad y razonabilidad en la determinación de las remuneraciones según el tiempo efectivamente trabajado.")
    add_p("Mayor detalle del texto citado, ver el siguiente enlace:")
    add_p("https://www2.congreso.gob.pe/sicr/cendocbib/con4_uibd.nsf/BE35EA4B0DF56C0A05257E2200538D4C/$FILE/1_DECRETO_SUPREMO_003_27_03_1997.pdf")
    add_p(f"Fecha de consulta: {fecha_hoy}.")
    add_p("Nota 2: Con esta determinación de los salarios se busca estimar la remuneración efectiva por hora trabajada, sin perder de vista la razonabilidad de los cálculos. Ello en virtud de que en un escenario de información asimétrica – el administrado, por la envergadura de sus actividades, a pesar de contar con información de comprobantes de pago, no la revela a la autoridad – este despacho no tiene acceso a los salarios reales en que incurre el administrado, cuyos rubros pueden incluir: remuneración básica, asignación familiar, vacaciones, otras bonificaciones extraordinarias, catorce sueldos al año, bonos por sindicato, entre otros.")
    add_p(f"2/ Para mayor detalle del costo, ver Anexo n.° {anexo_num}.")
    add_p(f"3/ El factor de ajuste permite actualizar los valores de fecha de costeo a fecha de incumplimiento. Para ello, dividimos el Índice de Precios al Consumidor (IPC) de fecha de incumplimiento ({fi_mes}, IPC= {fi_ipc_str}) entre el IPC disponible a la fecha de costeo. Los IPC empleados son:")
    add_p(f"- Referencia 1/: {ref_ipc}.")
    add_p("- Referencia 2/: Setiembre 2024, IPC= 114.050464.")
    add_p("El resultado final fue expresado en dos decimales como se aprecia en la tabla.")
    add_p("Nota: Índice de precios Lima Metropolitana (índice Dic.2021 = 100).")
    add_p(f"4/ Banco central de Reserva del Perú (BCRP), 2023. Series Estadísticas. Tipo de Cambio Nominal Bancario – Promedio a la fecha de incumplimiento ({fi_mes}, TC={fi_tc_str}).")
    add_p("Disponible en:")
    add_p("https://estadisticas.bcrp.gob.pe/estadisticas/series/mensuales/resultados/PN01210PM/html")
    add_p(f"Fecha de consulta: {fecha_hoy}.")
    add_p("(*) A fecha de incumplimiento. Cabe mencionar que, para el presente caso, se consideran el IPC y el TC correspondiente al mes de incumplimiento.")
    add_p("Elaboración: Subdirección de Sanción y Gestión de Incentivos (SSAG) – DFAI.")


def redactar_anexo_ce(doc, datos_hecho, num_hecho):
    resultados = datos_hecho.get('resultados', {})
    resultados_app = resultados.get('resultados_para_app', {})
    
    # AQUÍ ESTÁ LA LLAVE: Rescatamos el baúl donde se guardaron las variables
    ctx_word = resultados.get('contexto_final_word', {}) 
    extremos = resultados_app.get('extremos', [])
    
    doc.add_paragraph() 
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"Hecho imputado n.° {num_hecho}")
    run_sub.bold = True
    run_sub.underline = True
    run_sub.font.size = Pt(11)
    run_sub.font.name = 'Arial'
    
    doc.add_paragraph()
    
    if extremos and isinstance(extremos, list) and len(extremos) > 1:
        for j, extremo_data in enumerate(extremos):
            ce_data = extremo_data.get('ce_data', [])
            dibujar_tabla_ce_inf004(doc, ce_data, "Costo Evitado")
            dibujar_fuente_anexo_inf004(doc, ctx_word)
            doc.add_paragraph()
    else:
        totales_finales = resultados_app.get('totales', resultados_app)
        ce_data_raw = totales_finales.get('ce_data_raw', [])
        if not ce_data_raw and extremos:
            ce_data_raw = extremos[0].get('ce_data', [])
            
        dibujar_tabla_ce_inf004(doc, ce_data_raw, "Costo Evitado")
        dibujar_fuente_anexo_inf004(doc, ctx_word)
        doc.add_paragraph()

# =======================================================
# 3. FUNCIONES DE DIBUJO DE TABLAS 
# =======================================================

def dibujar_tabla_bi(doc, bi_data, titulo):
    if not bi_data: return
        
    generar_titulo_cuadro(doc, titulo)
    
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = 'Table Grid'
    tabla.autofit = False
    
    tabla.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = tabla._tblPr
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '680') 
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    ancho_desc = Cm(10.5)
    ancho_monto = Cm(3.5)
    
    headers = ["Descripción", "Monto"]
    for i, header in enumerate(headers):
        celda = tabla.rows[0].cells[i]
        celda.width = ancho_desc if i == 0 else ancho_monto
        dar_color_celda(celda, 'E7E6E6')
        centrar_verticalmente_celda(celda) 
        
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        
    for idx, item in enumerate(bi_data):
        row_cells = tabla.add_row().cells
        row_cells[0].width = ancho_desc
        row_cells[1].width = ancho_monto
        
        es_ultima_fila = (idx == len(bi_data) - 1)
        
        for celda in row_cells:
            centrar_verticalmente_celda(celda) 
            if es_ultima_fila:
                dar_color_celda(celda, 'E7E6E6')
            
        desc = str(item.get("descripcion_texto", item.get("descripcion", "")))
        superindice = str(item.get("ref", item.get("descripcion_superindice", "")))
        if superindice == "None": superindice = ""
        
        p_desc = row_cells[0].paragraphs[0]
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        formatear_celda_descripcion(p_desc, desc, es_negrita=es_ultima_fila)
            
        if superindice:
            run_sup = p_desc.add_run(f"({superindice})")
            run_sup.font.superscript = True
            run_sup.font.name = 'Arial'
            run_sup.font.size = Pt(10)
            run_sup.bold = es_ultima_fila
            
        monto = item.get("monto", "")
        p_monto = row_cells[1].paragraphs[0]
        p_monto.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_monto = p_monto.add_run(str(monto))
        run_monto.font.name = 'Arial'
        run_monto.font.size = Pt(10)
        run_monto.bold = es_ultima_fila

def dibujar_tabla_ce_inf004(doc, ce_data, titulo):
    if not ce_data: return
        
    p_tit = doc.add_paragraph()
    p_tit.paragraph_format.left_indent = Cm(0) 
    run_tit = p_tit.add_run("1. Costo de Sistematización y remisión de información - CE")
    run_tit.bold = True
    run_tit.font.size = Pt(8)
    run_tit.font.name = 'Arial'
    
    headers = ["Descripción", "Unidad", "Cantidad", "Monto (S/)", "Factor de ajuste 3/", "Monto (*)\n(S/)", "Monto (*)\n(US$) 4/"]
    llaves = ["descripcion", "unidad", "cantidad_real", "precio_soles", "factor_ajuste", "monto_soles", "monto_dolares"]
    
    tabla = doc.add_table(rows=1, cols=len(headers))
    
    # ---------------------------------------------------------
    # MAGIA XML: SOLO BORDES HORIZONTALES
    # ---------------------------------------------------------
    tblPr = tabla._tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    for b_name in ['top', 'bottom', 'insideH']:
        bdr = OxmlElement(f'w:{b_name}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '4') 
        bdr.set(qn('w:space'), '0')
        bdr.set(qn('w:color'), 'auto')
        tblBorders.append(bdr)
        
    for b_name in ['left', 'right', 'insideV']:
        bdr = OxmlElement(f'w:{b_name}')
        bdr.set(qn('w:val'), 'none')
        tblBorders.append(bdr)
        
    tblPr.append(tblBorders)
    
    # ---------------------------------------------------------
    # DISEÑO FIJO: Aseguramos los 15 cm exactos sin que Word se autoexpanda
    # ---------------------------------------------------------
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    # Centramos la tabla. Al medir exactamente 15 cm (el ancho de la página), 
    # se alineará perfecta y simétricamente con los márgenes sin ese desfase a la izquierda.
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Anchos calculados: Reducimos Descripción y ampliamos Cantidad (Total = 15.0 cm)
    anchos = [Cm(3.0), Cm(1.5), Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.25), Cm(2.25)]
    
    for i, header in enumerate(headers):
        celda = tabla.rows[0].cells[i]
        celda.width = anchos[i]
        centrar_verticalmente_celda(celda) 
        
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_h = p.add_run(header)
        run_h.bold = True
        run_h.font.size = Pt(8)
        run_h.font.name = 'Arial'
        
    total_soles = 0.0
    total_dolares = 0.0
    
    for idx, item in enumerate(ce_data):
        row_cells = tabla.add_row().cells
        for i, llave in enumerate(llaves):
            celda = row_cells[i]
            celda.width = anchos[i]
            centrar_verticalmente_celda(celda) 
            
            if llave == "unidad":
                valor = item.get("unidad", "horas") 
            elif llave == "cantidad_real":
                valor = item.get("horas", item.get("cantidad", ""))
            else:
                valor = item.get(llave, "")
                
            texto_celda = ""
            if valor != "" and valor is not None:
                if llave == "precio_soles": texto_celda = f"S/ {float(valor):,.3f}"
                elif llave == "monto_soles": texto_celda = f"S/ {float(valor):,.3f}"
                elif llave == "monto_dolares": texto_celda = f"US$ {float(valor):,.3f}"
                elif llave == "factor_ajuste": texto_celda = f"{float(valor):,.3f}"
                elif llave == "cantidad_real":
                    val_float = float(valor)
                    texto_celda = f"{int(val_float)}" if val_float.is_integer() else f"{val_float:,.2f}"
                else:
                    texto_celda = str(valor)
                
            p = celda.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            
            if i == 0:
                # INYECTAMOS EL " 1/" Y " 2/" AUTOMÁTICAMENTE
                texto_celda += f" {idx + 1}/" 
                
                formatear_celda_descripcion(p, texto_celda)
                for r in p.runs:
                    r.font.size = Pt(8)
            else:
                run_c = p.add_run(texto_celda)
                run_c.font.size = Pt(8)
                run_c.font.name = 'Arial'
            
        total_soles += item.get("monto_soles", 0.0)
        total_dolares += item.get("monto_dolares", 0.0)

    # ==================================
    # FILA TOTAL (Celdas Combinadas)
    # ==================================
    row_total = tabla.add_row().cells
    
    # Asignamos anchos antes de combinar para que Word no se pierda
    row_total[0].width = anchos[0]
    row_total[1].width = anchos[1]
    row_total[2].width = anchos[2]
    row_total[3].width = anchos[3]
    row_total[4].width = anchos[4]
    
    # Combinamos
    row_total[0].merge(row_total[1])
    row_total[0].merge(row_total[2])
    row_total[0].merge(row_total[3])
    row_total[0].merge(row_total[4])
    
    # Reforzamos el ancho de las celdas combinadas (3.5 + 1.5 + 1.5 + 2.0 + 2.0 = 10.5 cm)
    row_total[0].width = Cm(10.5)
    row_total[5].width = Cm(2.25)
    row_total[6].width = Cm(2.25)
    
    p_tot = row_total[0].paragraphs[0]
    p_tot.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_tot = p_tot.add_run("TOTAL")
    run_tot.bold = True
    run_tot.font.size = Pt(8)
    run_tot.font.name = 'Arial'
    
    p_ts = row_total[5].paragraphs[0]
    p_ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ts = p_ts.add_run(f"S/ {total_soles:,.3f}")
    run_ts.bold = True
    run_ts.font.size = Pt(8)
    run_ts.font.name = 'Arial'
    
    p_td = row_total[6].paragraphs[0]
    p_td.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_td = p_td.add_run(f"US$ {total_dolares:,.3f}")
    run_td.bold = True
    run_td.font.size = Pt(8)
    run_td.font.name = 'Arial'